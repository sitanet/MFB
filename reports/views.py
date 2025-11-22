import logging
from django.shortcuts import get_object_or_404, render
# from docx import Document
from django.template.loader import render_to_string
# Create your views here.
from django.http import HttpResponse
from django.template.loader import get_template
from django.template import Context
# from weasyprint import HTML
from accounts.views import check_role_admin
from accounts_admin.models import Account, Account_Officer
from company.models import Company, Branch
from django.db.models import Sum
from loans.models import Loans
from profit_solutions.settings import BASE_DIR
from reports.forms import StatementForm
from transactions.models import Memtrans
# from docxtpl import DocxTemplate
from io import BytesIO
import tempfile
import os
from django.contrib.auth.decorators import login_required, user_passes_test
# import cairo
from django.views import View


def generate_pdf(request, id):
    customer = get_object_or_404(Loans, id=id)

    # Access the related Customer instance
    customers = customer.customer
    cust_data = Account.objects.filter(gl_no__startswith='20').exclude(gl_no='20100').exclude(
        gl_no='20200').exclude(gl_no='20000')
    gl_no = Account.objects.all().values_list('gl_no', flat=True).filter(gl_no__startswith='200')

    # Filter Memtrans records based on the ac_no of the customer
    ac_no_list = Memtrans.objects.filter(ac_no=customer.ac_no).values_list('ac_no', flat=True).distinct()
    cust_branch = Company.objects.first()
    # Calculate the total amount for each ac_no
    amounts = Memtrans.objects.filter(ac_no=customer.ac_no).values('gl_no').annotate(total_amount=Sum('amount'))
    loan_schedule = customer.calculate_loan_schedule()
    officer = Account_Officer.objects.all()
    total_interest_sum = sum(payment['interest_payment'] for payment in loan_schedule)
    total_principal_sum = sum(payment['principal_payment'] for payment in loan_schedule)
    total_payments_sum = sum(payment['total_payment'] for payment in loan_schedule)

    context = {
        'customers': customers,
        'customer': customer,
        'cust_data': cust_data,
        'cust_branch': cust_branch,
        'gl_no': gl_no,
        'officer': officer,
        'ac_no_list': ac_no_list,
        'amounts': amounts,
        'loan_schedule': loan_schedule,
        'total_interest_sum': total_interest_sum,
        'total_principal_sum': total_principal_sum,
        'total_payments_sum': total_payments_sum,
        }  # Replace with your data

    html_string = render(request, 'reports/loans/generate_pdf.html', context).content.decode('utf-8')
    pdf_file = HTML(string=html_string).write_pdf()

    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = 'filename="output.pdf"'

    return response



    
from django.http import HttpResponse
from docx import Document
from django.shortcuts import render, get_object_or_404
from django.template.loader import render_to_string
import os

def generate_doc(request, id):
    customer = get_object_or_404(Loans, id=id)

    # Access the related Customer instance
    customers = customer.customer
    cust_data = Account.objects.filter(gl_no__startswith='20').exclude(gl_no='20100').exclude(
        gl_no='20200').exclude(gl_no='20000')
    gl_no = Account.objects.all().values_list('gl_no', flat=True).filter(gl_no__startswith='200')

    # Filter Memtrans records based on the ac_no of the customer
    ac_no_list = Memtrans.objects.filter(ac_no=customer.ac_no).values_list('ac_no', flat=True).distinct()
    cust_branch = Company.objects.all()
    # Calculate the total amount for each ac_no
    amounts = Memtrans.objects.filter(ac_no=customer.ac_no).values('gl_no').annotate(total_amount=Sum('amount'))
    loan_schedule = customer.calculate_loan_schedule()
    officer = Account_Officer.objects.all()
    total_interest_sum = sum(payment['interest_payment'] for payment in loan_schedule)
    total_principal_sum = sum(payment['principal_payment'] for payment in loan_schedule)
    total_payments_sum = sum(payment['total_payment'] for payment in loan_schedule)

    context = {
        'customers': customers,
        'customer': customer,
        'cust_data': cust_data,
        'cust_branch': cust_branch,
        'gl_no': gl_no,
        'officer': officer,
        'ac_no_list': ac_no_list,
        'amounts': amounts,
        'loan_schedule': loan_schedule,
        'total_interest_sum': total_interest_sum,
        'total_principal_sum': total_principal_sum,
        'total_payments_sum': total_payments_sum,
    }  # Replace with your data

    # Render HTML content from the template
    html_string = render_to_string('reports/loans/generate_doc.html', context)

    # Create a new Document
    document = Document()

    # Add paragraphs to the document
    for paragraph in html_string.split('<p>'):
        document.add_paragraph(paragraph.strip('</p>'))

    # Save the document to a response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'filename="output.docx"'
    document.save(response)

    return response




from openpyxl import Workbook

from django.http import HttpResponse
from io import BytesIO
from django.template.loader import render_to_string

def generate_excel(request, id):
    customer = get_object_or_404(Loans, id=id)

    # Access the related Customer instance
    customers = customer.customer
    cust_data = Account.objects.filter(gl_no__startswith='20').exclude(gl_no='20100').exclude(
        gl_no='20200').exclude(gl_no='20000')
    gl_no = Account.objects.all().values_list('gl_no', flat=True).filter(gl_no__startswith='200')

    # Filter Memtrans records based on the ac_no of the customer
    ac_no_list = Memtrans.objects.filter(ac_no=customer.ac_no).values_list('ac_no', flat=True).distinct()
    cust_branch = Company.objects.all()
    # Calculate the total amount for each ac_no
    amounts = Memtrans.objects.filter(ac_no=customer.ac_no).values('gl_no').annotate(total_amount=Sum('amount'))
    loan_schedule = customer.calculate_loan_schedule()
    officer = Account_Officer.objects.all()
    total_interest_sum = sum(payment['interest_payment'] for payment in loan_schedule)
    total_principal_sum = sum(payment['principal_payment'] for payment in loan_schedule)
    total_payments_sum = sum(payment['total_payment'] for payment in loan_schedule)

    context = {
        'customers': customers,
        'customer': customer,
        'cust_data': cust_data,
        'cust_branch': cust_branch,
        'gl_no': gl_no,
        'officer': officer,
        'ac_no_list': ac_no_list,
        'amounts': amounts,
        'loan_schedule': loan_schedule,
        'total_interest_sum': total_interest_sum,
        'total_principal_sum': total_principal_sum,
        'total_payments_sum': total_payments_sum,
    }  # Replace with your data

  

#     # Render HTML content from the template with the context
    html_string = render_to_string('reports/loans/generate_excel.html', context)

    # Create a new Workbook
    workbook = Workbook()

    # Get the active sheet
    sheet = workbook.active

    # Add headings to the sheet
    sheet.append(['Period', 'Payment Date', 'Principal Payment', 'Interest Payment', 'Total Payment', 'Principal Remaining'])

    # Add data rows to the sheet
    for payment in loan_schedule:
        sheet.append([
            payment['period'],
            payment['payment_date'],
            payment['principal_payment'],
            payment['interest_payment'],
            payment['total_payment'],
            payment['principal_remaining'],
        ])

    # Save the workbook to a BytesIO buffer
    buffer = BytesIO()
    workbook.save(buffer)
    buffer.seek(0)

    # Save the workbook to a response
    response = HttpResponse(buffer.read(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=loan_schedule.xlsx'

    return response







import csv
from django.http import HttpResponse
from io import StringIO

def generate_csv(request, id):
    customer = get_object_or_404(Loans, id=id)

    # Access the related Customer instance
    customers = customer.customer
    cust_data = Account.objects.filter(gl_no__startswith='20').exclude(gl_no='20100').exclude(
        gl_no='20200').exclude(gl_no='20000')
    gl_no = Account.objects.all().values_list('gl_no', flat=True).filter(gl_no__startswith='200')

    # Filter Memtrans records based on the ac_no of the customer
    ac_no_list = Memtrans.objects.filter(ac_no=customer.ac_no).values_list('ac_no', flat=True).distinct()
    cust_branch = Company.objects.all()
    # Calculate the total amount for each ac_no
    amounts = Memtrans.objects.filter(ac_no=customer.ac_no).values('gl_no').annotate(total_amount=Sum('amount'))
    loan_schedule = customer.calculate_loan_schedule()
    officer = Account_Officer.objects.all()
    total_interest_sum = sum(payment['interest_payment'] for payment in loan_schedule)
    total_principal_sum = sum(payment['principal_payment'] for payment in loan_schedule)
    total_payments_sum = sum(payment['total_payment'] for payment in loan_schedule)

    context = {
        'customers': customers,
        'customer': customer,
        'cust_data': cust_data,
        'cust_branch': cust_branch,
        'gl_no': gl_no,
        'officer': officer,
        'ac_no_list': ac_no_list,
        'amounts': amounts,
        'loan_schedule': loan_schedule,
        'total_interest_sum': total_interest_sum,
        'total_principal_sum': total_principal_sum,
        'total_payments_sum': total_payments_sum,
    }  # Replace with your data

  
    csv_content = render_to_string('reports/loans/generate_csv.html', {'loan_schedule': loan_schedule})
     # Create a CSV response
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename=loan_schedule.csv'

    # Create a CSV writer
    csv_writer = csv.writer(response)

    # Write header row
    csv_writer.writerow(['Period', 'Payment Date', 'Principal Payment', 'Interest Payment', 'Total Payment', 'Principal Remaining'])

    # Write data rows
    for payment in loan_schedule:
        csv_writer.writerow([
            payment['period'],
            payment['payment_date'],
            payment['principal_payment'],
            payment['interest_payment'],
            payment['total_payment'],
            payment['principal_remaining'],
        ])
    response.write(csv_content)

    return response




# views.py
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.shortcuts import render, redirect
from django.urls import reverse



from django.shortcuts import render





from django.shortcuts import render
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage



def generate_statement_data(gl_no, ac_no, start_date, end_date):
    transactions = Memtrans.objects.filter(
        gl_no=gl_no,
        ac_no=ac_no,
        ses_date__range=(start_date, end_date)
    ).order_by('ses_date')

    statement_data = []
    running_balance = 0
    total_debit = 0
    total_credit = 0
    opening_balance = 0
    closing_balance = 0

    for transaction in transactions:
        if transaction.ses_date < start_date:
            # Accumulate transactions before the start date for opening balance
            if transaction.amount > 0:
                opening_balance += transaction.amount
            else:
                opening_balance -= abs(transaction.amount)

    # Include a single entry for opening balance in the running balance
    running_balance += opening_balance

    # Include a separate entry for opening balance in the statement data
    statement_data.append({
        'date': start_date,
        'description': 'Opening Balance',
        'trx_no': '',
        'debit': 0,
        'credit': 0,
        'running_balance': running_balance,
    })

    for transaction in transactions:
        if transaction.ses_date >= start_date:
            if transaction.amount > 0:
                credit_amount = transaction.amount
                debit_amount = 0
            else:
                credit_amount = 0
                debit_amount = abs(transaction.amount)

            running_balance += (credit_amount - debit_amount)
            total_debit += debit_amount
            total_credit += credit_amount

            description = transaction.description or ''
            statement_data.append({
                'date': transaction.ses_date,
                'description': description,
                'trx_no': transaction.trx_no,
                'debit': debit_amount,
                'credit': credit_amount,
                'running_balance': running_balance,
            })

    closing_balance = running_balance


# views.py

@login_required(login_url='login')
@user_passes_test(check_role_admin)
def front_office_report(request):
   
    return render(request, 'reports/front_office_report.html')


@login_required(login_url='login')
@user_passes_test(check_role_admin)
def back_office_report(request):
   
    return render(request, 'reports/back_office_report.html')


@login_required(login_url='login')
@user_passes_test(check_role_admin)
def financial_report(request):
   
    return render(request, 'reports/financials/financial_report.html')



@login_required(login_url='login')
@user_passes_test(check_role_admin)
def savings_report(request):
   
    return render(request, 'reports/savings/savings_report.html')

from django.shortcuts import render
from django.db.models import Sum
from .forms import LoanOutstandingBalanceForm, StatementForm

from django.db.models import Q

from django.db.models import F
from django.db.models.functions import Now

# Additional imports for export functionality
import csv
import io
from datetime import datetime
from django.http import HttpResponse
from django.template.loader import render_to_string
from django.template.loader import get_template
from xhtml2pdf import pisa
from django.http import HttpResponse

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import tempfile
import os
# ...
@login_required(login_url='login')
@user_passes_test(check_role_admin)
def generate_statement_view(request):
    # ✅ Always use the logged-in user's branch
    branch = request.user.branch  
    branch_date = branch.session_date.strftime('%Y-%m-%d') if branch.session_date else ''

    if branch.session_status == 'Closed':
        return HttpResponse("You can not post any transaction. Session is closed.") 
    
    # Handle export requests
    export_format = request.GET.get('export')
    if export_format and request.GET.get('start_date'):
        return handle_statement_export(request, export_format)
    
    if request.method == 'POST':
        form = StatementForm(request.POST)
        if form.is_valid():
            start_date = form.cleaned_data['start_date']
            end_date = form.cleaned_data['end_date']
            gl_no = form.cleaned_data['gl_no']
            ac_no = form.cleaned_data['ac_no']

            # Retrieve transactions within the specified date range
            transactions = Memtrans.objects.filter(
                ses_date__range=[start_date, end_date],
                gl_no=gl_no,
                ac_no=ac_no
            ).exclude(error='H').order_by('ses_date', 'trx_no').annotate(
                current_time=Now()
            )

            # Opening balance
            opening_balance = Memtrans.objects.filter(
                ses_date__lt=start_date,
                gl_no=gl_no,
                ac_no=ac_no
            ).exclude(error='H').aggregate(
                opening_balance=Sum('amount')
            )['opening_balance'] or 0

            # Closing balance
            closing_balance = Memtrans.objects.filter(
                ses_date__lte=end_date,
                gl_no=gl_no,
                ac_no=ac_no
            ).exclude(error='H').aggregate(
                closing_balance=Sum('amount')
            )['closing_balance'] or 0

            # Debit & Credit
            debit_amount = transactions.filter(type='D').aggregate(
                debit_amount=Sum('amount')
            )['debit_amount'] or 0

            credit_amount = transactions.filter(type='C').aggregate(
                credit_amount=Sum('amount')
            )['credit_amount'] or 0

            # Customer name
            first_transaction = transactions.first()
            full_name = (
                first_transaction.customer.get_full_name()
                if first_transaction and first_transaction.customer
                else ''
            )

            # Statement details
            statement_data = []
            running_balance = opening_balance
            for transaction in transactions:
                running_balance += transaction.amount if transaction.type == 'D' else -transaction.amount
                entry = {
                    'date': transaction.ses_date,
                    'trx_no': transaction.trx_no,
                    'description': transaction.description,
                    'debit': transaction.amount if transaction.type == 'D' else 0,
                    'credit': transaction.amount if transaction.type == 'C' else 0,
                    'running_balance': running_balance,
                }
                statement_data.append(entry)

            context = {
                'start_date': start_date,
                'end_date': end_date,
                'gl_no': gl_no,
                'ac_no': ac_no,
                'opening_balance': opening_balance,
                'closing_balance': closing_balance,
                'debit_amount': debit_amount,
                'credit_amount': credit_amount,
                'statement_data': statement_data,
                'form': form,
                'full_name': full_name,
                'branch': branch,
                'company': branch.company,  # ✅ Branch must be linked to company
                'branch_date': branch_date,
            }
            return render(request, 'reports/accounts/statement_of_account.html', context)

    else:
        form = StatementForm()

    return render(request, 'reports/accounts/input_form.html', {
        'form': form,
        'branch': branch,
        'company': branch.company,
        'branch_date': branch_date
    })

from django.shortcuts import render
from transactions.models import Memtrans, Account
from datetime import datetime
from collections import defaultdict
from .forms import TrialBalanceForm


def generate_trial_balance(start_date, end_date, branch_code=None):
    """
    Generate a trial balance report between the specified start and end dates for a given branch code.

    Args:
        start_date (date): The start date for the report.
        end_date (date): The end date for the report.
        branch_code (str or None): The branch code to filter the transactions, or None to include all branches.

    Returns:
        tuple: Contains sorted trial balance data, subtotals for different GL number prefixes, and totals for debit, credit, and balance.
    """
    # Query all relevant Memtrans entries for the given branch
    memtrans_entries = Memtrans.objects.filter(ses_date__range=[start_date, end_date], error='A')

    # Apply branch filtering only if a specific branch code is provided
    if branch_code:
        memtrans_entries = memtrans_entries.filter(branch=branch_code)

    # Initialize a dictionary to hold GL account balances
    gl_customer_balance = defaultdict(lambda: {'debit': 0, 'credit': 0, 'balance': 0})

    # Process each entry
    for entry in memtrans_entries:
        gl_no = entry.gl_no
        amount = entry.amount

        if entry.type == 'N':
            if amount < 0:
                gl_customer_balance[gl_no]['debit'] += abs(amount)
            else:
                gl_customer_balance[gl_no]['credit'] += amount
        else:
            if amount < 0:
                gl_customer_balance[gl_no]['credit'] += abs(amount)
            else:
                gl_customer_balance[gl_no]['debit'] += amount

    # Fetch all necessary accounts in a single query to improve performance
    accounts = Account.objects.filter(gl_no__in=gl_customer_balance.keys()).values('gl_no', 'gl_name')
    account_map = {account['gl_no']: account['gl_name'] for account in accounts}

    # Sort GL numbers and prepare trial balance data
    sorted_keys = sorted(gl_customer_balance.keys())
    sorted_trial_balance_data = []
    subtotal_1 = subtotal_2 = subtotal_3 = subtotal_4 = subtotal_5 = 0

    for gl_no in sorted_keys:
        balance_data = gl_customer_balance[gl_no]
        debit = balance_data['debit']
        credit = balance_data['credit']
        balance = debit - credit
        gl_name = account_map.get(gl_no, '')  # Default to empty string if GL number not found
        balance_data.update({'gl_no': gl_no, 'gl_name': gl_name, 'balance': balance})
        sorted_trial_balance_data.append(balance_data)

        # Calculate subtotals for GL numbers starting with specific prefixes
        if gl_no.startswith("1"):
            subtotal_1 += balance
        elif gl_no.startswith("2"):
            subtotal_2 += balance
        elif gl_no.startswith("3"):
            subtotal_3 += balance
        elif gl_no.startswith("4"):
            subtotal_4 += balance
        elif gl_no.startswith("5"):
            subtotal_5 += balance

    # Calculate total debit, credit, and balance
    total_debit = sum(entry['debit'] for entry in sorted_trial_balance_data)
    total_credit = sum(entry['credit'] for entry in sorted_trial_balance_data)
    total_balance = total_debit - total_credit

    return sorted_trial_balance_data, subtotal_1, subtotal_2, subtotal_3, subtotal_4, subtotal_5, total_debit, total_credit, total_balance


def trial_balance(request):
    user_branch = request.user.branch  # assuming User has a branch FK
    branches = Branch.objects.all() if user_branch.head_office else Branch.objects.filter(id=user_branch.id)
    head_office = user_branch.head_office 
    if request.method == 'POST':
        form = TrialBalanceForm(request.POST)
        form.fields['branch'].queryset = branches  # ✅ limit branch choices

        if form.is_valid():
            start_date = form.cleaned_data['start_date']
            end_date = form.cleaned_data['end_date']
            branch = form.cleaned_data['branch']

            trial_balance_data, subtotal_1, subtotal_2, subtotal_3, subtotal_4, subtotal_5, total_debit, total_credit, total_balance = generate_trial_balance(
                start_date, end_date, branch
            )

            if not trial_balance_data:
                messages.warning(request, "No trial balance data found for the selected branch and date range.")

            return render(request, 'reports/financials/trial_balance.html', {
                'form': form,
                'branches': branches,
                'trial_balance_data': trial_balance_data,
                'subtotal_1': subtotal_1,
                'subtotal_2': subtotal_2,
                'subtotal_3': subtotal_3,
                'subtotal_4': subtotal_4,
                'subtotal_5': subtotal_5,
                'total_debit': total_debit,
                'total_credit': total_credit,
                'total_balance': total_balance,
                'start_date': start_date,
                'end_date': end_date,
                'selected_branch': branch.id if branch else None,
                'branch': branch,
            })
    else:
        form = TrialBalanceForm()
        form.fields['branch'].queryset = branches  

    return render(request, 'reports/financials/trial_balance.html', {
        'form': form,
        'head_office': head_office,
        'branches': branches,
        'selected_branch': None,
    })


from django.shortcuts import render
from django.contrib import messages
from django.db.models import Subquery, OuterRef, Value, CharField, Sum
from django.db.models.functions import Concat
from django.utils import timezone
from django.contrib.auth.decorators import login_required

from transactions.models import Memtrans
from company.models import Branch
from customers.models import Customer
from accounts.models import User
from .forms import TransactionForm


@login_required
def transaction_sequence_by_ses_date(request):
    current_datetime = timezone.now()
    report_data = None
    selected_branch = None
    start_date = None
    end_date = None
    session_date = None
    total_debit = 0
    total_credit = 0

    user_branch = request.user.branch
    head_office = user_branch.head_office 

    # Determine branches user can access
    if request.user.branch.head_office:
        # Head office: can see all branches in the company
        branches = Branch.objects.filter(company=request.user.branch.company)
    else:
        # Regular branch user: only their branch
        branches = Branch.objects.filter(id=request.user.branch.id)

    # ✅ Users should always be from all branches in the form
    users = User.objects.filter(branch__in=branches).distinct()

    # Build the form
    form = TransactionForm(request.POST or None, branches=branches, users=users)

    if request.method == 'POST' and form.is_valid():
        start_date = form.cleaned_data['start_date']
        end_date = form.cleaned_data['end_date']
        branch_obj = form.cleaned_data['branch']
        user_obj = form.cleaned_data['user']
        code = form.cleaned_data['code']

        # Base queryset
        report_data = Memtrans.objects.filter(
            ses_date__range=(start_date, end_date)
        ).select_related('user', 'branch').order_by('ses_date')

        # Branch filtering
        if branch_obj:
            if not branch_obj.head_office:  
                selected_branch = branch_obj
                report_data = report_data.filter(branch=selected_branch)
                session_date = selected_branch.session_date
        else:
            selected_branch = None
            session_date = None

        # User filtering
        if user_obj:
            report_data = report_data.filter(user_id=user_obj.id)

        # Code filtering
        if code:
            report_data = report_data.filter(code=code)

        # Annotate customer full name
        report_data = report_data.annotate(
            customer_name=Subquery(
                Customer.objects.filter(
                    gl_no=OuterRef('gl_no'),
                    ac_no=OuterRef('ac_no')
                ).annotate(
                    full_name=Concat(
                        'first_name', Value(' '),
                        'middle_name', Value(' '),
                        'last_name',
                        output_field=CharField()
                    )
                ).values('full_name')[:1]
            )
        )

        # Totals
        total_debit = report_data.filter(amount__lt=0).aggregate(total_debit=Sum('amount'))['total_debit'] or 0
        total_credit = report_data.filter(amount__gte=0).aggregate(total_credit=Sum('amount'))['total_credit'] or 0

        if not report_data.exists():
            messages.warning(request, "No transactions found for the selected criteria.")

    return render(request, 'reports/savings/transaction_sequence_by_ses_date.html', {
        'form': form,
        'head_office':head_office,
        'start_date': start_date,
        'end_date': end_date,
        'report_data': report_data,
        'branches': branches,
        'users': users,
        'selected_branch': selected_branch,
        'session_date': session_date,
        'total_debit': total_debit,
        'total_credit': total_credit,
        'current_datetime': current_datetime,
    })


@login_required
def transaction_sequence_by_trx_date(request):
    user_branch = request.user.branch
    head_office = bool(user_branch.head_office)
    print("HEAD OFFICE FLAG:", head_office, type(head_office))
    current_datetime = timezone.now()
    report_data = None
    selected_branch = None
    start_date = None
    end_date = None
    session_date = None
    total_debit = 0
    total_credit = 0


    # Determine branches and users based on head office vs branch
    # Determine branches and users based on head office vs branch
    if request.user.branch.head_office:
        # Head office user: can see all branches
        branches = Branch.objects.filter(company=request.user.branch.company)

        # All users assigned to these branches
        users = User.objects.filter(branch__company=request.user.branch.company).distinct()
    else:
        # Regular branch user: only their own branch
        branches = Branch.objects.filter(id=request.user.branch.id)
        users = User.objects.filter(branch=request.user.branch)





    form = TransactionForm(request.POST or None, branches=branches, users=users)

    if request.method == 'POST' and form.is_valid():
        start_date = form.cleaned_data['start_date']
        end_date = form.cleaned_data['end_date']
        branch_obj = form.cleaned_data['branch']
        user_obj = form.cleaned_data['user']
        code = form.cleaned_data['code']

        report_data = Memtrans.objects.filter(
            sys_date__range=(start_date, end_date)
        ).select_related('user', 'branch').order_by('sys_date')

        # Branch filtering
        if branch_obj:
            if not branch_obj.head_office:
                selected_branch = branch_obj
                report_data = report_data.filter(branch=selected_branch)
                session_date = selected_branch.session_date
        else:
            selected_branch = None
            session_date = None

        # User filtering
        if user_obj:
            report_data = report_data.filter(user_id=user_obj.id)

        # Code filtering
        if code:
            report_data = report_data.filter(code=code)

        # Annotate customer name
        report_data = report_data.annotate(
            customer_name=Subquery(
                Customer.objects.filter(
                    gl_no=OuterRef('gl_no'),
                    ac_no=OuterRef('ac_no')
                ).annotate(
                    full_name=Concat(
                        'first_name', Value(' '),
                        'middle_name', Value(' '),
                        'last_name',
                        output_field=CharField()
                    )
                ).values('full_name')[:1]
            )
        )
        

        # Totals
        total_debit = report_data.filter(amount__lt=0).aggregate(total_debit=Sum('amount'))['total_debit'] or 0
        total_credit = report_data.filter(amount__gte=0).aggregate(total_credit=Sum('amount'))['total_credit'] or 0

        if not report_data.exists():
            messages.warning(request, "No transactions found for the selected criteria.")

    return render(request, 'reports/savings/transaction_sequence_by_trx_date.html', {
        'form': form,
        'head_office':head_office,
        'start_date': start_date,
        'end_date': end_date,
        'report_data': report_data,
        'branches': branches,
        'users': users,
        'selected_branch': selected_branch,
        'session_date': session_date,
        'total_debit': total_debit,
        'total_credit': total_credit,
        'current_datetime': current_datetime,
    })



from accounts.models import User
from django.contrib.auth.decorators import login_required
from django.db.models import Sum, Subquery, OuterRef, Value, CharField
from django.db.models.functions import Concat
from django.utils import timezone
from django.shortcuts import render, get_object_or_404
from django.contrib import messages

from transactions.models import Memtrans
from company.models import Branch
from customers.models import Customer
from .forms import TransactionForm


@login_required
def transaction_journal_listing_by_ses_date(request):
    current_datetime = timezone.now()
    report_data = None
    selected_branch = None
    start_date = None
    end_date = None
    session_date = None
    total_debit = 0
    total_credit = 0

    # ✅ Determine accessible branches
# ✅ Determine accessible branches
# ✅ Determine accessible branches
    if request.user.branch.head_office:
        # Head office user: see all company branches
        branches = Branch.objects.filter(company=request.user.branch.company)
        head_office = True
    else:
        # Regular branch user: see only their branch
        branches = Branch.objects.filter(id=request.user.branch.id)
        head_office = False

    # ✅ All users belonging to those branches
    users = User.objects.filter(branch__in=branches).distinct()

    # ✅ Build form with branch/user options
    form = TransactionForm(request.POST or None, branches=branches, users=users)


    if request.method == 'POST' and form.is_valid():
        start_date = form.cleaned_data['start_date']
        end_date = form.cleaned_data['end_date']
        branch_obj = form.cleaned_data['branch']
        user_obj = form.cleaned_data['user']
        code = form.cleaned_data['code']

        # Base queryset
        report_data = Memtrans.objects.all()

        # Date filtering
        if start_date and end_date:
            report_data = report_data.filter(ses_date__range=(start_date, end_date))

        # Branch filtering
        if branch_obj:
            selected_branch = branch_obj
            report_data = report_data.filter(branch=selected_branch)
            session_date = selected_branch.session_date

        # User filtering
        if user_obj:
            report_data = report_data.filter(user_id=user_obj.id)

        # Code filtering
        if code:
            report_data = report_data.filter(code=code)

        # Optimize and order
        report_data = report_data.select_related('user', 'branch').order_by('ses_date')

        # Annotate customer full name
        report_data = report_data.annotate(
            customer_name=Subquery(
                Customer.objects.filter(
                    gl_no=OuterRef('gl_no'),
                    ac_no=OuterRef('ac_no')
                ).annotate(
                    full_name=Concat(
                        'first_name', Value(' '),
                        'middle_name', Value(' '),
                        'last_name',
                        output_field=CharField()
                    )
                ).values('full_name')[:1]
            )
        )

        # Totals
        total_debit = report_data.filter(type='D').aggregate(total=Sum('amount'))['total'] or 0
        total_credit = report_data.filter(type='C').aggregate(total=Sum('amount'))['total'] or 0

        # ✅ Show message if no data
        if not report_data.exists():
            messages.warning(request, "No transactions found for the selected criteria.")

    return render(request, 'reports/savings/transaction_journal_listing_by_ses_date.html', {
        'form': form,
        'head_office': head_office, 
        'start_date': start_date,
        'end_date': end_date,
        'report_data': report_data,
        'branches': branches,
        'users': users,
        'selected_branch': selected_branch,
        'session_date': session_date,
        'total_debit': total_debit,
        'total_credit': total_credit,
        'current_datetime': current_datetime,
    })


from django.shortcuts import render, get_object_or_404
from django.utils import timezone
from django.db.models import Sum, Q, Value, OuterRef, Subquery, CharField
from django.db.models.functions import Concat
from django.contrib import messages
from django.contrib.auth.decorators import login_required

from .forms import TransactionForm
from transactions.models import Memtrans
from company.models import Branch
from customers.models import Customer
from accounts.models import User   # Use your custom User model


@login_required
@login_required
def transaction_journal_listing_by_trx_date(request):
    current_datetime = timezone.now()
    report_data = None
    selected_branch = None
    start_date = None
    end_date = None
    session_date = None
    total_debit = 0
    total_credit = 0

    # ✅ Determine accessible branches
    if request.user.branch.head_office:  # <-- correct field
        # Head office user: see all company branches
        branches = Branch.objects.filter(company=request.user.branch.company)
        head_office = True
    else:
        # Regular branch user: only their branch
        branches = Branch.objects.filter(id=request.user.branch.id)
        head_office = False

    # ✅ All users belonging to those branches
    users = User.objects.filter(branch__in=branches).distinct()

    # Pass branches & users to form
    form = TransactionForm(request.POST or None, branches=branches, users=users)

    if request.method == 'POST' and form.is_valid():
        start_date = form.cleaned_data['start_date']
        end_date = form.cleaned_data['end_date']
        branch_obj = form.cleaned_data['branch']
        user_obj = form.cleaned_data['user']
        code = form.cleaned_data['code']

        # Base queryset
        report_data = Memtrans.objects.select_related('user', 'branch').only(
            'app_date', 'trx_no', 'code', 'gl_no', 'ac_no',
            'amount', 'user__username', 'branch__branch_name'
        )

        # Date filtering
        if start_date and end_date:
            report_data = report_data.filter(app_date__range=(start_date, end_date))

        # Branch filtering
        if branch_obj:
            selected_branch = branch_obj
            report_data = report_data.filter(branch=selected_branch)
            session_date = selected_branch.session_date

        # User filtering
        if user_obj:
            report_data = report_data.filter(user_id=user_obj.id)

        # Code filtering
        if code:
            report_data = report_data.filter(code=code)

        # Annotate customer full name
        report_data = report_data.annotate(
            customer_name=Subquery(
                Customer.objects.filter(
                    gl_no=OuterRef('gl_no'),
                    ac_no=OuterRef('ac_no')
                ).annotate(
                    full_name=Concat(
                        'first_name', Value(' '),
                        'middle_name', Value(' '),
                        'last_name',
                        output_field=CharField()
                    )
                ).values('full_name')[:1],
                output_field=CharField()
            )
        ).order_by('app_date')

        # Totals
        totals = report_data.aggregate(
            total_debit=Sum('amount', filter=Q(amount__lt=0)),
            total_credit=Sum('amount', filter=Q(amount__gte=0))
        )

        total_debit = abs(totals['total_debit'] or 0)
        total_credit = totals['total_credit'] or 0

        # ✅ Show warning if no data
        if not report_data.exists():
            messages.warning(request, "No transactions found for the selected criteria.")

    return render(request, 'reports/savings/transaction_journal_listing_by_trx_date.html', {
        'form': form,
        'branches': branches,
        'users': users,
        'report_data': report_data,
        'selected_branch': selected_branch,
        'start_date': start_date,
        'end_date': end_date,
        'session_date': session_date,
        'total_debit': total_debit,
        'total_credit': total_credit,
        'current_datetime': current_datetime,
        'head_office': head_office,  # ✅ useful for template condition
    })


from django.shortcuts import render, get_object_or_404
from django.utils import timezone
from django.db.models import Sum, Q, Value, OuterRef, Subquery, CharField
from django.db.models.functions import Concat
from django.contrib import messages
from django.contrib.auth.decorators import login_required

from .forms import TransactionForm
from transactions.models import Memtrans
from company.models import Branch
from customers.models import Customer
from accounts.models import User   # ✅ custom User model


@login_required
@login_required
def transaction_day_sheet_by_session_date(request):
    current_datetime = timezone.now()
    report_data = None
    selected_branch = None
    start_date = None
    end_date = None
    session_date = None
    total_debit = 0
    total_credit = 0

    # ✅ Determine accessible branches
    if request.user.branch.head_office:  # <-- use head_office
        # Head office user: all branches under the company
        branches = Branch.objects.filter(company=request.user.branch.company)
        head_office = True
    else:
        # Regular branch user: only their branch
        branches = Branch.objects.filter(id=request.user.branch.id)
        head_office = False

    # ✅ All users from those branches
    users = User.objects.filter(branch__in=branches).distinct()

    # Pass into form
    form = TransactionForm(request.POST or None, branches=branches, users=users)

    if request.method == 'POST' and form.is_valid():
        start_date = form.cleaned_data['start_date']
        end_date = form.cleaned_data['end_date']
        branch_obj = form.cleaned_data['branch']
        user_obj = form.cleaned_data['user']
        code = form.cleaned_data['code']

        # Base queryset
        report_data = Memtrans.objects.filter(
            ses_date__range=(start_date, end_date)
        ).select_related('user', 'branch').order_by('ses_date', 'app_date')

        # Branch filtering
        if branch_obj:
            selected_branch = branch_obj
            report_data = report_data.filter(branch=selected_branch)
            session_date = selected_branch.session_date

        # User filtering
        if user_obj:
            report_data = report_data.filter(user_id=user_obj.id)

        # Code filtering
        if code:
            report_data = report_data.filter(code=code)

        # Annotate customer name
        report_data = report_data.annotate(
            customer_name=Subquery(
                Customer.objects.filter(
                    gl_no=OuterRef('gl_no'),
                    ac_no=OuterRef('ac_no')
                ).annotate(
                    full_name=Concat(
                        'first_name', Value(' '),
                        'middle_name', Value(' '),
                        'last_name',
                        output_field=CharField()
                    )
                ).values('full_name')[:1],
                output_field=CharField()
            )
        )

        # Totals
        totals = report_data.aggregate(
            total_debit=Sum('amount', filter=Q(type='D')),
            total_credit=Sum('amount', filter=Q(type='C'))
        )
        total_debit = totals['total_debit'] or 0
        total_credit = totals['total_credit'] or 0

        # ✅ Show warning if no records
        if not report_data.exists():
            messages.warning(request, "No transactions found for the selected criteria.")

    return render(request, 'reports/savings/transaction_day_sheet_by_session_date.html', {
        'form': form,
        'start_date': start_date,
        'end_date': end_date,
        'report_data': report_data,
        'branches': branches,
        'users': users,
        'selected_branch': selected_branch,
        'session_date': session_date,
        'total_debit': total_debit,
        'total_credit': total_credit,
        'current_datetime': current_datetime,
        'head_office': head_office,  # ✅ pass flag to template
    })


from django.shortcuts import render, get_object_or_404
from django.utils import timezone
from django.db.models import Sum, Q, Value, OuterRef, Subquery, CharField
from django.db.models.functions import Concat
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.contrib.auth import get_user_model

from .forms import TransactionForm
from transactions.models import Memtrans
from company.models import Branch
from customers.models import Customer

User = get_user_model()


@login_required
@login_required
def transaction_day_sheet_by_trx_date(request):
    current_datetime = timezone.now()
    report_data = None
    selected_branch = None
    start_date = None
    end_date = None
    session_date = None
    total_debit = 0
    total_credit = 0

    # ✅ Determine accessible branches
    if request.user.branch.head_office:  # <-- use head_office
        # Head office: all company branches
        branches = Branch.objects.filter(company=request.user.branch.company)
        head_office = True
    else:
        # Regular branch: only their branch
        branches = Branch.objects.filter(id=request.user.branch.id)
        head_office = False

    # ✅ Users from those branches
    users = User.objects.filter(branch__in=branches).distinct()

    # Build form with branch/user options
    form = TransactionForm(request.POST or None, branches=branches, users=users)

    if request.method == 'POST' and form.is_valid():
        start_date = form.cleaned_data['start_date']
        end_date = form.cleaned_data['end_date']
        branch_obj = form.cleaned_data['branch']
        user_obj = form.cleaned_data['user']
        code = form.cleaned_data['code']

        # Base queryset
        report_data = Memtrans.objects.filter(
            app_date__range=(start_date, end_date)
        ).select_related('user', 'branch').order_by('app_date', 'ses_date')

        # Branch filtering
        if branch_obj:
            selected_branch = branch_obj
            report_data = report_data.filter(branch=selected_branch)
            session_date = selected_branch.session_date

        # User filtering
        if user_obj:
            report_data = report_data.filter(user_id=user_obj.id)

        # Code filtering
        if code:
            report_data = report_data.filter(code=code)

        # Annotate customer name
        report_data = report_data.annotate(
            customer_name=Subquery(
                Customer.objects.filter(
                    gl_no=OuterRef('gl_no'),
                    ac_no=OuterRef('ac_no')
                ).annotate(
                    full_name=Concat(
                        'first_name', Value(' '),
                        'middle_name', Value(' '),
                        'last_name',
                        output_field=CharField()
                    )
                ).values('full_name')[:1],
                output_field=CharField()
            )
        )

        # Totals
        totals = report_data.aggregate(
            total_debit=Sum('amount', filter=Q(type='D')),
            total_credit=Sum('amount', filter=Q(type='C'))
        )
        total_debit = totals['total_debit'] or 0
        total_credit = totals['total_credit'] or 0

        # ✅ Show warning if no results
        if not report_data.exists():
            messages.warning(request, "No transactions found for the selected criteria.")

    return render(request, 'reports/savings/transaction_day_sheet_by_trx_date.html', {
        'form': form,
        'start_date': start_date,
        'end_date': end_date,
        'report_data': report_data,
        'branches': branches,
        'users': users,
        'selected_branch': selected_branch,
        'session_date': session_date,
        'total_debit': total_debit,
        'total_credit': total_credit,
        'current_datetime': current_datetime,
        'head_office': head_office,  # ✅ flag for template
    })



from django.shortcuts import render, get_object_or_404
from django.utils import timezone
from django.db.models import Sum, Q, Value, OuterRef, Subquery, CharField
from django.db.models.functions import Concat
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.contrib.auth import get_user_model

from .forms import TransactionForm
from transactions.models import Memtrans
from company.models import Branch
from customers.models import Customer

User = get_user_model()


@login_required
def general_trx_register_by_session_date(request):
    current_datetime = timezone.now()
    report_data = None
    selected_branch = None
    start_date = None
    end_date = None
    session_date = None
    total_debit = 0
    total_credit = 0

    # ✅ Determine accessible branches
    if request.user.branch.head_office:
        # Head office: all branches of the same company
        branches = Branch.objects.filter(company=request.user.branch.company)
    else:
        # Regular branch: only their branch
        branches = Branch.objects.filter(id=request.user.branch.id)

    # ✅ Users from those branches
    users = User.objects.filter(branch__in=branches).distinct()

    # Pass branches/users into the form
    form = TransactionForm(request.POST or None, branches=branches, users=users)

    if request.method == 'POST' and form.is_valid():
        start_date = form.cleaned_data['start_date']
        end_date = form.cleaned_data['end_date']
        branch_obj = form.cleaned_data['branch']
        user_obj = form.cleaned_data['user']
        code = form.cleaned_data['code']

        # Base queryset
        report_data = Memtrans.objects.filter(
            ses_date__range=(start_date, end_date)
        ).select_related('user', 'branch').order_by('ses_date')

        # Branch filtering
        if branch_obj:
            selected_branch = branch_obj
            report_data = report_data.filter(branch=selected_branch)
            session_date = selected_branch.session_date

        # User filtering
        if user_obj:
            report_data = report_data.filter(user_id=user_obj.id)

        # Code filtering
        if code:
            report_data = report_data.filter(code=code)

        # Annotate customer name
        report_data = report_data.annotate(
            customer_name=Subquery(
                Customer.objects.filter(
                    gl_no=OuterRef('gl_no'),
                    ac_no=OuterRef('ac_no')
                ).annotate(
                    full_name=Concat(
                        'first_name', Value(' '),
                        'middle_name', Value(' '),
                        'last_name',
                        output_field=CharField()
                    )
                ).values('full_name')[:1],
                output_field=CharField()
            )
        )

        # Totals
        totals = report_data.aggregate(
            total_debit=Sum('amount', filter=Q(type='D')),
            total_credit=Sum('amount', filter=Q(type='C'))
        )
        total_debit = totals['total_debit'] or 0
        total_credit = totals['total_credit'] or 0

        # ✅ Show warning if no results
        if not report_data.exists():
            messages.warning(request, "No transactions found for the selected criteria.")

    return render(request, 'reports/savings/general_trx_register_by_session_date.html', {
        'form': form,
        'branches': branches,
        'users': users,
        'current_datetime': current_datetime,
        'report_data': report_data,
        'selected_branch': selected_branch,
        'start_date': start_date,
        'end_date': end_date,
        'session_date': session_date,
        'total_debit': total_debit,
        'total_credit': total_credit,
    })

from django.shortcuts import render, get_object_or_404
from django.utils import timezone
from django.db.models import Sum, Q, Value, OuterRef, Subquery, CharField
from django.db.models.functions import Concat
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.contrib.auth import get_user_model

from .forms import TransactionForm
from transactions.models import Memtrans
from company.models import Branch
from customers.models import Customer

User = get_user_model()


@login_required
def general_trx_register_by_trx_date(request):
    current_datetime = timezone.now()
    report_data = None
    selected_branch = None
    start_date = None
    end_date = None
    session_date = None
    total_debit = 0
    total_credit = 0

    # ✅ Determine accessible branches
    if request.user.branch.head_office:
        branches = Branch.objects.filter(company=request.user.branch.company)
    else:
        branches = Branch.objects.filter(id=request.user.branch.id)

    # ✅ Users from those branches
    users = User.objects.filter(branch__in=branches).distinct()

    # ✅ Pass filtered branches/users into form
    form = TransactionForm(request.POST or None, branches=branches, users=users)

    if request.method == 'POST' and form.is_valid():
        start_date = form.cleaned_data['start_date']
        end_date = form.cleaned_data['end_date']
        branch_obj = form.cleaned_data['branch']
        user_obj = form.cleaned_data['user']
        code = form.cleaned_data['code']

        # Base queryset
        report_data = Memtrans.objects.filter(
            app_date__range=(start_date, end_date)
        ).select_related('user', 'branch').order_by('app_date')

        # Branch filtering
        if branch_obj:
            selected_branch = branch_obj
            report_data = report_data.filter(branch=selected_branch)
            session_date = selected_branch.session_date

        # User filtering
        if user_obj:
            report_data = report_data.filter(user_id=user_obj.id)

        # Code filtering
        if code:
            report_data = report_data.filter(code=code)

        # Annotate customer name
        report_data = report_data.annotate(
            customer_name=Subquery(
                Customer.objects.filter(
                    gl_no=OuterRef('gl_no'),
                    ac_no=OuterRef('ac_no')
                ).annotate(
                    full_name=Concat(
                        'first_name', Value(' '),
                        'middle_name', Value(' '),
                        'last_name',
                        output_field=CharField()
                    )
                ).values('full_name')[:1],
                output_field=CharField()
            )
        )

        # Totals
        totals = report_data.aggregate(
            total_debit=Sum('amount', filter=Q(type='D')),
            total_credit=Sum('amount', filter=Q(type='C'))
        )
        total_debit = totals['total_debit'] or 0
        total_credit = totals['total_credit'] or 0

        # ✅ Warn if no results
        if not report_data.exists():
            messages.warning(request, "No transactions found for the selected criteria.")

    return render(request, 'reports/savings/general_trx_register_by_trx_date.html', {
        'form': form,
        'branches': branches,
        'users': users,
        'current_datetime': current_datetime,
        'report_data': report_data,
        'selected_branch': selected_branch,
        'start_date': start_date,
        'end_date': end_date,
        'session_date': session_date,
        'total_debit': total_debit,
        'total_credit': total_credit,
    })


from django.db.models import Sum, F, ExpressionWrapper, FloatField, Value, CharField, Subquery, OuterRef
from django.db.models.functions import Concat
from django.shortcuts import render, get_object_or_404
from django.utils import timezone
from django.core.paginator import Paginator
from django.contrib import messages
from django.contrib.auth.decorators import login_required

from accounts_admin.models import Account
from transactions.models import Memtrans
from company.models import Branch
from customers.models import Customer
from django.contrib.auth import get_user_model

User = get_user_model()


@login_required
def cashier_teller_cash_status_by_session_date(request):
    current_datetime = timezone.now()

    # ✅ Determine accessible branches and users
    if request.user.branch.head_office:
        branches = Branch.objects.filter(company=request.user.branch.company)
    else:
        branches = Branch.objects.filter(id=request.user.branch.id)

    users = User.objects.filter(branch__in=branches).distinct()
    accounts = Account.objects.filter(gl_no__isnull=False).values('gl_no').distinct().order_by('gl_no')

    context = {
        'branches': branches,
        'users': users,
        'accounts': accounts,
        'account_gl_numbers': [acc['gl_no'] for acc in accounts],
        'current_datetime': current_datetime,
        'company_date': timezone.now().date(),
        'report_data': None,
        'total_debit': 0,
        'total_credit': 0,
        'start_date': None,
        'end_date': None,
    }

    if request.method == 'POST':
        try:
            gl_no = request.POST.get('gl_no')
            if gl_no == 'custom':
                gl_no = request.POST.get('gl_no_custom', '').strip() or None

            # ✅ Required dates
            start_date = request.POST.get('start_date')
            end_date = request.POST.get('end_date')
            if not start_date or not end_date:
                messages.error(request, "Both start date and end date are required")
                return render(request, 'reports/savings/cashier_teller_cash_status_by_session_date.html', context)

            try:
                start_date = timezone.datetime.strptime(start_date, '%Y-%m-%d').date()
                end_date = timezone.datetime.strptime(end_date, '%Y-%m-%d').date()
            except (ValueError, TypeError):
                messages.error(request, "Invalid date format")
                return render(request, 'reports/savings/cashier_teller_cash_status_by_session_date.html', context)

            # ✅ Base queryset
            report_data = Memtrans.objects.filter(
                ses_date__range=(start_date, end_date)
            ).select_related('user', 'branch').order_by('ses_date', 'id')

            # ✅ Branch filter
            branch_id = request.POST.get('branch')
            if branch_id:
                selected_branch = get_object_or_404(Branch, id=branch_id, company=request.user.branch.company)
                report_data = report_data.filter(branch=selected_branch)
                context['selected_branch_obj'] = selected_branch

            # ✅ User filter
            user_id = request.POST.get('user')
            if user_id:
                get_object_or_404(User, id=user_id, branch__in=branches)
                report_data = report_data.filter(user_id=user_id)

            # ✅ Other filters
            code = request.POST.get('code')
            if code:
                report_data = report_data.filter(code=code)

            if gl_no:
                report_data = report_data.filter(gl_no=gl_no)

            ac_no = request.POST.get('ac_no')
            if ac_no:
                report_data = report_data.filter(ac_no=ac_no)

            # ✅ No results check
            if not report_data.exists():
                messages.info(request, "No transactions found matching your criteria")
                context.update({
                    'start_date': start_date,
                    'end_date': end_date,
                    'gl_no': gl_no,
                    'ac_no': ac_no,
                    'code': code,
                    'user_id': user_id,
                    'branch_id': branch_id,
                })
                return render(request, 'reports/savings/cashier_teller_cash_status_by_session_date.html', context)

            # ✅ Annotate customer
            report_data = report_data.annotate(
                customer_name=Subquery(
                    Customer.objects.filter(
                        gl_no=OuterRef('gl_no'),
                        ac_no=OuterRef('ac_no')
                    ).annotate(
                        full_name=Concat(
                            'first_name', Value(' '),
                            'middle_name', Value(' '),
                            'last_name',
                            output_field=CharField()
                        )
                    ).values('full_name')[:1]
                )
            )

            # ✅ Running balance
            running_balance = 0
            for trx in report_data:
                running_balance += trx.amount
                trx.running_balance = running_balance

            # ✅ Pagination
            paginator = Paginator(report_data, 50)
            page_number = request.GET.get('page', 1)
            page_obj = paginator.get_page(page_number)

            # ✅ Totals
            total_credit = report_data.filter(amount__gt=0).aggregate(Sum('amount'))['amount__sum'] or 0
            total_debit = report_data.filter(amount__lt=0).aggregate(
                total_debit=Sum(ExpressionWrapper(F('amount') * -1, output_field=FloatField()))
            )['total_debit'] or 0

            context.update({
                'start_date': start_date,
                'end_date': end_date,
                'gl_no': gl_no,
                'ac_no': ac_no,
                'code': code,
                'user_id': user_id,
                'branch_id': branch_id,
                'report_data': page_obj,
                'page_obj': page_obj,
                'total_debit': abs(total_debit),
                'total_credit': abs(total_credit),
            })

        except Exception as e:
            messages.error(request, f"An error occurred: {str(e)}")

    return render(request, 'reports/savings/cashier_teller_cash_status_by_session_date.html', context)



from django.db.models import Sum, F, ExpressionWrapper, FloatField, Value, CharField, Subquery, OuterRef
from django.db.models.functions import Concat
from django.shortcuts import render, get_object_or_404
from django.utils import timezone
from django.core.paginator import Paginator
from django.contrib import messages
from django.contrib.auth.decorators import login_required

from accounts_admin.models import Account
from transactions.models import Memtrans
from company.models import Branch
from customers.models import Customer
from django.contrib.auth import get_user_model

User = get_user_model()


@login_required
def cashier_teller_cash_status_by_trx_date(request):
    current_datetime = timezone.now()

    # ✅ Branch & User filtering (head office vs branch user)
    if request.user.branch.head_office:
        branches = Branch.objects.filter(company=request.user.branch.company)
    else:
        branches = Branch.objects.filter(id=request.user.branch.id)

    users = User.objects.filter(branch__in=branches).distinct()
    accounts = Account.objects.filter(gl_no__isnull=False).values('gl_no').distinct().order_by('gl_no')

    context = {
        'branches': branches,
        'users': users,
        'accounts': accounts,
        'account_gl_numbers': [acc['gl_no'] for acc in accounts],
        'current_datetime': current_datetime,
        'report_data': None,
        'total_debit': 0,
        'total_credit': 0,
        'start_date': None,
        'end_date': None,
    }

    if request.method == 'POST':
        try:
            # ✅ Extract POST params
            start_date = request.POST.get('start_date')
            end_date = request.POST.get('end_date')
            branch_id = request.POST.get('branch')
            user_id = request.POST.get('user')
            code = request.POST.get('code')

            # GL number (custom handling)
            gl_no = request.POST.get('gl_no')
            if gl_no == 'custom':
                gl_no = request.POST.get('gl_no_custom', '').strip()

            ac_no = request.POST.get('ac_no')
            page = request.GET.get('page', 1)

            # ✅ Validate dates
            if not start_date or not end_date:
                messages.error(request, "Both start date and end date are required")
                return render(request, 'reports/savings/cashier_teller_cash_status_by_trx_date.html', context)

            try:
                start_date = timezone.datetime.strptime(start_date, '%Y-%m-%d').date()
                end_date = timezone.datetime.strptime(end_date, '%Y-%m-%d').date()
            except (ValueError, TypeError):
                messages.error(request, "Invalid date format")
                return render(request, 'reports/savings/cashier_teller_cash_status_by_trx_date.html', context)

            # ✅ Base queryset (transaction date = app_date)
            report_data = Memtrans.objects.filter(
                app_date__range=(start_date, end_date)
            ).select_related('user', 'branch').order_by('app_date', 'id')

            # ✅ Branch filter (ensure belongs to same company)
            if branch_id:
                selected_branch = get_object_or_404(Branch, id=branch_id, company=request.user.branch.company)
                report_data = report_data.filter(branch=selected_branch)
                context['selected_branch'] = selected_branch
                context['selected_branch_obj'] = selected_branch

            # ✅ User filter (ensure belongs to accessible branches)
            if user_id:
                get_object_or_404(User, id=user_id, branch__in=branches)
                report_data = report_data.filter(user_id=user_id)

            # ✅ Other filters
            if code:
                report_data = report_data.filter(code=code)
            if gl_no:
                report_data = report_data.filter(gl_no=gl_no)
            if ac_no:
                report_data = report_data.filter(ac_no=ac_no)

            # ✅ No results
            if not report_data.exists():
                messages.info(request, "No transactions found matching your criteria")
                context.update({
                    'start_date': start_date,
                    'end_date': end_date,
                    'gl_no': gl_no,
                    'ac_no': ac_no,
                    'code': code,
                    'user_id': user_id,
                    'branch_id': branch_id,
                })
                return render(request, 'reports/savings/cashier_teller_cash_status_by_trx_date.html', context)

            # ✅ Annotate customer name
            report_data = report_data.annotate(
                customer_name=Subquery(
                    Customer.objects.filter(
                        gl_no=OuterRef('gl_no'),
                        ac_no=OuterRef('ac_no')
                    ).annotate(
                        full_name=Concat(
                            'first_name', Value(' '),
                            'middle_name', Value(' '),
                            'last_name',
                            output_field=CharField()
                        )
                    ).values('full_name')[:1]
                )
            )

            # ✅ Running balance
            running_balance = 0
            for trx in report_data:
                running_balance += trx.amount
                trx.running_balance = running_balance

            # ✅ Pagination
            paginator = Paginator(report_data, 50)
            page_obj = paginator.get_page(page)

            # ✅ Totals
            total_credit = report_data.filter(amount__gt=0).aggregate(Sum('amount'))['amount__sum'] or 0
            total_debit = report_data.filter(amount__lt=0).aggregate(
                total_debit=Sum(ExpressionWrapper(F('amount') * -1, output_field=FloatField()))
            )['total_debit'] or 0

            context.update({
                'start_date': start_date,
                'end_date': end_date,
                'gl_no': gl_no,
                'ac_no': ac_no,
                'code': code,
                'user_id': user_id,
                'branch_id': branch_id,
                'report_data': page_obj,
                'page_obj': page_obj,
                'total_debit': abs(total_debit),
                'total_credit': abs(total_credit),
            })

        except Exception as e:
            messages.error(request, f"An error occurred: {str(e)}")

    return render(request, 'reports/savings/cashier_teller_cash_status_by_trx_date.html', context)


# views.py
# views.py
from django.shortcuts import render, get_object_or_404
from django.utils import timezone
from django.db.models import Sum
from datetime import datetime
from django.contrib import messages

from accounts_admin.models import Account
from transactions.models import Memtrans
from company.models import Branch
from customers.models import Customer


def account_statement(request):
    logged_in_user = request.user

    # ✅ Branch filtering: head office sees all, branch user sees only their branch
    if logged_in_user.branch.head_office:
        branches = Branch.objects.filter(company=logged_in_user.branch.company)
    else:
        branches = Branch.objects.filter(id=logged_in_user.branch.id)

    # ✅ Accounts filtered by accessible branches
    gl_data = Account.objects.all().values('gl_no', 'gl_name').distinct()
    gl_nos = [(item['gl_no'], item['gl_name']) for item in gl_data]

    # Form params
    gl_no = request.POST.get('gl_no')
    ac_no = request.POST.get('ac_no')
    start_date = request.POST.get('start_date')
    end_date = request.POST.get('end_date')
    branch_id = request.POST.get('branch')

    # Defaults
    transactions = []
    statement_data = []
    opening_balance = total_debit = total_credit = reporting_balance = 0
    selected_branch = None
    customer = None

    # ✅ Date parsing
    if start_date:
        try:
            start_date = datetime.strptime(start_date, '%Y-%m-%d').date()
        except ValueError:
            messages.error(request, "Invalid start date format")
            start_date = None

    if end_date:
        try:
            end_date = datetime.strptime(end_date, '%Y-%m-%d').date()
        except ValueError:
            messages.error(request, "Invalid end date format")
            end_date = None

    # ✅ If all required fields provided
    if gl_no and ac_no and start_date and end_date:
        # Base queryset
        transactions = Memtrans.objects.filter(
            gl_no=gl_no,
            ac_no=ac_no,
            app_date__range=[start_date, end_date],
            branch__in=branches
        )

        # ✅ Branch filter
        if branch_id:
            selected_branch = get_object_or_404(Branch, id=branch_id, company=logged_in_user.branch.company)
            transactions = transactions.filter(branch=selected_branch)
        else:
            selected_branch = branches.first()

        # ✅ Customer filter
        try:
            customer = Customer.objects.get(gl_no=gl_no, ac_no=ac_no, branch__in=branches)
        except Customer.DoesNotExist:
            messages.warning(request, "Customer not found for the provided GL/AC number")
            customer = None

        # ✅ Opening balance
        opening_balance = Memtrans.objects.filter(
            gl_no=gl_no,
            ac_no=ac_no,
            app_date__lt=start_date,
            branch__in=(branches if not branch_id else [selected_branch])
        ).aggregate(total=Sum('amount'))['total'] or 0

        balance = opening_balance
        previous_transaction_date = None

        # Add opening balance entry
        statement_data.append({
            'branch': selected_branch.branch_name if selected_branch else 'All Your Branches',
            'trx_date': start_date,
            'trx_no': 'Opening Balance',
            'description': 'Opening Balance',
            'debit': 0,
            'credit': 0,
            'balance': balance,
            'days_without_activity': '',
        })


        # ✅ Transaction loop
        for trx in transactions.order_by('app_date'):
            debit = abs(trx.amount) if trx.amount < 0 else 0
            credit = trx.amount if trx.amount > 0 else 0
            balance += (credit - debit)

            total_debit += debit
            total_credit += credit

            days_without_activity = (
                (trx.app_date - previous_transaction_date).days if previous_transaction_date else 0
            )

        statement_data.append({
            'branch': trx.branch.branch_name,
            'ses_date': trx.ses_date,
            'trx_date': trx.app_date,
            'trx_no': trx.trx_no,
            'description': trx.description,
            'debit': debit if debit else 0,
            'credit': credit if credit else 0,
            'balance': balance,
            'days_without_activity': days_without_activity,
        })


        previous_transaction_date = trx.app_date

        reporting_balance = balance

        if not transactions.exists():
            messages.info(request, "No transactions found for this account in the given date range.")

    context = {
        'branches': branches,
        'gl_nos': gl_nos,
        'statement_data': statement_data,
        'gl_no': gl_no,
        'ac_no': ac_no,
        'start_date': start_date.strftime('%Y-%m-%d') if start_date else '',
        'end_date': end_date.strftime('%Y-%m-%d') if end_date else '',
        'branch': selected_branch,
        'company': selected_branch.company if selected_branch else None,
        'current_datetime': timezone.now(),
        'opening_balance': f"{opening_balance:,}",
        'closing_balance': f"{reporting_balance:,}",
        'total_debit': f"{total_debit:,}",
        'total_credit': f"{total_credit:,}",
        'reporting_balance': f"{reporting_balance:,}",
        'full_name': customer.get_full_name() if customer else '',
        'customer': customer,
    }

    return render(request, 'reports/savings_report/savings_account_statement.html', context)


    
from django.shortcuts import render, get_object_or_404
from django.db.models import Sum
from django.utils import timezone
from customers.models import Customer
from transactions.models import Memtrans
from django.db.models import Sum, Q, Min
from company.models import Company
from datetime import datetime
from accounts_admin.models import Account, Region, Account_Officer
from company.models import Company



from datetime import timedelta

from django.utils import timezone





from django.utils import timezone

import pytz  # Optional: for timezone-aware datetime

from django.shortcuts import render
from django.db.models import Sum
# from .models import Customer, Memtrans, Branch, Region, Account_Officer, Account
from django.utils import timezone
from datetime import datetime






from django.contrib import messages
from django.db.models import Sum
from django.utils import timezone
from django.core.exceptions import PermissionDenied
from datetime import datetime

def all_customers_account_balances(request):
    customer_data = {}
    default_reporting_date = timezone.now().date()
    current_datetime = timezone.now()

    selected_branch = None
    selected_gl_no = None
    selected_region = None
    selected_officer = None
    reporting_date = default_reporting_date
    include_non_zero = False
    exclude_ac_no_one = False
    grand_total = 0

    # ✅ Get user branch safely
    user_branch = getattr(request.user, "branch", None) if request.user.is_authenticated else None

    # Defaults
    branches = []
    regions = []
    account_officers = []
    gl_accounts = []
    customers = Customer.objects.none()

    # ✅ Handle branch access rules
    if user_branch:
        if user_branch.head_office:
            # Head office sees everything
            branches = Branch.objects.all()
            regions = Region.objects.all()
            account_officers = Account_Officer.objects.all()
            gl_accounts = Account.objects.all().distinct("gl_no")
            customers = Customer.objects.all()
        else:
            # Normal branch user sees only their branch
            branches = [user_branch]
            regions = Region.objects.filter(branch=user_branch)
            account_officers = Account_Officer.objects.filter(branch=user_branch)
            gl_accounts = Account.objects.filter(branch=user_branch).distinct("gl_no")
            customers = Customer.objects.filter(branch=user_branch)

    if request.method == "POST":
        reporting_date_str = request.POST.get("reporting_date")
        branch_id = request.POST.get("branch")
        gl_no = request.POST.get("gl_no")
        region_id = request.POST.get("region")
        officer_id = request.POST.get("credit_officer")
        include_non_zero = request.POST.get("include_non_zero") == "on"
        exclude_ac_no_one = request.POST.get("exclude_ac_no_one") == "on"

        # ✅ Reporting date
        if reporting_date_str:
            reporting_date = datetime.strptime(reporting_date_str, "%Y-%m-%d").date()
        else:
            reporting_date = default_reporting_date

        filtered_customers = customers

        # ✅ Branch filter with permission check
        if branch_id:
            try:
                selected_branch = Branch.objects.get(id=branch_id)
                if not user_branch.head_office and selected_branch != user_branch:
                    raise PermissionDenied("You do not have access to this branch.")
                filtered_customers = filtered_customers.filter(branch=selected_branch)
            except Branch.DoesNotExist:
                selected_branch = None
        else:
            selected_branch = user_branch if user_branch and not user_branch.head_office else None

        # ✅ GL filter
        if gl_no:
            filtered_customers = filtered_customers.filter(gl_no=gl_no)
            selected_gl_no = gl_no

        # ✅ Region filter
        if region_id:
            try:
                if user_branch.head_office:
                    selected_region = Region.objects.get(id=region_id)
                else:
                    selected_region = Region.objects.get(id=region_id, branch=user_branch)
                filtered_customers = filtered_customers.filter(region=selected_region)
            except Region.DoesNotExist:
                selected_region = None

        # ✅ Officer filter
        if officer_id:
            try:
                if user_branch.head_office:
                    selected_officer = Account_Officer.objects.get(id=officer_id)
                else:
                    selected_officer = Account_Officer.objects.get(id=officer_id, branch=user_branch)
                filtered_customers = filtered_customers.filter(credit_officer=selected_officer)
            except Account_Officer.DoesNotExist:
                selected_officer = None

        # ✅ Customer dict for grouping
        customer_dict = {}
        for customer in filtered_customers:
            if exclude_ac_no_one and customer.ac_no == "1":
                continue
            customer_dict[(customer.gl_no, customer.ac_no)] = customer

        # ✅ Transactions up to reporting date
        transactions = Memtrans.objects.filter(app_date__lte=reporting_date)
        if not user_branch.head_office:
            transactions = transactions.filter(customer__branch=user_branch)
        elif selected_branch:
            transactions = transactions.filter(branch=selected_branch)

        # ✅ GL map
        gl_map = {a.gl_no: a.gl_name for a in gl_accounts}

        # ✅ Compute balances
        for (gl_no, ac_no), customer in customer_dict.items():
            customer_transactions = transactions.filter(gl_no=gl_no, ac_no=ac_no)
            account_balance = customer_transactions.aggregate(total=Sum("amount"))["total"] or 0

            if include_non_zero and account_balance == 0:
                continue

            gl_name = gl_map.get(gl_no, "Unknown Account")

            if gl_no not in customer_data:
                customer_data[gl_no] = {
                    "gl_name": gl_name,
                    "customers": [],
                    "subtotal": 0,
                    "count": 0,
                }

            customer_data[gl_no]["customers"].append({
                "first_name": customer.first_name,
                "middle_name": customer.middle_name,
                "last_name": customer.last_name,
                "address": customer.address,
                "account_balance": account_balance,
                "gl_name": gl_name,
            })

            customer_data[gl_no]["subtotal"] += account_balance
            customer_data[gl_no]["count"] += 1
            grand_total += account_balance

        # ✅ If no data found, notify user
        if not customer_data:
            messages.info(request, "No data found for the selected filters.")

    # ✅ Company for header
    company = user_branch.company if user_branch else None
    username = request.user.username if request.user.is_authenticated else None

    context = {
        "branches": branches,
        "regions": regions,
        "account_officers": account_officers,
        "gl_accounts": gl_accounts,
        "customer_data": customer_data,
        "selected_branch": selected_branch,
        "selected_gl_no": selected_gl_no,
        "selected_region": selected_region,
        "selected_officer": selected_officer,
        "reporting_date": reporting_date,
        "include_non_zero": include_non_zero,
        "exclude_ac_no_one": exclude_ac_no_one,
        "grand_total": grand_total,
        "current_datetime": current_datetime,
        "company": company,
        "username": username,
        "user": request.user,
        "user_branch": user_branch,
    }

    return render(request, "reports/savings_report/savings_account_balance_report.html", context)
from django.shortcuts import render, redirect
from django.utils import timezone
from datetime import datetime
from django.db.models import Sum
from django.contrib import messages
import logging

def savings_transaction_report(request):
    logger = logging.getLogger(__name__)
    user = request.user
    user_branch = getattr(user, 'branch', None)  # Corrected

    if not user_branch:
        return render(request, 'error.html', {'message': 'User has no branch assigned.'})
    
    user_company_name = user_branch.company_name
    is_head_office = user_branch.head_office

    try:
        # Initialize filter options based on head office or branch
        if is_head_office:
            branches = Branch.objects.filter(company_name=user_company_name).order_by('branch_name')
            regions = Region.objects.filter(branch__company_name=user_company_name)
            account_officers = Account_Officer.objects.filter(region__branch__company_name=user_company_name)
        else:
            branches = Branch.objects.filter(id=user_branch.id)
            regions = Region.objects.filter(branch=user_branch)
            account_officers = Account_Officer.objects.filter(region__branch=user_branch)
        
        gl_accounts = Account.objects.filter(
            gl_no__in=Customer.objects.filter(branch__company_name=user_company_name).values('gl_no').distinct(),
            branch__company_name=user_company_name
        )
        current_datetime = timezone.now()

        # Default values
        transaction_data = {}
        default_start_date = timezone.now().date()
        default_end_date = timezone.now().date()
        form_data = {
            'selected_branch': None,
            'selected_gl_no': None,
            'selected_region': None,
            'selected_officer': None,
            'start_date': default_start_date,
            'end_date': default_end_date,
            'include_non_zero': False,
            'exclude_ac_no_one': False,
        }
        grand_total_debit = 0
        grand_total_credit = 0

        if request.method == 'POST':
            # Form dates
            form_data['start_date'] = datetime.strptime(request.POST.get('start_date', ''), '%Y-%m-%d').date() \
                if request.POST.get('start_date') else default_start_date
            form_data['end_date'] = datetime.strptime(request.POST.get('end_date', ''), '%Y-%m-%d').date() \
                if request.POST.get('end_date') else default_end_date

            form_data['include_non_zero'] = request.POST.get('include_non_zero') == 'on'
            form_data['exclude_ac_no_one'] = request.POST.get('exclude_ac_no_one') == 'on'

            # Base customer queryset
            customers = Customer.objects.filter(branch__company_name=user_company_name)

            # Branch filter
            branch_id = request.POST.get('branch')
            if branch_id:
                try:
                    branch = Branch.objects.get(id=branch_id, company_name=user_company_name)
                    customers = customers.filter(branch=branch.branch_code)
                    form_data['selected_branch'] = branch
                except Branch.DoesNotExist:
                    messages.warning(request, "Selected branch not found")
            elif not is_head_office and user_branch:
                customers = customers.filter(branch=user_branch.branch_code)
                form_data['selected_branch'] = user_branch

            # GL No filter
            if gl_no := request.POST.get('gl_no'):
                customers = customers.filter(gl_no=gl_no)
                form_data['selected_gl_no'] = gl_no

            # Region filter
            if region_id := request.POST.get('region'):
                try:
                    region = Region.objects.get(id=region_id, branch__company_name=user_company_name)
                    branch_codes = Branch.objects.filter(region=region, company_name=user_company_name).values_list('branch_code', flat=True)
                    customers = customers.filter(branch__in=branch_codes)
                    form_data['selected_region'] = region
                except Region.DoesNotExist:
                    messages.warning(request, "Selected region not found")

            # Account Officer filter
            if officer_id := request.POST.get('credit_officer'):
                try:
                    officer = Account_Officer.objects.get(id=officer_id, region__branch__company_name=user_company_name)
                    branch_codes = Branch.objects.filter(account_officer=officer, company_name=user_company_name).values_list('branch_code', flat=True)
                    customers = customers.filter(branch__in=branch_codes)
                    form_data['selected_officer'] = officer
                except Account_Officer.DoesNotExist:
                    messages.warning(request, "Selected account officer not found")

            # GL name lookup dictionary
            gl_name_dict = dict(Account.objects.filter(branch__company_name=user_company_name).values_list('gl_no', 'gl_name'))

            # Customer dictionary for lookup
            customer_dict = {
                (c.gl_no, c.ac_no): c for c in customers
                if not (form_data['exclude_ac_no_one'] and c.ac_no == '1')
            }

            # Transactions in date range
            transactions = Memtrans.objects.filter(
                app_date__gte=form_data['start_date'],
                app_date__lte=form_data['end_date'],
                branch__company_name=user_company_name
            )

            if form_data['selected_branch']:
                transactions = transactions.filter(branch=form_data['selected_branch'].branch_code)

            # Process transactions per customer
            for (gl_no, ac_no), customer in customer_dict.items():
                cust_transactions = transactions.filter(gl_no=gl_no, ac_no=ac_no)

                if not form_data['include_non_zero']:
                    cust_transactions = cust_transactions.exclude(amount=0)

                if not cust_transactions.exists():
                    continue

                if gl_no not in transaction_data:
                    transaction_data[gl_no] = {
                        'gl_name': gl_name_dict.get(gl_no, 'Unknown'),
                        'transactions': [],
                        'subtotal_debit': 0,
                        'subtotal_credit': 0,
                        'count': 0
                    }

                for trx in cust_transactions:
                    debit = abs(trx.amount) if trx.amount < 0 else 0
                    credit = trx.amount if trx.amount > 0 else 0

                    transaction_data[gl_no]['transactions'].append({
                        'trx_no': trx.trx_no,
                        'ses_date': trx.ses_date,
                        'app_date': trx.app_date,
                        'description': trx.description,
                        'debit': debit,
                        'credit': credit,
                        'ac_no': customer.ac_no,
                        'customer_name': f"{customer.first_name} {customer.last_name}"
                    })

                    transaction_data[gl_no]['subtotal_debit'] += debit
                    transaction_data[gl_no]['subtotal_credit'] += credit
                    transaction_data[gl_no]['count'] += 1

                    grand_total_debit += debit
                    grand_total_credit += credit

        context = {
            'branches': branches,
            'regions': regions,
            'account_officers': account_officers,
            'gl_accounts': gl_accounts,
            'transaction_data': transaction_data,
            'grand_total_debit': grand_total_debit,
            'grand_total_credit': grand_total_credit,
            'current_datetime': current_datetime,
            'user_branch': user_branch,
            'user_branch_is_head_office': is_head_office,
            'user_company_name': user_company_name,
            **form_data
        }

        return render(request, 'reports/savings_report/savings_transactions_report.html', context)

    except Exception as e:
        logger.error(f"Error in savings_transaction_report: {str(e)}", exc_info=True)
        messages.error(request, f"An error occurred while generating the transaction report: {str(e)}")
        return render(request, 'reports/savings_report/savings_transactions_report.html', {
            'branches': Branch.objects.filter(company_name=user_company_name).order_by('branch_name'),
            'regions': Region.objects.filter(branch__company_name=user_company_name),
            'account_officers': Account_Officer.objects.filter(region__branch__company_name=user_company_name),
            'gl_accounts': Account.objects.filter(
                gl_no__in=Customer.objects.filter(branch__company_name=user_company_name).values('gl_no').distinct(),
                branch__company_name=user_company_name
            ),
            'transaction_data': {},
            'grand_total_debit': 0,
            'grand_total_credit': 0,
            'current_datetime': timezone.now(),
            'user_branch': user_branch,
            'user_branch_is_head_office': is_head_office,
            'user_company_name': user_company_name,
            **form_data
        })



from datetime import datetime
from django.db.models import Min, Max



from datetime import timedelta
from django.db.models import Sum



from django.db.models import Sum, Max
from django.shortcuts import render
from django.utils import timezone
from datetime import datetime, timedelta
from django.db.models import Sum, Max
from django.shortcuts import render
from django.utils import timezone
from datetime import datetime, timedelta



from datetime import datetime, timedelta
from django.db.models import Sum, Max
from django.utils import timezone
from django.shortcuts import render
from django.core.exceptions import PermissionDenied

def savings_account_listing(request):
    # Initialize variables
    start_date = end_date = None
    selected_branch = selected_gl_no = selected_region = selected_officer = None
    include_non_zero = exclude_ac_no_one = False
    customer_data = {}
    grand_total = 0

    # ✅ Get user's branch safely
    user_branch = getattr(request.user, "branch", None)
    user_company = user_branch.company if user_branch else None

    # Default dates
    default_end_date = timezone.now().date()
    default_start_date = default_end_date

    if request.method == "POST":
        # Get form data
        start_date_str = request.POST.get("start_date")
        end_date_str = request.POST.get("end_date")
        branch_id = request.POST.get("branch")
        selected_gl_no = request.POST.get("gl_no")
        region_id = request.POST.get("region")
        officer_id = request.POST.get("credit_officer")
        include_non_zero = request.POST.get("include_non_zero") == "on"
        exclude_ac_no_one = request.POST.get("exclude_ac_no_one") == "on"

        # Convert dates
        start_date = datetime.strptime(start_date_str, "%Y-%m-%d").date() if start_date_str else default_start_date
        end_date = datetime.strptime(end_date_str, "%Y-%m-%d").date() if end_date_str else default_end_date

        # ✅ Handle head office vs normal branch
        if user_branch and user_branch.head_office:
            # Head office user → can see all customers
            customers = Customer.objects.all()
        elif user_branch:
            # Normal branch user → only customers in their branch
            customers = Customer.objects.filter(branch=user_branch)
        else:
            customers = Customer.objects.none()

        # Branch filter
        if branch_id:
            try:
                selected_branch = Branch.objects.get(id=branch_id)
                if not (user_branch and user_branch.head_office) and selected_branch != user_branch:
                    raise PermissionDenied("You do not have access to this branch.")
                customers = customers.filter(branch=selected_branch)
            except Branch.DoesNotExist:
                selected_branch = None

        # GL filter
        if selected_gl_no:
            customers = customers.filter(gl_no=selected_gl_no)

        # Region filter
        if region_id:
            try:
                selected_region = Region.objects.get(id=region_id)
                if not (user_branch and user_branch.head_office):
                    selected_region = Region.objects.get(id=region_id, branch=user_branch)
                customers = customers.filter(region=selected_region)
            except Region.DoesNotExist:
                selected_region = None

        # Officer filter
        if officer_id:
            try:
                selected_officer = Account_Officer.objects.get(id=officer_id)
                if not (user_branch and user_branch.head_office):
                    selected_officer = Account_Officer.objects.get(id=officer_id, branch=user_branch)
                customers = customers.filter(credit_officer=selected_officer)
            except Account_Officer.DoesNotExist:
                selected_officer = None

        # Create customer dictionary
        customer_dict = {}
        for customer in customers:
            if exclude_ac_no_one and customer.ac_no == "1":
                continue
            customer_dict[(customer.gl_no, customer.ac_no)] = customer

        # Filter transactions
        transactions = Memtrans.objects.filter(app_date__gte=start_date, app_date__lte=end_date)
        if user_branch and not user_branch.head_office:
            transactions = transactions.filter(customer__branch=user_branch)
        elif selected_branch:
            transactions = transactions.filter(branch=selected_branch)

        # Process customer data
        for (gl_no, ac_no), customer in customer_dict.items():
            customer_transactions = transactions.filter(gl_no=gl_no, ac_no=ac_no)
            total_balance = customer_transactions.aggregate(total=Sum("amount"))["total"] or 0

            if include_non_zero and total_balance == 0:
                continue

            if gl_no not in customer_data:
                account_obj = Account.objects.filter(gl_no=gl_no).first()
                gl_name = account_obj.gl_name if account_obj else "Unknown"
                customer_data[gl_no] = {"gl_name": gl_name, "customers": [], "subtotal": 0, "count": 0}

            last_trx_date = customer_transactions.aggregate(latest_date=Max("app_date"))["latest_date"]
            trx_dates = customer_transactions.values_list("sys_date", flat=True).distinct()

            days_without_activity = sum(
                1
                for current_date in (start_date + timedelta(days=i) for i in range((end_date - start_date).days + 1))
                if current_date not in {trx_date.date() for trx_date in trx_dates}
            )

            customer_info = {
                "gl_no": gl_no,
                "first_name": customer.first_name,
                "middle_name": customer.middle_name,
                "last_name": customer.last_name,
                "address": customer.address,
                "account_balance": total_balance,
                "last_trx_date": last_trx_date,
                "days_without_activity": days_without_activity,
            }

            customer_data[gl_no]["customers"].append(customer_info)
            customer_data[gl_no]["subtotal"] += total_balance
            customer_data[gl_no]["count"] += 1
            grand_total += total_balance

    # Prepare dropdown data
    if user_branch and user_branch.head_office:
        branches = Branch.objects.all()
        gl_accounts = Account.objects.exclude(gl_no="").order_by("gl_no").distinct()
        regions = Region.objects.all()
        account_officers = Account_Officer.objects.all()
    elif user_branch:
        branches = [user_branch]
        gl_accounts = Account.objects.filter(branch=user_branch).exclude(gl_no="").order_by("gl_no").distinct()
        regions = Region.objects.filter(branch=user_branch)
        account_officers = Account_Officer.objects.filter(branch=user_branch)
    else:
        branches = []
        gl_accounts = []
        regions = []
        account_officers = []

    context = {
        "customer_data": customer_data,
        "start_date": start_date,
        "end_date": end_date,
        "selected_branch": selected_branch,
        "selected_gl_no": selected_gl_no,
        "selected_region": selected_region,
        "selected_officer": selected_officer,
        "include_non_zero": include_non_zero,
        "exclude_ac_no_one": exclude_ac_no_one,
        "branches": branches,
        "gl_accounts": gl_accounts,
        "regions": regions,
        "account_officers": account_officers,
        "grand_total": grand_total,
        "current_datetime": timezone.now(),
        "company": user_branch.company if user_branch else None,
        'user_branch': user_branch, 
    }

    return render(request, "reports/savings_report/savings_account_listing.html", context)




@login_required
def savings_account_status(request):
    user = request.user

    # Get user's branch
    try:
        user_branch = user.branch
    except Branch.DoesNotExist:
        return render(request, 'error.html', {'message': 'User has no branch assigned.'})

    user_company_name = user_branch.company_name

    # Determine accessible branches
    if user_branch.head_office:
        branches = Branch.objects.filter(company_name=user_company_name)
    else:
        branches = Branch.objects.filter(id=user_branch.id)

    # Base querysets
    gl_accounts = Account.objects.filter(branch__company_name=user_company_name)
    customers = Customer.objects.filter(branch__company_name=user_company_name, label='C')
    memtrans = Memtrans.objects.filter(branch__company_name=user_company_name)
    regions = Region.objects.filter(branch__company_name=user_company_name)
    account_officers = Account_Officer.objects.filter(region__branch__company_name=user_company_name)

    # Initialize context
    context = {
        'branches': branches,
        'gl_accounts': gl_accounts,
        'regions': regions,
        'account_officers': account_officers,
        'customer_data': [],
        'user_company_name': user_company_name,
        'user_branch': user_branch,
        'include_non_zero': False,
        'exclude_ac_no_one': False,
        'grand_total': 0,
        'selected_branch': None,
        'selected_gl_no': None,
        'selected_region': None,
        'selected_officer': None,
        'start_date': request.GET.get('start_date', ''),
        'end_date': request.GET.get('end_date', ''),
    }

    # Only process if form submitted (GET parameters present)
    if request.GET:
        # Read checkboxes (GET)
        include_non_zero = request.GET.get("include_non_zero") == "on"
        exclude_ac_no_one = request.GET.get("exclude_ac_no_one") == "on"

        # Read filters
        branch_filter = request.GET.get('branch') or None
        region_filter = request.GET.get('region') or None
        officer_filter = request.GET.get('account_officer') or None
        gl_no_filter = request.GET.get('gl_no') or None

        # Filter by branch
        selected_branch = None
        if branch_filter:
            selected_branch = Branch.objects.filter(id=branch_filter).first()
            if selected_branch:
                gl_accounts = gl_accounts.filter(branch=selected_branch)
                customers = customers.filter(branch=selected_branch)
                memtrans = memtrans.filter(branch=selected_branch)
                regions = regions.filter(branch=selected_branch)
                account_officers = account_officers.filter(region__branch=selected_branch)

        # Filter by region
        if region_filter:
            customers = customers.filter(region_id=region_filter)
            account_officers = account_officers.filter(region_id=region_filter)

        # Filter by account officer
        if officer_filter:
            customers = customers.filter(account_officer_id=officer_filter)

        # Filter by GL number
        if gl_no_filter:
            customers = customers.filter(gl_no=gl_no_filter)

        # Optional: filter by start and end date
        start_date_str = request.GET.get('start_date')
        end_date_str = request.GET.get('end_date')
        try:
            start_date = datetime.strptime(start_date_str, "%Y-%m-%d").date() if start_date_str else None
            end_date = datetime.strptime(end_date_str, "%Y-%m-%d").date() if end_date_str else None
        except ValueError:
            start_date = end_date = None

        # Prepare customer data
        customer_data = []
        for cust in customers:
            cust_trans = memtrans.filter(ac_no=cust.ac_no)
            if start_date:
                cust_trans = cust_trans.filter(sys_date__date__gte=start_date)
            if end_date:
                cust_trans = cust_trans.filter(sys_date__date__lte=end_date)

            balance = cust_trans.aggregate(total_balance=Sum('amount'))['total_balance'] or 0
            last_trx = cust_trans.aggregate(last_date=Max('sys_date'))['last_date']
            days_without_activity = (now().date() - last_trx.date()).days if last_trx else None

            customer_data.append({
                'customer': cust,
                'balance': balance,
                'last_transaction_date': last_trx,
                'days_without_activity': days_without_activity,
                'reg_date': cust.reg_date,
            })

        # Apply checkbox filters
        if include_non_zero:
            customer_data = [c for c in customer_data if c["balance"] != 0]
        if exclude_ac_no_one:
            customer_data = [c for c in customer_data if getattr(c["customer"], "label", "C") == "C"]

        # Calculate grand total
        grand_total = sum(c["balance"] for c in customer_data)

        # Update context
        context.update({
            'customer_data': customer_data,
            'include_non_zero': include_non_zero,
            'exclude_ac_no_one': exclude_ac_no_one,
            'grand_total': grand_total,
            'selected_branch': selected_branch,
            'selected_gl_no': gl_no_filter,
            'selected_region': region_filter,
            'selected_officer': officer_filter,
            'start_date': start_date_str,
            'end_date': end_date_str,
        })

    return render(request, 'reports/savings_report/savings_account_status.html', context)

from django.shortcuts import render, redirect
from django.utils import timezone
from datetime import datetime
from django.db.models import Max, Sum
from django.contrib import messages

def savings_account_with_zero_balance(request):
    user = request.user

    # Get the user's branch/company_name
    user_branch = getattr(user, 'branch', None)
    if not user_branch:
        return render(request, 'error.html', {'message': 'User has no branch assigned.'})
    user_company_name = user_branch.company_name

    # Initialize filter options - filtered by company_name
    branches = Branch.objects.filter(company_name=user_company_name)
    regions = Region.objects.filter(branch__company_name=user_company_name)
    account_officers = Account_Officer.objects.filter(region__branch__company_name=user_company_name)
    gl_accounts = Account.objects.filter(
        gl_no__in=Customer.objects.filter(branch__company_name=user_company_name).values('gl_no').distinct(),
        branch__company_name=user_company_name
    )
    current_datetime = timezone.now()

    try:
        # Default values
        customer_data = {}
        default_reporting_date = timezone.now().date()
        form_data = {
            'selected_branch': None,
            'selected_gl_no': None,
            'selected_region': None,
            'selected_officer': None,
            'reporting_date': default_reporting_date,
            'exclude_ac_no_one': False,
            'include_non_zero': False,
        }
        grand_total = 0

        if request.method == 'POST':
            # Process form data
            reporting_date_str = request.POST.get('reporting_date')
            form_data['reporting_date'] = datetime.strptime(
                reporting_date_str, '%Y-%m-%d'
            ).date() if reporting_date_str else default_reporting_date

            form_data['exclude_ac_no_one'] = request.POST.get('exclude_ac_no_one') == 'on'
            form_data['include_non_zero'] = request.POST.get('include_non_zero') == 'on'

            # Apply filters - start with company_name base filter
            customers = Customer.objects.filter(branch__company_name=user_company_name)

            # Branch filter
            if branch_id := request.POST.get('branch'):
                try:
                    branch = Branch.objects.get(id=branch_id, company_name=user_company_name)
                    customers = customers.filter(branch=branch)
                    form_data['selected_branch'] = branch
                except Branch.DoesNotExist:
                    messages.warning(request, "Selected branch not found")

            # GL No filter
            if gl_no := request.POST.get('gl_no'):
                customers = customers.filter(gl_no=gl_no)
                form_data['selected_gl_no'] = gl_no

            # Region filter
            if region_id := request.POST.get('region'):
                try:
                    region = Region.objects.get(id=region_id, branch__company_name=user_company_name)
                    branch_ids = Branch.objects.filter(region=region, company_name=user_company_name).values_list('id', flat=True)
                    customers = customers.filter(branch__id__in=branch_ids)
                    form_data['selected_region'] = region
                except Region.DoesNotExist:
                    messages.warning(request, "Selected region not found")

            # Account Officer filter
            if officer_id := request.POST.get('credit_officer'):
                try:
                    officer = Account_Officer.objects.get(id=officer_id, region__branch__company_name=user_company_name)
                    branch_ids = Branch.objects.filter(account_officer=officer, company_name=user_company_name).values_list('id', flat=True)
                    customers = customers.filter(branch__id__in=branch_ids)
                    form_data['selected_officer'] = officer
                except Account_Officer.DoesNotExist:
                    messages.warning(request, "Selected account officer not found")

            # Process customers with balance filter
            gl_name_dict = dict(Account.objects.filter(branch__company_name=user_company_name).values_list('gl_no', 'gl_name'))

            for customer in customers:
                if form_data['exclude_ac_no_one'] and customer.ac_no == '1':
                    continue

                transactions = Memtrans.objects.filter(
                    gl_no=customer.gl_no,
                    ac_no=customer.ac_no,
                    app_date__lte=form_data['reporting_date'],
                    branch__company_name=user_company_name
                )

                account_balance = transactions.aggregate(total=Sum('amount'))['total'] or 0

                # Filtering logic
                if form_data['include_non_zero']:
                    if account_balance == 0:
                        continue   # skip zero balance if "include non-zero" is checked
                else:
                    if account_balance != 0:
                        continue   # skip non-zero if unchecked (default zero balance report)

                last_transaction_date = transactions.aggregate(max_date=Max('app_date'))['max_date'] or form_data['reporting_date']
                days_inactive = (form_data['reporting_date'] - last_transaction_date).days

                # Organize data by GL account
                if customer.gl_no not in customer_data:
                    customer_data[customer.gl_no] = {
                        'gl_name': gl_name_dict.get(customer.gl_no, 'Unknown'),
                        'customers': [],
                        'subtotal': 0,
                        'count': 0
                    }

                customer_data[customer.gl_no]['customers'].append({
                    'gl_no': customer.gl_no,
                    'ac_no': customer.ac_no,
                    'first_name': customer.first_name,
                    'middle_name': customer.middle_name,
                    'last_name': customer.last_name,
                    'address': getattr(customer, 'address', '-'),
                    'account_balance': account_balance,
                    'last_trx_date': last_transaction_date,
                    'days_without_activity': days_inactive,
                })

                customer_data[customer.gl_no]['subtotal'] += account_balance
                customer_data[customer.gl_no]['count'] += 1
                grand_total += account_balance

        context = {
            'branches': branches,
            'regions': regions,
            'account_officers': account_officers,
            'gl_accounts': gl_accounts,
            'customer_data': customer_data,
            'grand_total': grand_total,
            'current_datetime': current_datetime,
            'user_company_name': user_company_name,
            **form_data
        }

        return render(request, 'reports/savings_report/savings_zero_balance_report.html', context)

    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f"Error in savings_account_with_zero_balance: {str(e)}", exc_info=True)
        messages.error(request, f"An error occurred while generating the report: {str(e)}")
        return redirect('savings_account_with_zero_balance')





from django.shortcuts import render
from django.utils import timezone
from datetime import datetime
from django.db.models import Sum
from django.contrib import messages
import logging

def savings_account_overdrawn(request):
    logger = logging.getLogger(__name__)

    try:
        user = request.user
        user_branch = getattr(user, 'branch', None)
        user_branch_is_head_office = user_branch.head_office if user_branch else False
        user_company = user_branch.company if user_branch else None

        # Filter options: branches, regions, account officers
        if user_branch_is_head_office:
            branches = Branch.objects.filter(company=user_company).order_by('branch_name')
            regions = Region.objects.filter(branch__company=user_company)
            account_officers = Account_Officer.objects.filter(region__branch__company=user_company)
        elif user_branch:
            branches = Branch.objects.filter(id=user_branch.id)
            regions = Region.objects.filter(branch=user_branch)
            account_officers = Account_Officer.objects.filter(region__branch=user_branch)
        else:
            branches = Branch.objects.none()
            regions = Region.objects.none()
            account_officers = Account_Officer.objects.none()

        # GL accounts for the company
        gl_accounts = Account.objects.filter(
            gl_no__in=Customer.objects.filter(branch__company=user_company).values('gl_no').distinct(),
            branch__company=user_company
        )

        # Default values
        customer_data = {}
        default_reporting_date = timezone.now().date()
        form_data = {
            'selected_branch': None,
            'selected_gl_no': None,
            'selected_region': None,
            'selected_officer': None,
            'reporting_date': default_reporting_date,
            'exclude_ac_no_one': False,
        }
        grand_total = 0
        current_datetime = timezone.now()

        if request.method == 'POST':
            # Get form values
            reporting_date_str = request.POST.get('reporting_date')
            form_data['reporting_date'] = datetime.strptime(
                reporting_date_str, '%Y-%m-%d'
            ).date() if reporting_date_str else default_reporting_date

            form_data['exclude_ac_no_one'] = request.POST.get('exclude_ac_no_one') == 'on'

            # Base customer queryset
            customers = Customer.objects.filter(branch__company=user_company)

            # Apply branch filter
            if branch_id := request.POST.get('branch'):
                try:
                    branch = Branch.objects.get(id=branch_id, company=user_company)
                    customers = customers.filter(branch=branch)
                    form_data['selected_branch'] = branch
                except Branch.DoesNotExist:
                    messages.warning(request, "Selected branch not found")
            elif not user_branch_is_head_office and user_branch:
                customers = customers.filter(branch=user_branch)
                form_data['selected_branch'] = user_branch

            # GL No filter
            if gl_no := request.POST.get('gl_no'):
                customers = customers.filter(gl_no=gl_no)
                form_data['selected_gl_no'] = gl_no

            # Region filter
            if region_id := request.POST.get('region'):
                try:
                    region = Region.objects.get(id=region_id, branch__company=user_company)
                    branch_ids = Branch.objects.filter(region=region, company=user_company).values_list('id', flat=True)
                    customers = customers.filter(branch__id__in=branch_ids)
                    form_data['selected_region'] = region
                except Region.DoesNotExist:
                    messages.warning(request, "Selected region not found")

            # Account Officer filter
            if officer_id := request.POST.get('credit_officer'):
                try:
                    officer = Account_Officer.objects.get(id=officer_id, region__branch__company=user_company)
                    branch_ids = Branch.objects.filter(account_officer=officer, company=user_company).values_list('id', flat=True)
                    customers = customers.filter(branch__id__in=branch_ids)
                    form_data['selected_officer'] = officer
                except Account_Officer.DoesNotExist:
                    messages.warning(request, "Selected account officer not found")

            # GL name lookup dictionary
            gl_name_dict = dict(Account.objects.filter(branch__company=user_company).values_list('gl_no', 'gl_name'))

            # Process overdrawn accounts
            for customer in customers:
                if form_data['exclude_ac_no_one'] and customer.ac_no == '1':
                    continue

                transactions = Memtrans.objects.filter(
                    gl_no=customer.gl_no,
                    ac_no=customer.ac_no,
                    branch__company=user_company,
                    app_date__lte=form_data['reporting_date']
                )

                account_balance = transactions.aggregate(total=Sum('amount'))['total'] or 0
                if account_balance >= 0:
                    continue

                last_transaction = transactions.order_by('-ses_date').first()
                last_trx_date = last_transaction.ses_date if last_transaction else None

                if customer.gl_no not in customer_data:
                    customer_data[customer.gl_no] = {
                        'gl_name': gl_name_dict.get(customer.gl_no, 'Unknown'),
                        'customers': [],
                        'subtotal': 0,
                        'count': 0
                    }

                customer_data[customer.gl_no]['customers'].append({
                    'gl_no': customer.gl_no,
                    'ac_no': customer.ac_no,
                    'first_name': customer.first_name,
                    'middle_name': customer.middle_name,
                    'last_name': customer.last_name,
                    'address': getattr(customer, 'address', '-'),
                    'account_balance': account_balance,
                    'last_trx_date': last_trx_date,
                })

                customer_data[customer.gl_no]['subtotal'] += account_balance
                customer_data[customer.gl_no]['count'] += 1
                grand_total += account_balance

        context = {
            'branches': branches,
            'regions': regions,
            'account_officers': account_officers,
            'gl_accounts': gl_accounts,
            'customer_data': customer_data,
            'grand_total': grand_total,
            'current_datetime': current_datetime,
            'user_branch': user_branch,
            'user_branch_is_head_office': user_branch_is_head_office,
            **form_data
        }

        return render(request, 'reports/savings_report/savings_overdrawn_account_status.html', context)

    except Exception as e:
        logger.error(f"Error in savings_account_overdrawn: {str(e)}", exc_info=True)
        messages.error(request, f"An error occurred: {str(e)}")
        return render(request, 'reports/savings_report/savings_overdrawn_account_status.html', {
            'branches': Branch.objects.all().order_by('branch_name'),
            'regions': Region.objects.all(),
            'account_officers': Account_Officer.objects.all(),
            'gl_accounts': Account.objects.filter(gl_no__in=Customer.objects.values('gl_no').distinct()),
            'customer_data': {},
            'grand_total': 0,
            'current_datetime': timezone.now(),
            'user_branch': user_branch,
            'user_branch_is_head_office': user_branch_is_head_office,
            **form_data
        })


from django.shortcuts import render
from django.utils import timezone
from datetime import datetime
from django.db.models import Sum
from django.contrib import messages
import logging

def savings_interest_paid(request):
    logger = logging.getLogger(__name__)
    user = request.user
    user_branch = getattr(user, 'branch', None)  # Corrected here
    
    if not user_branch:
        return render(request, 'error.html', {'message': 'User has no branch assigned.'})
    
    user_company_name = user_branch.company_name
    is_head_office = user_branch.head_office

    try:
        # Initialize filter options
        if is_head_office:
            branches = Branch.objects.filter(company_name=user_company_name).order_by('branch_name')
            regions = Region.objects.filter(branch__company_name=user_company_name)
            account_officers = Account_Officer.objects.filter(region__branch__company_name=user_company_name)
        else:
            branches = Branch.objects.filter(id=user_branch.id)
            regions = Region.objects.filter(branch=user_branch)
            account_officers = Account_Officer.objects.filter(region__branch=user_branch)
        
        gl_accounts = Account.objects.filter(
            gl_no__in=Customer.objects.filter(branch__company_name=user_company_name).values('gl_no').distinct(),
            branch__company_name=user_company_name
        )

        # Default values
        customer_data = {}
        default_reporting_date = timezone.now().date()
        form_data = {
            'selected_branch': None,
            'selected_gl_no': None,
            'selected_region': None,
            'selected_officer': None,
            'reporting_date': default_reporting_date,
            'exclude_ac_no_one': False,
        }
        grand_total = 0
        current_datetime = timezone.now()

        if request.method == 'POST':
            reporting_date_str = request.POST.get('reporting_date')
            form_data['reporting_date'] = datetime.strptime(
                reporting_date_str, '%Y-%m-%d'
            ).date() if reporting_date_str else default_reporting_date

            form_data['exclude_ac_no_one'] = request.POST.get('exclude_ac_no_one') == 'on'

            customers = Customer.objects.filter(branch__company_name=user_company_name)

            # Apply branch filter
            branch_id = request.POST.get('branch')
            if branch_id:
                try:
                    branch = Branch.objects.get(id=branch_id, company_name=user_company_name)
                    customers = customers.filter(branch=branch.branch_code)
                    form_data['selected_branch'] = branch
                except Branch.DoesNotExist:
                    messages.warning(request, "Selected branch not found")
            elif not is_head_office and user_branch:
                customers = customers.filter(branch=user_branch.branch_code)
                form_data['selected_branch'] = user_branch

            # GL No filter
            if gl_no := request.POST.get('gl_no'):
                customers = customers.filter(gl_no=gl_no)
                form_data['selected_gl_no'] = gl_no

            # Region filter
            if region_id := request.POST.get('region'):
                try:
                    region = Region.objects.get(id=region_id, branch__company_name=user_company_name)
                    branch_codes = Branch.objects.filter(region=region, company_name=user_company_name).values_list('branch_code', flat=True)
                    customers = customers.filter(branch__in=branch_codes)
                    form_data['selected_region'] = region
                except Region.DoesNotExist:
                    messages.warning(request, "Selected region not found")

            # Account Officer filter
            if officer_id := request.POST.get('credit_officer'):
                try:
                    officer = Account_Officer.objects.get(id=officer_id, region__branch__company_name=user_company_name)
                    branch_codes = Branch.objects.filter(account_officer=officer, company_name=user_company_name).values_list('branch_code', flat=True)
                    customers = customers.filter(branch__in=branch_codes)
                    form_data['selected_officer'] = officer
                except Account_Officer.DoesNotExist:
                    messages.warning(request, "Selected account officer not found")

            # GL name lookup dictionary
            gl_name_dict = dict(Account.objects.filter(branch__company_name=user_company_name).values_list('gl_no', 'gl_name'))

            # Process interest payments
            for customer in customers:
                if form_data['exclude_ac_no_one'] and customer.ac_no == '1':
                    continue

                transactions = Memtrans.objects.filter(
                    gl_no=customer.gl_no,
                    ac_no=customer.ac_no,
                    code='MSI',
                    app_date__lte=form_data['reporting_date'],
                    branch__company_name=user_company_name
                )

                interest_paid = transactions.aggregate(total=Sum('amount'))['total'] or 0
                if interest_paid == 0:
                    continue

                last_transaction = transactions.order_by('-ses_date').first()
                last_trx_date = last_transaction.ses_date if last_transaction else None
                days_since_last = (form_data['reporting_date'] - last_trx_date).days if last_trx_date else None

                if customer.gl_no not in customer_data:
                    customer_data[customer.gl_no] = {
                        'gl_name': gl_name_dict.get(customer.gl_no, 'Unknown'),
                        'customers': [],
                        'subtotal': 0,
                        'count': 0
                    }

                customer_data[customer.gl_no]['customers'].append({
                    'gl_no': customer.gl_no,
                    'ac_no': customer.ac_no,
                    'first_name': customer.first_name,
                    'middle_name': customer.middle_name,
                    'last_name': customer.last_name,
                    'address': getattr(customer, 'address', '-'),
                    'interest_paid': interest_paid,
                    'last_interest_date': last_trx_date,
                    'days_since_last_interest': days_since_last,
                })

                customer_data[customer.gl_no]['subtotal'] += interest_paid
                customer_data[customer.gl_no]['count'] += 1
                grand_total += interest_paid

        context = {
            'branches': branches,
            'regions': regions,
            'account_officers': account_officers,
            'gl_accounts': gl_accounts,
            'customer_data': customer_data,
            'grand_total': grand_total,
            'current_datetime': current_datetime,
            'user_branch': user_branch,
            'user_branch_is_head_office': is_head_office,
            'user_company_name': user_company_name,
            **form_data
        }

        return render(request, 'reports/savings_report/savings_interest_paid.html', context)

    except Exception as e:
        logger.error(f"Error in savings_interest_paid: {str(e)}", exc_info=True)
        messages.error(request, f"An error occurred while generating the interest paid report: {str(e)}")
        return render(request, 'reports/savings_report/savings_interest_paid.html', {
            'branches': Branch.objects.filter(company_name=user_company_name).order_by('branch_name'),
            'regions': Region.objects.filter(branch__company_name=user_company_name),
            'account_officers': Account_Officer.objects.filter(region__branch__company_name=user_company_name),
            'gl_accounts': Account.objects.filter(
                gl_no__in=Customer.objects.filter(branch__company_name=user_company_name).values('gl_no').distinct(),
                branch__company_name=user_company_name
            ),
            'customer_data': {},
            'grand_total': 0,
            'current_datetime': timezone.now(),
            'user_branch': user_branch,
            'user_branch_is_head_office': is_head_office,
            'user_company_name': user_company_name,
            **form_data
        })




def savings_account_credit_balance(request):
    logger = logging.getLogger(__name__)
    user = request.user

    # Get the user's branch
    user_branch = getattr(user, 'branch', None)
    if not user_branch:
        return render(request, 'error.html', {'message': 'User has no branch assigned.'})
    user_company_name = user_branch.company_name

    try:
        # Determine if user is head office
        is_head_office = user_branch.head_office

        # Initialize filter options
        if is_head_office:
            branches = Branch.objects.filter(company_name=user_company_name).order_by('branch_name')
            regions = Region.objects.filter(branch__company_name=user_company_name)
            account_officers = Account_Officer.objects.filter(region__branch__company_name=user_company_name)
        else:
            # Non-head-office users: only their branch
            branches = Branch.objects.filter(id=user_branch.id)
            regions = Region.objects.filter(branch=user_branch)
            account_officers = Account_Officer.objects.filter(region__branch=user_branch)

        gl_accounts = Account.objects.filter(
            gl_no__in=Customer.objects.filter(branch__company_name=user_company_name).values('gl_no').distinct(),
            branch__company_name=user_company_name
        )
        current_datetime = timezone.now()

        # Default values
        customer_data = {}
        default_reporting_date = timezone.now().date()
        form_data = {
            'selected_branch': None,
            'selected_gl_no': None,
            'selected_region': None,
            'selected_officer': None,
            'reporting_date': default_reporting_date,
            'exclude_ac_no_one': False,
        }
        grand_total = 0

        if request.method == 'POST':
            # Process form data
            reporting_date_str = request.POST.get('reporting_date')
            form_data['reporting_date'] = datetime.strptime(
                reporting_date_str, '%Y-%m-%d'
            ).date() if reporting_date_str else default_reporting_date
            form_data['exclude_ac_no_one'] = request.POST.get('exclude_ac_no_one') == 'on'

            # Base customer queryset
            customers = Customer.objects.filter(branch__company_name=user_company_name)

            # Branch filter
            if branch_id := request.POST.get('branch'):
                try:
                    branch = Branch.objects.get(id=branch_id, company_name=user_company_name)
                    customers = customers.filter(branch=branch)
                    form_data['selected_branch'] = branch
                except Branch.DoesNotExist:
                    messages.warning(request, "Selected branch not found")
            elif not is_head_office:
                # Non-head-office user: force filter to their branch
                customers = customers.filter(branch=user_branch)
                form_data['selected_branch'] = user_branch

            # GL No filter
            if gl_no := request.POST.get('gl_no'):
                customers = customers.filter(gl_no=gl_no)
                form_data['selected_gl_no'] = gl_no

            # Region filter
            if region_id := request.POST.get('region'):
                try:
                    region = Region.objects.get(id=region_id, branch__company_name=user_company_name)
                    branch_ids = Branch.objects.filter(region=region, company_name=user_company_name).values_list('id', flat=True)
                    customers = customers.filter(branch__id__in=branch_ids)
                    form_data['selected_region'] = region
                except Region.DoesNotExist:
                    messages.warning(request, "Selected region not found")

            # Account Officer filter
            if officer_id := request.POST.get('credit_officer'):
                try:
                    officer = Account_Officer.objects.get(id=officer_id, region__branch__company_name=user_company_name)
                    branch_ids = Branch.objects.filter(account_officer=officer, company_name=user_company_name).values_list('id', flat=True)
                    customers = customers.filter(branch__id__in=branch_ids)
                    form_data['selected_officer'] = officer
                except Account_Officer.DoesNotExist:
                    messages.warning(request, "Selected account officer not found")

            # Process accounts with credit balance (>0)
            gl_name_dict = dict(Account.objects.filter(branch__company_name=user_company_name).values_list('gl_no', 'gl_name'))

            for customer in customers:
                if form_data['exclude_ac_no_one'] and customer.ac_no == '1':
                    continue

                transactions = Memtrans.objects.filter(
                    gl_no=customer.gl_no,
                    ac_no=customer.ac_no,
                    app_date__lte=form_data['reporting_date'],
                    branch__company_name=user_company_name
                )

                # Calculate account balance
                account_balance = transactions.aggregate(total=Sum('amount'))['total'] or 0

                # Skip accounts with balance <= 0
                if account_balance <= 0:
                    continue

                # Last transaction date
                last_transaction = transactions.order_by('-ses_date').first()
                last_trx_date = last_transaction.ses_date if last_transaction else None

                # Organize by GL account
                if customer.gl_no not in customer_data:
                    customer_data[customer.gl_no] = {
                        'gl_name': gl_name_dict.get(customer.gl_no, 'Unknown'),
                        'customers': [],
                        'subtotal': 0,
                        'count': 0
                    }

                customer_data[customer.gl_no]['customers'].append({
                    'gl_no': customer.gl_no,
                    'ac_no': customer.ac_no,
                    'first_name': customer.first_name,
                    'middle_name': customer.middle_name,
                    'last_name': customer.last_name,
                    'address': getattr(customer, 'address', '-'),
                    'account_balance': account_balance,
                    'last_trx_date': last_trx_date,
                })

                customer_data[customer.gl_no]['subtotal'] += account_balance
                customer_data[customer.gl_no]['count'] += 1
                grand_total += account_balance

        context = {
            'branches': branches,
            'regions': regions,
            'account_officers': account_officers,
            'gl_accounts': gl_accounts,
            'customer_data': customer_data,
            'grand_total': grand_total,
            'current_datetime': current_datetime,
            'user_company_name': user_company_name,
            'user_branch': user_branch,  # important for template
            'is_head_office': is_head_office,
            **form_data
        }

        return render(request, 'reports/savings_report/savings_account_with_credit_balance_report.html', context)

    except Exception as e:
        logger.error(f"Error in savings_account_credit_balance: {str(e)}", exc_info=True)
        messages.error(request, f"An error occurred while generating the credit balance report: {str(e)}")
        return render(request, 'reports/savings_report/savings_account_with_credit_balance_report.html', {
            'branches': Branch.objects.filter(company_name=user_company_name).order_by('branch_name'),
            'regions': Region.objects.filter(branch__company_name=user_company_name),
            'account_officers': Account_Officer.objects.filter(region__branch__company_name=user_company_name),
            'gl_accounts': Account.objects.filter(
                gl_no__in=Customer.objects.filter(branch__company_name=user_company_name).values('gl_no').distinct(),
                branch__company_name=user_company_name
            ),
            'current_datetime': timezone.now(),
            'user_company_name': user_company_name,
            'user_branch': user_branch,
            'is_head_office': is_head_office
        })

# from django.shortcuts import render
# from django.utils import timezone
# from django.db.models import Sum, Min
# from django.contrib import messages
# from datetime import datetime

# def savings_overdrawn_account_status(request):
#     user = request.user

#     # Get the user's branch
#     user_branch = getattr(user, 'branch', None)
#     if not user_branch:
#         return render(request, 'error.html', {'message': 'User has no branch assigned.'})

#     user_company = user_branch.company

#     # Determine if user is head office
#     is_head_office = user_branch.head_office

#     # Filter options
#     if is_head_office:
#         branches = Branch.objects.filter(company=user_company).order_by('branch_name')
#         regions = Region.objects.filter(branch__company=user_company)
#         account_officers = Account_Officer.objects.filter(region__branch__company=user_company)
#     else:
#         # Non-head-office users can only see their branch
#         branches = Branch.objects.filter(id=user_branch.id)
#         regions = Region.objects.filter(branch=user_branch)
#         account_officers = Account_Officer.objects.filter(region__branch=user_branch)

#     gl_accounts = Account.objects.filter(
#         gl_no__in=Customer.objects.filter(branch__company=user_company).values('gl_no').distinct(),
#         branch__company=user_company
#     )

#     # Default reporting date
#     reporting_date = timezone.now().date()
#     customer_data = {}
#     grand_total = 0
#     selected_branch = None
#     selected_gl_no = None
#     selected_region = None
#     selected_officer = None
#     exclude_ac_no_one = False
#     current_datetime = timezone.now()

#     if request.method == 'POST':
#         # Get form data
#         reporting_date_str = request.POST.get('reporting_date')
#         branch_id = request.POST.get('branch')
#         gl_no = request.POST.get('gl_no')
#         region_id = request.POST.get('region')
#         officer_id = request.POST.get('credit_officer')
#         exclude_ac_no_one = request.POST.get('exclude_ac_no_one') == 'on'

#         if reporting_date_str:
#             reporting_date = datetime.strptime(reporting_date_str, '%Y-%m-%d').date()

#         # Base queryset filtered by company
#         customers = Customer.objects.filter(branch__company=user_company)

#         # Branch filter
#         if branch_id:
#             try:
#                 selected_branch = Branch.objects.get(id=branch_id, company=user_company)
#                 customers = customers.filter(branch=selected_branch)
#             except Branch.DoesNotExist:
#                 messages.warning(request, "Selected branch not found")
#         elif not is_head_office:
#             # Non-head-office users can only access their branch
#             customers = customers.filter(branch=user_branch)
#             selected_branch = user_branch

#         # GL filter
#         if gl_no:
#             customers = customers.filter(gl_no=gl_no)
#             selected_gl_no = gl_no

#         # Region filter
#         if region_id:
#             try:
#                 selected_region = Region.objects.get(id=region_id, branch__company=user_company)
#                 branch_ids = Branch.objects.filter(region=selected_region, company=user_company).values_list('id', flat=True)
#                 customers = customers.filter(branch__id__in=branch_ids)
#             except Region.DoesNotExist:
#                 messages.warning(request, "Selected region not found")

#         # Account officer filter
#         if officer_id:
#             try:
#                 selected_officer = Account_Officer.objects.get(id=officer_id, region__branch__company=user_company)
#                 branch_ids = Branch.objects.filter(account_officer=selected_officer, company=user_company).values_list('id', flat=True)
#                 customers = customers.filter(branch__id__in=branch_ids)
#             except Account_Officer.DoesNotExist:
#                 messages.warning(request, "Selected account officer not found")

#         # GL name dictionary
#         gl_name_dict = dict(Account.objects.filter(branch__company=user_company).values_list('gl_no', 'gl_name'))

#         # Process customers
#         for customer in customers:
#             if exclude_ac_no_one and customer.ac_no == '1':
#                 continue

#             transactions = Memtrans.objects.filter(
#                 gl_no=customer.gl_no,
#                 ac_no=customer.ac_no,
#                 branch__company=user_company
#             )

#             transactions = transactions.filter(app_date__lte=reporting_date)
#             account_balance = transactions.aggregate(total=Sum('amount'))['total'] or 0

#             if account_balance >= 0:
#                 continue

#             last_transaction = transactions.order_by('-ses_date').first()
#             last_trx_date = last_transaction.ses_date if last_transaction else None

#             if customer.gl_no not in customer_data:
#                 customer_data[customer.gl_no] = {
#                     'gl_name': gl_name_dict.get(customer.gl_no, 'Unknown'),
#                     'customers': [],
#                     'subtotal': 0,
#                     'count': 0
#                 }

#             customer_data[customer.gl_no]['customers'].append({
#                 'gl_no': customer.gl_no,
#                 'ac_no': customer.ac_no,
#                 'first_name': customer.first_name,
#                 'middle_name': customer.middle_name,
#                 'last_name': customer.last_name,
#                 'address': getattr(customer, 'address', '-'),
#                 'account_balance': account_balance,
#                 'last_trx_date': last_trx_date,
#             })

#             customer_data[customer.gl_no]['subtotal'] += account_balance
#             customer_data[customer.gl_no]['count'] += 1
#             grand_total += account_balance

#     context = {
#         'branches': branches,
#         'regions': regions,
#         'account_officers': account_officers,
#         'gl_accounts': gl_accounts,
#         'customer_data': customer_data,
#         'selected_branch': selected_branch,
#         'selected_gl_no': selected_gl_no,
#         'selected_region': selected_region,
#         'selected_officer': selected_officer,
#         'reporting_date': reporting_date,
#         'exclude_ac_no_one': exclude_ac_no_one,
#         'grand_total': grand_total,
#         'current_datetime': current_datetime,
#         'user_branch': user_branch,
#         'is_head_office': is_head_office,
#         'user_company': user_company,
#     }

#     return render(request, 'reports/savings_report/savings_overdrawn_account_status.html', context)




from django.shortcuts import render
from transactions.models import Memtrans
from accounts.models import Account
from company.models import Company
from datetime import datetime
from collections import defaultdict
from django.http import JsonResponse


from collections import defaultdict
from django.shortcuts import render, get_object_or_404

from .forms import TrialBalanceForm

def generate_balance_sheet(start_date, end_date, branch_code=None):
    """
    Generate balance sheet data for a given date range and optionally filter by branch.
    """
    # Query all relevant Memtrans entries for the given branch
    memtrans_entries = Memtrans.objects.filter(ses_date__range=[start_date, end_date], error='A')

    # Apply branch filtering only if a specific branch code is provided
    if branch_code:
        memtrans_entries = memtrans_entries.filter(branch=branch_code)

    # Initialize a dictionary to hold GL account balances
    gl_customer_balance = defaultdict(lambda: {'debit': 0, 'credit': 0, 'balance': 0})

    # Process each entry
    for entry in memtrans_entries:
        gl_no = entry.gl_no
        amount = entry.amount

        if entry.type == 'N':
            if amount < 0:
                gl_customer_balance[gl_no]['debit'] += abs(amount)
            else:
                gl_customer_balance[gl_no]['credit'] += amount
        else:
            if amount < 0:
                gl_customer_balance[gl_no]['credit'] += abs(amount)
            else:
                gl_customer_balance[gl_no]['debit'] += amount

    # Fetch all necessary accounts in a single query to improve performance
    accounts = Account.objects.filter(gl_no__in=gl_customer_balance.keys()).values('gl_no', 'gl_name')
    account_map = {account['gl_no']: account['gl_name'] for account in accounts}

    # Sort GL numbers and prepare balance sheet data
    sorted_keys = sorted(gl_customer_balance.keys())
    sorted_balance_sheet_data = []
    subtotal_4 = subtotal_5 = 0

    for gl_no in sorted_keys:
        if gl_no.startswith("4") or gl_no.startswith("5"):
            amount = gl_customer_balance[gl_no]['debit'] - gl_customer_balance[gl_no]['credit']
            if gl_no.startswith("4"):
                subtotal_4 += amount
            else:
                subtotal_5 += amount
        else:
            balance_data = gl_customer_balance[gl_no]
            debit = balance_data['debit']
            credit = balance_data['credit']
            balance = debit - credit
            gl_name = account_map.get(gl_no, '')  # Default to empty string if GL number not found
            balance_data.update({'gl_no': gl_no, 'gl_name': gl_name, 'balance': balance})
            sorted_balance_sheet_data.append(balance_data)

    # Calculate net income, total debit, credit, and balance
    net_income = subtotal_4 + subtotal_5
    total_debit = sum(entry['debit'] for entry in sorted_balance_sheet_data)
    total_credit = sum(entry['credit'] for entry in sorted_balance_sheet_data)
    total_balance = total_credit - total_debit

    return sorted_balance_sheet_data, subtotal_4, subtotal_5, total_debit, total_credit, total_balance, net_income



def balance_sheet(request):
    user = request.user

    # Get user's branch
    user_branch = getattr(user, 'branch', None)
    if not user_branch:
        return render(request, 'error.html', {'message': 'User has no branch assigned.'})

    # Determine if user is head office
    is_head_office = getattr(user_branch, 'head_office', False)

    # Branch list: head office sees all, others see only their branch
    if is_head_office:
        branches = Branch.objects.all().order_by('branch_name')
    else:
        branches = Branch.objects.filter(id=user_branch.id)

    if request.method == 'POST':
        form = TrialBalanceForm(request.POST)
        if form.is_valid():
            start_date = form.cleaned_data['start_date']
            end_date = form.cleaned_data['end_date']

            selected_branch_obj = form.cleaned_data['branch'] if form.cleaned_data['branch'] else None
            selected_branch_id = selected_branch_obj.id if selected_branch_obj else None

            # Detect "all branches"
            is_all_branches = True if not selected_branch_obj else False

            # ⚠️ Generate report: pass branch id or None for consolidated
            branch_filter = selected_branch_obj.id if selected_branch_obj else None

            balance_sheet_data, subtotal_4, subtotal_5, total_debit, total_credit, total_balance, net_income = generate_balance_sheet(
                start_date, end_date, branch_filter
            )

            context = {
                'form': form,
                'branches': branches,
                'balance_sheet_data': balance_sheet_data,
                'subtotal_4': subtotal_4,
                'subtotal_5': subtotal_5,
                'total_debit': total_debit,
                'total_credit': total_credit,
                'total_balance': total_balance,
                'net_income': net_income,
                'start_date': start_date,
                'end_date': end_date,
                'selected_branch': selected_branch_id,
                'selected_branch_obj': selected_branch_obj,
                'is_all_branches': is_all_branches,  # ✅ NEW
                'user_branch': user_branch,
                'is_head_office': is_head_office,
            }

            return render(request, 'reports/financials/balance_sheet.html', context)

    else:
        form = TrialBalanceForm()

    return render(request, 'reports/financials/balance_sheet.html', {
        'form': form,
        'branches': branches,
        'selected_branch': None,
        'selected_branch_obj': None,
        'is_all_branches': False,
        'user_branch': user_branch,
        'is_head_office': is_head_office,
    })




from django.shortcuts import render, get_object_or_404
from collections import defaultdict
from .forms import TrialBalanceForm
from transactions.models import Memtrans
from accounts.models import Account
from company.models import Company

def generate_profit_and_loss(start_date, end_date, branch_code=None):
    """
    Generate profit and loss statement data for a given date range and optionally filter by branch.
    Only include GL numbers starting with 4 (revenue) or 5 (expenses).
    """
    # Query all relevant Memtrans entries for the given date range and branch
    memtrans_entries = Memtrans.objects.filter(ses_date__range=[start_date, end_date], error='A')

    # Apply branch filtering only if a specific branch code is provided
    if branch_code:
        memtrans_entries = memtrans_entries.filter(branch=branch_code)

    # Initialize a dictionary to hold GL account balances
    gl_customer_balance = defaultdict(lambda: {'debit': 0, 'credit': 0, 'balance': 0})

    # Process each entry
    for entry in memtrans_entries:
        gl_no = entry.gl_no

        # Filter GL numbers starting with 4 or 5
        if gl_no.startswith('4') or gl_no.startswith('5'):
            amount = entry.amount

            # Determine the type of entry and update the balances accordingly
            if entry.type == 'N':
                if amount < 0:
                    gl_customer_balance[gl_no]['debit'] += amount  # Use amount directly to retain sign
                else:
                    gl_customer_balance[gl_no]['credit'] += amount
            else:
                if amount < 0:
                    gl_customer_balance[gl_no]['credit'] += amount  # Use amount directly to retain sign
                else:
                    gl_customer_balance[gl_no]['debit'] += amount

    # Fetch all necessary accounts in a single query to improve performance
    accounts = Account.objects.filter(gl_no__in=gl_customer_balance.keys()).values('gl_no', 'gl_name')
    account_map = {account['gl_no']: account['gl_name'] for account in accounts}

    # Prepare profit and loss data
    sorted_profit_and_loss_data = []
    subtotal_4 = subtotal_5 = 0

    for gl_no, balance_data in gl_customer_balance.items():
        debit = balance_data['debit']
        credit = balance_data['credit']
        # Calculate balance by adding debit and credit
        balance = debit + credit
        gl_name = account_map.get(gl_no, '')  # Default to empty string if GL number not found
        balance_data.update({'gl_no': gl_no, 'gl_name': gl_name, 'balance': balance})

        # Calculate subtotals for revenues and expenses
        if gl_no.startswith("4"):
            subtotal_4 += balance
        elif gl_no.startswith("5"):
            subtotal_5 += balance

        sorted_profit_and_loss_data.append(balance_data)

    # Calculate total debit, credit, and net income
    total_debit = sum(entry['debit'] for entry in sorted_profit_and_loss_data)
    total_credit = sum(entry['credit'] for entry in sorted_profit_and_loss_data)
    net_income = subtotal_4 + subtotal_5

    return sorted_profit_and_loss_data, subtotal_4, subtotal_5, total_debit, total_credit, net_income

def profit_and_loss(request):
    branches = Branch.objects.all()  # Changed from Company to Branch

    if request.method == 'POST':
        form = TrialBalanceForm(request.POST)
        if form.is_valid():
            start_date = form.cleaned_data['start_date']
            end_date = form.cleaned_data['end_date']
            branch_id = form.cleaned_data['branch'].id if form.cleaned_data['branch'] else None

            # Determine branch code if a specific branch is selected
            branch_code = None
            if branch_id is not None:
                branch = get_object_or_404(Branch, id=branch_id)  # Changed from Company to Branch
                branch_code = branch.branch_code

            # Generate profit and loss data
            profit_and_loss_data, subtotal_4, subtotal_5, total_debit, total_credit, net_income = generate_profit_and_loss(
                start_date, end_date, branch_code
            )

            return render(request, 'reports/financials/profit_and_loss.html', {
                'form': form,
                'branches': branches,  # Changed from companies to branches
                'balance_sheet_data': profit_and_loss_data,
                'subtotal_4': subtotal_4,
                'subtotal_5': subtotal_5,
                'total_debit': total_debit,
                'total_credit': total_credit,
                'net_income': net_income,
                'start_date': start_date,
                'end_date': end_date,
                'selected_branch': branch_id,
                'branch': branch if branch_id else branches.first(),  # Changed from company to branch
            })
    else:
        form = TrialBalanceForm()

    return render(request, 'reports/financials/profit_and_loss.html', {
        'form': form,
        'branches': branches,  # Changed from companies to branches
        'selected_branch': None,
    })

    
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from loans.models import LoanHist

def loan_hist(request):
    if request.method == "POST":
        id = request.POST.get('id')
        loanhist_entry = get_object_or_404(LoanHist, id=id)
        loanhist_entry.delete()
        messages.success(request, 'LoanHist entry deleted successfully!')
        return redirect('loan_hist')  # Assuming 'loan_hist' is the name of the URL pattern for this view

    gl_no = request.GET.get('gl_no')
    ac_no = request.GET.get('ac_no')
    cycle = request.GET.get('cycle')

    loan_hist = LoanHist.objects.all()

    if gl_no:
        loan_hist = loan_hist.filter(gl_no=gl_no)
    if ac_no:
        loan_hist = loan_hist.filter(ac_no=ac_no)
    if cycle:
        loan_hist = loan_hist.filter(cycle=cycle)

    return render(request, 'reports/loans/loan_hist.html', {
        'loan_hist': loan_hist,
        'gl_no': gl_no,
        'ac_no': ac_no,
        'cycle': cycle,
    })




@login_required(login_url='login')
@user_passes_test(check_role_admin)
def loan_report(request):
    return render(request, 'reports/loans/loan_report.html')





from django.shortcuts import render, get_object_or_404
from company.models import Company
from reports.forms import LoanLedgerCardForm  # Ensure this form is correctly defined

from django.shortcuts import render
from django.db.models import Sum

def loan_ledger_card_view(request):
    branches = Branch.objects.all()
    accounts = Account.objects.filter(gl_no__startswith="104")
    error_message = None

    if request.method == 'POST':
        form = LoanLedgerCardForm(request.POST)  # ✅ no user_branch
        if form.is_valid():
            branch = form.cleaned_data['branch']
            account = form.cleaned_data['account']
            ac_no = form.cleaned_data.get('ac_no')
            cycle = form.cleaned_data.get('cycle')

            try:
                loan = Loans.objects.get(
                    branch=branch, gl_no=account.gl_no, ac_no=ac_no, cycle=cycle
                )
                disbursement_amount = loan.loan_amount
                total_interest = LoanHist.objects.filter(
                    branch=branch, gl_no=account.gl_no, ac_no=ac_no, cycle=cycle, trx_type='LD'
                ).aggregate(Sum('interest'))['interest__sum'] or 0
                disbursement_date = loan.disbursement_date
                num_installments = loan.num_install
                loan_officer = loan.loan_officer
                annual_interest_rate = loan.interest_rate

                customer = Customer.objects.get(
                    branch=branch, gl_no=account.gl_no, ac_no=ac_no
                )

                principal_balance = disbursement_amount
                interest_balance = total_interest
                ledger_card = LoanHist.objects.filter(
                    branch=branch, gl_no=account.gl_no, ac_no=ac_no, cycle=cycle
                ).order_by('trx_date')

                total_payment = 0
                penalty_balance = 0

                for entry in ledger_card:
                    total_payment += entry.principal + entry.interest + entry.penalty
                    if entry.trx_type == 'LP':
                        principal_balance += entry.principal
                        interest_balance += entry.interest

                    penalty_balance += entry.penalty
                    total_balance = principal_balance + interest_balance + penalty_balance

                    entry.total_payment = total_payment
                    entry.principal_balance = principal_balance
                    entry.interest_balance = interest_balance
                    entry.penalty_balance = penalty_balance
                    entry.total_balance = total_balance

                return render(request, 'reports/loans/loan_ledger_card.html', {
                    'form': form,
                    'ledger_card': ledger_card,
                    'total_payment': total_payment,
                    'principal_balance': principal_balance,
                    'interest_balance': interest_balance,
                    'penalty_balance': penalty_balance,
                    'total_balance': total_balance,
                    'customer': customer,
                    'loan': loan,
                    'disbursement_amount': disbursement_amount,
                    'disbursement_date': disbursement_date,
                    'num_installments': num_installments,
                    'loan_officer': loan_officer,
                    'annual_interest_rate': annual_interest_rate,
                    'branches': branches,
                    'accounts': accounts,
                    'branch': branch,
                    'form_submitted': True,
                })

            except Loans.DoesNotExist:
                error_message = "⚠️ No loan record found for the provided criteria."
            except Customer.DoesNotExist:
                error_message = "⚠️ No customer record found for the provided criteria."

    else:
        form = LoanLedgerCardForm()

    return render(request, 'reports/loans/loan_ledger_card.html', {
        'form': form,
        'branches': branches,
        'accounts': accounts,
        'form_submitted': False,
        'error_message': error_message,
    })

from django.shortcuts import render
from django.db.models import Sum
from reports.forms import LoanLedgerCardForm
from company.models import Branch
from accounts_admin.models import Account
from loans.models import Loans, LoanHist, Customer
from datetime import date
# views.py


# views.py



from datetime import date
from django.db.models import Sum
from django.shortcuts import render
from .forms import LoanLedgerCardForm
from company.models import Branch
from accounts_admin.models import Account
from loans.models import Loans, LoanHist
from customers.models import Customer


def loan_repayment_schedule(request):
    """
    Full, robust view that:
      - handles missing loan/customer gracefully
      - computes running balances per ledger entry
      - computes grand totals and final balances
      - provides safe default values for form inputs (ac_no_value, cycle_value)
    """

    # Basic lookups and safe defaults
    branches = Branch.objects.all()
    accounts = Account.objects.filter(gl_no__startswith="104")
    ledger_card = []
    error_message = None
    selected_branch = None

    # Context-safe defaults
    branch_name = ""
    company_name = ""
    current_date = date.today()

    loan = None
    customer = None
    customer_name = ""

    # Loan detail defaults
    disbursement_amount = 0
    loan_period = 0
    num_installments = 0
    disbursement_date = None
    annual_interest_rate = 0

    # Totals and final balances defaults
    grand_total_principal = 0
    grand_total_interest = 0
    grand_total_penalty = 0
    total_paid_sum = 0
    final_principal_balance = 0
    final_interest_balance = 0
    final_penalty_balance = 0
    final_total_balance = 0
    grand_total_disbursed = 0
    outstanding_amount = 0

    # Safe default values for input fields (pre-fill)
    ac_no_value = ""
    cycle_value = ""

    if request.method == "POST":
        form = LoanLedgerCardForm(request.POST)
        if form.is_valid():
            branch = form.cleaned_data["branch"]
            account = form.cleaned_data["account"]
            ac_no = form.cleaned_data.get("ac_no")
            cycle = form.cleaned_data.get("cycle")
            start_date = form.cleaned_data.get("start_date")
            end_date = form.cleaned_data.get("end_date")

            # set safe prefilled values for the form inputs
            ac_no_value = ac_no or ""
            cycle_value = cycle or ""

            # set branch/company info for header
            selected_branch = branch
            branch_name = getattr(branch, "branch_name", "") or ""
            company_name = getattr(getattr(branch, "company", None), "company_name", "") or ""

            try:
                # find loan (may raise Loans.DoesNotExist)
                loan = Loans.objects.get(branch=branch, gl_no=account.gl_no, ac_no=ac_no, cycle=cycle)

                # loan details
                disbursement_amount = loan.loan_amount or 0
                loan_period = loan.num_install or 0
                num_installments = loan.num_install or 0
                disbursement_date = loan.disbursement_date
                annual_interest_rate = loan.interest_rate or 0

                # customer (may raise Customer.DoesNotExist)
                try:
                    customer = Customer.objects.get(branch=branch, gl_no=account.gl_no, ac_no=ac_no)
                    customer_name = " ".join(
                        filter(None, [customer.first_name, customer.middle_name, customer.last_name])
                    )
                except Customer.DoesNotExist:
                    customer = None
                    customer_name = ""
                    error_message = "No customer record found for the provided criteria."

                # fetch ledger entries
                ledger_entries = LoanHist.objects.filter(
                    branch=branch, gl_no=account.gl_no, ac_no=ac_no, cycle=cycle
                )
                if start_date:
                    ledger_entries = ledger_entries.filter(trx_date__gte=start_date)
                if end_date:
                    ledger_entries = ledger_entries.filter(trx_date__lte=end_date)

                ledger_entries = ledger_entries.order_by("trx_date")

                # initial running balances
                principal_balance = disbursement_amount
                disbursed_interest = ledger_entries.filter(trx_type="LD").aggregate(total=Sum("interest"))["total"] or 0
                interest_balance = disbursed_interest
                penalty_balance = 0
                total_payment_running = 0

                ledger_card = []
                counter = 1
                for entry in ledger_entries:
                    # guard None decimals
                    entry_principal = entry.principal or 0
                    entry_interest = entry.interest or 0
                    entry_penalty = entry.penalty or 0

                    if entry.trx_type == "LP":  # Loan Payment
                        principal_balance = max(0, principal_balance - entry_principal)
                        interest_balance = max(0, interest_balance - entry_interest)
                        total_payment_running += entry_principal + entry_interest + entry_penalty
                    else:
                        # For other trx types, we still accumulate penalty and leave principal/interest adjustments aside
                        total_payment_running += entry_principal + entry_interest + entry_penalty

                    penalty_balance += entry_penalty
                    total_balance = principal_balance + interest_balance + penalty_balance

                    # attach computed fields for template
                    entry.counter = counter
                    entry.principal_balance = principal_balance
                    entry.interest_balance = interest_balance
                    entry.penalty_balance = penalty_balance
                    entry.total_balance = total_balance
                    entry.total_payment = total_payment_running
                    entry.customer_name = customer_name

                    ledger_card.append(entry)
                    counter += 1

                # compute grand totals from ledger_card
                grand_total_principal = sum((e.principal or 0) for e in ledger_card) if ledger_card else 0
                grand_total_interest = sum((e.interest or 0) for e in ledger_card) if ledger_card else 0
                grand_total_penalty = sum((e.penalty or 0) for e in ledger_card) if ledger_card else 0
                total_paid_sum = sum(( (e.principal or 0) + (e.interest or 0) + (e.penalty or 0)) for e in ledger_card) if ledger_card else 0

                # final balances (from running balances or from last ledger row)
                final_principal_balance = principal_balance
                final_interest_balance = interest_balance
                final_penalty_balance = penalty_balance
                final_total_balance = final_principal_balance + final_interest_balance + final_penalty_balance

                # totals disbursed/outstanding
                grand_total_disbursed = (disbursement_amount or 0) + (disbursed_interest or 0)
                outstanding_amount = final_total_balance

            except Loans.DoesNotExist:
                loan = None
                error_message = "No loan record found for the provided criteria."
        else:
            error_message = "Please correct the errors on the form."
    else:
        form = LoanLedgerCardForm()

    # Ensure form variable exists even for GET
    if request.method != "POST":
        form = LoanLedgerCardForm()

    # Prepare context and render
    context = {
        "form": form,
        "branches": branches,
        "accounts": accounts,
        "form_submitted": request.method == "POST",
        "error_message": error_message,
        "selected_branch": selected_branch,
        "branch_name": branch_name,
        "company_name": company_name,
        "current_date": current_date,
        "loan": loan,
        "customer": customer,
        "customer_name": customer_name,
        "ledger_card": ledger_card,
        # totals & balances
        "grand_total_principal": grand_total_principal,
        "grand_total_interest": grand_total_interest,
        "grand_total_penalty": grand_total_penalty,
        "total_paid_sum": total_paid_sum,
        "final_principal_balance": final_principal_balance,
        "final_interest_balance": final_interest_balance,
        "final_penalty_balance": final_penalty_balance,
        "final_total_balance": final_total_balance,
        "grand_total_disbursed": grand_total_disbursed,
        "outstanding_amount": outstanding_amount,
        # loan detail fields
        "disbursement_amount": disbursement_amount,
        "loan_period": loan_period,
        "num_installments": num_installments,
        "disbursement_date": disbursement_date,
        "annual_interest_rate": annual_interest_rate,
        # safe prefill values for inputs
        "ac_no_value": ac_no_value,
        "cycle_value": cycle_value,
    }

    return render(request, "reports/loans/loan_repayment_schedule.html", context)



from .forms import LoanDisbursementReportForm

from django.shortcuts import render

from django.shortcuts import render
from django.db.models import Sum
from .forms import LoanDisbursementReportForm
from loans.models import Loans
from company.models import Company

from django.shortcuts import render
from django.utils import timezone
from django.db.models import Sum
from django.contrib.auth.decorators import login_required
from .forms import LoanDisbursementReportForm
from loans.models import Loans
from company.models import Branch

@login_required

def loan_disbursement_report(request):
    # Get current user's branch (default)
    user_branch = request.user.branch

    # Initialize form with user's branch
    form = LoanDisbursementReportForm(
        request.POST or None,
        user_branch=None,   # 🔹 force form to load ALL branches
        initial={'reporting_date': date.today()}
    )


    # Initialize context
    context = {
        'form': form,
        'expected_repayments': [],
        'branches': Branch.objects.all(),  # 🔹 allow all branches
        'gl_accounts': Account.objects.filter(gl_no__startswith="104"), # 🔹 allow all accounts
        'current_date': timezone.now().strftime('%d/%m/%Y'),
        'reporting_date': None,
        'selected_branch': None,
        'selected_gl_no': None,
        'grand_total_loan_amount': 0,
        'grand_total_interest': 0,
        'grand_total_principal_paid': 0,
        'grand_total_interest_paid': 0,
        'grand_total_repayment': 0,
    }

    if request.method == 'POST':
        if form.is_valid():
            reporting_date = form.cleaned_data['reporting_date']
            branch = form.cleaned_data.get('branch') or user_branch
            gl_no = form.cleaned_data.get('gl_no')

            # 🔹 Base query - do NOT force only user_branch
            loans = Loans.objects.filter(
                disbursement_date__lte=reporting_date
            ).select_related('customer')

            # Apply branch filter if selected
            if branch:
                loans = loans.filter(branch=branch)

            # Apply GL filter if selected
            if gl_no:
                loans = loans.filter(gl_no=gl_no.gl_no)

            expected_repayments = []
            grand_totals = {
                'loan_amount': 0,
                'total_interest': 0,
                'total_principal_paid': 0,
                'total_interest_paid': 0,
                'expected_principal_repayment': 0,
                'expected_interest_repayment': 0,
            }

            for loan in loans:
                # Simplified repayment calc (adjust with real logic)
                expected_principal = loan.loan_amount / loan.num_install if loan.num_install else 0
                expected_interest = (loan.loan_amount * loan.interest_rate / 100) / 12  

                payments = LoanHist.objects.filter(
                    gl_no=loan.gl_no,
                    ac_no=loan.ac_no,
                    cycle=loan.cycle,
                    trx_type='LP',
                    trx_date__lte=reporting_date
                ).aggregate(
                    total_principal=Sum('principal'),
                    total_interest=Sum('interest')
                )

                principal_paid = payments['total_principal'] or 0
                interest_paid = payments['total_interest'] or 0

                principal_due = max(0, expected_principal - principal_paid)
                interest_due = max(0, expected_interest - interest_paid)

                repayment_data = {
                    'gl_no': loan.gl_no,
                    'ac_no': loan.ac_no,
                    'customer_name': loan.customer.get_full_name() if loan.customer else "No Customer",
                    'loan_amount': loan.loan_amount,
                    'total_interest': expected_interest,
                    'total_principal_paid': principal_paid,
                    'total_interest_paid': interest_paid,
                    'expected_principal_repayment': principal_due,
                    'expected_interest_repayment': interest_due,
                }

                expected_repayments.append(repayment_data)

                # Update grand totals
                grand_totals['loan_amount'] += loan.loan_amount
                grand_totals['total_interest'] += expected_interest
                grand_totals['total_principal_paid'] += principal_paid
                grand_totals['total_interest_paid'] += interest_paid
                grand_totals['expected_principal_repayment'] += principal_due
                grand_totals['expected_interest_repayment'] += interest_due

            context.update({
                'expected_repayments': expected_repayments,
                'reporting_date': reporting_date,
                'selected_branch': branch.branch_code if branch else None,
                'selected_gl_no': gl_no.gl_no if gl_no else None,
                'grand_totals': grand_totals,
                'grand_total_loan_amount': grand_totals['loan_amount'],
                'grand_total_interest': grand_totals['total_interest'],
                'grand_total_principal_paid': grand_totals['total_principal_paid'],
                'grand_total_interest_paid': grand_totals['total_interest_paid'],
                'grand_total_repayment': grand_totals['expected_principal_repayment'],
            })
        else:
            context['error_message'] = "Please correct the errors below."

    return render(request, 'reports/loans/loan_disbursement_report.html', context)



@login_required


def loan_repayment_report(request, loan_id=None):
    # Get current user's branch and company name
    user_branch = request.user.branch
    user_company_name = user_branch.company_name if user_branch else None

    # Initialize form with user's branch
    form = LoanRepaymentReportForm(request.POST or None, user_branch=user_branch)

    # Ensure loan is always defined
    loan = None

    # Initialize context
    context = {
        'form': form,
        'repayment_list': [],
        'disbursements': [],
        'loan': None,
        'selected_branch': None,
        'company_name': user_company_name,
        'branch_name': user_branch.branch_name if user_branch else None,
        'grand_total_principal': 0,
        'grand_total_interest': 0,
        'grand_total_penalty': 0,
        'total_paid_sum': 0,
        'subtotals': [],
        'current_date': timezone.now().strftime('%d/%m/%Y'),
        'start_date': None,
        'end_date': None,
        'grand_total_disbursed': 0,
        'outstanding_amount': 0,
    }

    # Get loan if ID is provided
    if loan_id:
        try:
            loan = Loans.objects.select_related(
                'customer', 'branch', 'loan_officer', 'business_sector'
            ).get(
                pk=loan_id,
                branch__company_name=user_company_name
            )
            context['loan'] = loan
        except Loans.DoesNotExist:
            messages.error(request, "Loan not found or not accessible")

    if request.method == 'POST' and form.is_valid():
        start_date = form.cleaned_data['start_date']
        end_date = form.cleaned_data['end_date']
        branch = form.cleaned_data.get('branch')
        gl_no = form.cleaned_data.get('gl_no')
        cycle = form.cleaned_data.get('cycle')

        repayments = LoanHist.objects.filter(
            trx_date__range=[start_date, end_date],
            trx_type='LP',
            branch__company_name=user_company_name
        ).select_related('branch')

        if loan:
            repayments = repayments.filter(
                gl_no=loan.gl_no,
                ac_no=loan.ac_no,
                cycle=loan.cycle
            )
            disbursements = LoanHist.objects.filter(
                trx_type='LD',
                gl_no=loan.gl_no,
                ac_no=loan.ac_no,
                cycle=loan.cycle,
                branch__company_name=user_company_name
            )
        elif branch:
            repayments = repayments.filter(branch=branch)
        if gl_no:
            repayments = repayments.filter(gl_no=gl_no.gl_no)
        if cycle:
            repayments = repayments.filter(cycle=cycle)

        loan_numbers = repayments.values_list('gl_no', 'ac_no', 'cycle').distinct()
        loans = Loans.objects.filter(
            gl_no__in=[x[0] for x in loan_numbers],
            ac_no__in=[x[1] for x in loan_numbers],
            cycle__in=[x[2] for x in loan_numbers]
        ).select_related('customer')

        loan_customer_map = {
            (loan.gl_no, loan.ac_no, loan.cycle): loan.customer.get_full_name() if loan.customer else "No Customer"
            for loan in loans
        }

        repayment_list = []
        for repayment in repayments:
            key = (repayment.gl_no, repayment.ac_no, repayment.cycle)
            repayment_list.append({
                'counter': len(repayment_list) + 1,
                'customer_name': loan_customer_map.get(key, "No Customer"),
                'branch_name': repayment.branch.branch_name,
                'gl_no': repayment.gl_no,
                'ac_no': repayment.ac_no,
                'cycle': repayment.cycle,
                'trx_date': repayment.trx_date,
                'trx_no': repayment.trx_no,
                'principal': float(repayment.principal or 0),
                'interest': float(repayment.interest or 0),
                'penalty': float(repayment.penalty or 0),
                'total_paid': float((repayment.principal or 0) + (repayment.interest or 0)),
            })

        repayment_aggregates = repayments.aggregate(
            total_principal=Sum('principal'),
            total_interest=Sum('interest'),
            total_penalty=Sum('penalty'),
            total_paid=Sum(F('principal') + F('interest'))
        )

        disbursed_total = disbursements.aggregate(
            total_disbursed=Sum('principal')
        ).get('total_disbursed', 0) or 0 if loan else 0

        subtotals = []
        for subtotal in repayments.values(
            'gl_no', 'cycle', 'branch__branch_name'
        ).annotate(
            subtotal_principal=Sum('principal'),
            subtotal_interest=Sum('interest'),
            subtotal_penalty=Sum('penalty'),
            subtotal_paid=Sum(F('principal') + F('interest'))
        ).order_by('branch__branch_name', 'gl_no', 'cycle'):
            subtotals.append({
                'branch_name': subtotal['branch__branch_name'],
                'gl_no': subtotal['gl_no'],
                'cycle': subtotal['cycle'],
                'subtotal_principal': float(subtotal['subtotal_principal'] or 0),
                'subtotal_interest': float(subtotal['subtotal_interest'] or 0),
                'subtotal_penalty': float(subtotal['subtotal_penalty'] or 0),
                'subtotal_paid': float(subtotal['subtotal_paid'] or 0),
            })

        outstanding_amount = max(0, float(loan.loan_amount) - float(repayment_aggregates.get('total_principal') or 0)) if loan else 0

        context.update({
            'repayment_list': repayment_list,
            'disbursements': list(disbursements) if loan else [],
            'loan': loan,
            'outstanding_amount': outstanding_amount,
            'grand_total_principal': float(repayment_aggregates.get('total_principal', 0) or 0),
            'grand_total_interest': float(repayment_aggregates.get('total_interest', 0) or 0),
            'grand_total_penalty': float(repayment_aggregates.get('total_penalty', 0) or 0),
            'total_paid_sum': float(repayment_aggregates.get('total_paid', 0) or 0),
            'grand_total_disbursed': float(disbursed_total),
            'selected_branch': branch,
            'start_date': start_date,
            'end_date': end_date,
            'subtotals': subtotals,
        })

    return render(request, 'reports/loans/loan_repayment_report.html', context)


# views.py
from django.shortcuts import render
from django.utils import timezone
from django.db.models import Sum, F
from .forms import LoanRepaymentReportForm
from loans.models import LoanHist, Loans  # Make sure Loans is imported
from company.models import Branch  # Import Branch if not already
from accounts.models import Account
from django.contrib.auth.decorators import login_required




from datetime import datetime, timedelta
from django.utils.timezone import now

@login_required
 # adjust imports as needed

def repayment_since_disbursement_report(request):
    user_branches = Branch.objects.filter(user=request.user)

    # Assuming Account has a ForeignKey to Branch as 'branch', if not, adjust accordingly
    user_gl_accounts = Account.objects.filter(branch__in=user_branches).distinct()

    context = {
        'report_title': 'Repayment Since Disbursement',
        'repayments': [],
        'grand_total_principal': 0,
        'grand_total_interest': 0,
        'grand_total_penalty': 0,
        'total_paid_sum': 0,
        'total_loan_amount': 0,
        'total_loan_interest': 0,
        'total_percentage_paid': 0,
        'current_date': timezone.now().strftime('%Y-%m-%d'),
        'start_date': '',
        'end_date': '',
        'branches': user_branches,
        'gl_accounts': user_gl_accounts,
        'selected_branch': '',
        'selected_gl_no': '',
        'company_name': 'Your Company',  # default, will update if branch selected
        'branch_name': 'All Branches',
    }

    if request.method == 'POST':
        start_date = request.POST.get('start_date', '')
        end_date = request.POST.get('end_date', '')
        selected_branch = request.POST.get('branch', '')
        selected_gl_no = request.POST.get('gl_no', '')

        context['start_date'] = start_date
        context['end_date'] = end_date
        context['selected_branch'] = selected_branch
        context['selected_gl_no'] = selected_gl_no

        if start_date and end_date:
            try:
                start_date_obj = datetime.strptime(start_date, '%Y-%m-%d').date()
                end_date_obj = datetime.strptime(end_date, '%Y-%m-%d').date()

                # Filter loans by user's branches only
                loans = Loans.objects.filter(
                    disbursement_date__lte=end_date_obj,
                    branch__in=user_branches,
                )

                if selected_branch:
                    loans = loans.filter(branch_id=selected_branch)
                    branch = user_branches.filter(id=selected_branch).first()
                    if branch:
                        context['company_name'] = branch.company_name
                        context['branch_name'] = branch.branch_name

                if selected_gl_no:
                    loans = loans.filter(gl_no=selected_gl_no)

                repayments = LoanHist.objects.filter(
                    trx_date__range=[start_date_obj, end_date_obj],
                    trx_type='LP',
                    gl_no__in=loans.values_list('gl_no', flat=True)
                )

                repayment_summaries = repayments.values('gl_no', 'ac_no', 'cycle').annotate(
                    total_principal=Sum('principal'),
                    total_interest=Sum('interest'),
                    total_penalty=Sum('penalty'),
                    total_paid=Sum(F('principal') + F('interest') + F('penalty'))
                )

                grand_total_principal = repayment_summaries.aggregate(Sum('total_principal'))['total_principal__sum'] or 0
                grand_total_interest = repayment_summaries.aggregate(Sum('total_interest'))['total_interest__sum'] or 0
                grand_total_penalty = repayment_summaries.aggregate(Sum('total_penalty'))['total_penalty__sum'] or 0

                repayments_with_percentage = []
                total_loan_amount = 0
                total_loan_interest = 0

                for summary in repayment_summaries:
                    loan = loans.filter(
                        gl_no=summary['gl_no'],
                        ac_no=summary['ac_no'],
                        cycle=summary['cycle']
                    ).first()

                    if loan:
                        loan_amount = loan.loan_amount
                        loan_interest = loan.total_interest
                        customer = loan.customer
                        customer_name = customer.get_full_name() if customer else 'Unknown'
                        branch_name = loan.branch.branch_name if loan.branch else 'N/A'
                    else:
                        loan_amount = 0
                        loan_interest = 0
                        customer_name = 'Unknown'
                        branch_name = 'N/A'

                    total_loan_amount += loan_amount
                    total_loan_interest += loan_interest
                    percentage_paid = (summary['total_principal'] / loan_amount) * 100 if loan_amount > 0 else 0

                    repayments_with_percentage.append({
                        'gl_no': summary['gl_no'],
                        'ac_no': summary['ac_no'],
                        'cycle': summary['cycle'],
                        'customer_name': customer_name,
                        'branch_name': branch_name,
                        'total_principal': summary['total_principal'],
                        'total_interest': summary['total_interest'],
                        'total_penalty': summary['total_penalty'],
                        'total_paid': summary['total_paid'],
                        'loan_amount': loan_amount,
                        'total_loan_interest': loan_interest,
                        'percentage_paid': round(percentage_paid, 2)
                    })

                total_paid_sum = sum(item['total_paid'] for item in repayments_with_percentage)
                total_percentage_paid = (grand_total_principal / total_loan_amount) * 100 if total_loan_amount > 0 else 0

                context.update({
                    'repayments': repayments_with_percentage,
                    'grand_total_principal': grand_total_principal,
                    'grand_total_interest': grand_total_interest,
                    'grand_total_penalty': grand_total_penalty,
                    'total_paid_sum': total_paid_sum,
                    'total_loan_amount': total_loan_amount,
                    'total_loan_interest': total_loan_interest,
                    'total_percentage_paid': total_percentage_paid,
                })

            except ValueError:
                messages.error(request, "Invalid date format. Please use YYYY-MM-DD format.")

    return render(request, 'reports/loans/repayment_since_disbursement_report.html', context)


from django.db.models import Sum
from django.utils import timezone
from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from decimal import Decimal
@login_required

def loan_outstanding_balance(request):
    user = request.user

    user_company_names = Branch.objects.filter(user=user).values_list('company_name', flat=True).distinct()
    branches = Branch.objects.filter(company_name__in=user_company_names)
    gl_accounts = Account.objects.filter(branch__company_name__in=user_company_names)

    outstanding_loans = []
    reporting_date = ''
    selected_branch = ''
    selected_gl_no = ''
    selected_branch_obj = None
    selected_gl_obj = None

    # Initialize grand totals dictionary
    grand_totals = {
        'loan_disbursement': 0,
        'outstanding_principal': 0,
        'outstanding_interest': 0,
        'outstanding_amount': 0,
    }

    if request.method == 'POST':
        reporting_date = request.POST.get('reporting_date', '')
        selected_branch = request.POST.get('branch', '')
        selected_gl_no = request.POST.get('gl_no', '')

        if reporting_date:
            loans = Loans.objects.filter(
                disbursement_date__lte=reporting_date,
                branch__company_name__in=user_company_names
            )

            if selected_branch:
                branch_qs = branches.filter(id=selected_branch)
                if branch_qs.exists():
                    selected_branch_obj = branch_qs.first()
                    loans = loans.filter(branch_id=selected_branch)

            if selected_gl_no:
                gl_qs = gl_accounts.filter(gl_no=selected_gl_no)
                if gl_qs.exists():
                    selected_gl_obj = gl_qs.first()
                    loans = loans.filter(gl_no=selected_gl_no)

            for loan in loans:
                latest_transaction = LoanHist.objects.filter(
                    gl_no=loan.gl_no,
                    ac_no=loan.ac_no,
                    cycle=loan.cycle,
                    trx_type='LD'
                ).order_by('-trx_date').first()

                expiry_date = latest_transaction.trx_date if latest_transaction else None

                total_principal_paid = LoanHist.objects.filter(
                    gl_no=loan.gl_no,
                    ac_no=loan.ac_no,
                    cycle=loan.cycle,
                    trx_type='LP'
                ).aggregate(total=Sum('principal'))['total'] or 0

                total_interest_paid = LoanHist.objects.filter(
                    gl_no=loan.gl_no,
                    ac_no=loan.ac_no,
                    cycle=loan.cycle,
                    trx_type='LP'
                ).aggregate(total=Sum('interest'))['total'] or 0

                customer_name = (
                    f"{loan.customer.first_name} {loan.customer.middle_name or ''} {loan.customer.last_name}".strip()
                    if loan.customer else 'N/A'
                )

                outstanding_principal = loan.loan_amount + total_principal_paid
                outstanding_interest = loan.total_interest - total_interest_paid
                outstanding_amount = outstanding_principal + outstanding_interest

                # Append loan record
                outstanding_loans.append({
                    'gl_no': loan.gl_no,
                    'ac_no': loan.ac_no,
                    'customer_name': customer_name,
                    'loan_amount': loan.loan_amount,
                    'total_interest': loan.total_interest,
                    'disbursement_date': loan.disbursement_date,
                    'total_principal_paid': total_principal_paid,
                    'total_interest_paid': total_interest_paid,
                    'outstanding_principal': outstanding_principal,
                    'outstanding_interest': outstanding_interest,
                    'outstanding_amount': outstanding_amount,
                    'expiry_date': expiry_date
                })

                # Update grand totals
                grand_totals['loan_disbursement'] += loan.loan_amount
                grand_totals['outstanding_principal'] += outstanding_principal
                grand_totals['outstanding_interest'] += outstanding_interest
                grand_totals['outstanding_amount'] += outstanding_amount

    context = {
        'report_title': 'Loan Outstanding Balance Report',
        'outstanding_loans': outstanding_loans,
        'grand_totals': grand_totals,
        'grand_total_loan_disbursement': grand_totals['loan_disbursement'],
        'grand_total_outstanding_principal': grand_totals['outstanding_principal'],
        'grand_total_outstanding_interest': grand_totals['outstanding_interest'],
        'grand_total_outstanding_amount': grand_totals['outstanding_amount'],
        'current_date': timezone.now(),
        'reporting_date': reporting_date,
        'branches': branches,
        'gl_accounts': gl_accounts,
        'selected_branch': selected_branch,
        'selected_gl_no': selected_gl_no,
        'selected_branch_obj': selected_branch_obj,
        'selected_gl_obj': selected_gl_obj,
    }

    return render(request, 'reports/loans/loan_outstanding_balance_report.html', context)



from django.core.exceptions import ValidationError

from django.shortcuts import render
from django.utils import timezone
from django.core.exceptions import ValidationError
from django.db.models import Sum
 # Update with the actual import path
from django.http import HttpResponseBadRequest
from django.shortcuts import render, redirect
from django.contrib import messages
from django.db.models import Sum
from django.utils import timezone

from django.shortcuts import render, redirect
from django.contrib import messages
from django.db.models import Sum
from django.utils import timezone



def expected_repayment(request):
    # Initialize variables
    expected_repayments = []
    grand_total_repayment = 0
    grand_total_interest = 0
    grand_total_principal_paid = 0
    grand_total_interest_paid = 0
    reporting_date = ''
    selected_branch = ''
    selected_gl_no = ''

    if request.user.is_authenticated:
        # Get branches linked to the logged-in user
        branches = Branch.objects.filter(user=request.user)
        # Get accounts linked to those branches
        gl_accounts = Account.objects.filter(branch__in=branches)
    else:
        branches = Branch.objects.none()
        gl_accounts = Account.objects.none()

    if request.method == 'POST':
        # Get form data
        reporting_date = request.POST.get('reporting_date', '')
        selected_branch = request.POST.get('branch', '')
        selected_gl_no = request.POST.get('gl_no', '')

        # Validate reporting_date
        if reporting_date:
            try:
                reporting_date = timezone.datetime.strptime(reporting_date, '%Y-%m-%d').date()
            except ValueError:
                messages.error(request, "Invalid date format. Please use YYYY-MM-DD.")
                return redirect('expected_repayment')
        else:
            reporting_date = timezone.now().date()

        # Filter loans disbursed on or before reporting_date
        loans = Loans.objects.filter(disbursement_date__lte=reporting_date)

        # Filter loans by selected branch if provided and the branch belongs to the user
        if selected_branch:
            if branches.filter(id=selected_branch).exists():
                loans = loans.filter(branch_id=selected_branch)
            else:
                loans = loans.none()  # User selected branch they don't own

        # Filter loans by selected GL number if provided and the GL account belongs to the user's branches
        if selected_gl_no:
            if gl_accounts.filter(gl_no=selected_gl_no).exists():
                loans = loans.filter(gl_no=selected_gl_no)
            else:
                loans = loans.none()  # User selected GL not in their branches

        # Prepare list for output
        for loan in loans:
            total_disbursements = LoanHist.objects.filter(
                gl_no=loan.gl_no,
                ac_no=loan.ac_no,
                cycle=loan.cycle,
                trx_type='LD',
                trx_date__lte=reporting_date
            ).aggregate(total_disbursements=Sum('principal'))['total_disbursements'] or 0

            total_principal_paid = LoanHist.objects.filter(
                gl_no=loan.gl_no,
                ac_no=loan.ac_no,
                cycle=loan.cycle,
                trx_type='LP',
                trx_date__lte=reporting_date
            ).aggregate(total_principal_paid=Sum('principal'))['total_principal_paid'] or 0

            total_interest_paid = LoanHist.objects.filter(
                gl_no=loan.gl_no,
                ac_no=loan.ac_no,
                cycle=loan.cycle,
                trx_type='LP',
                trx_date__lte=reporting_date
            ).aggregate(total_interest_paid=Sum('interest'))['total_interest_paid'] or 0

            total_interest = LoanHist.objects.filter(
                gl_no=loan.gl_no,
                ac_no=loan.ac_no,
                cycle=loan.cycle,
                trx_type='LD',
                trx_date__lte=reporting_date
            ).aggregate(total_interest=Sum('interest'))['total_interest'] or 0

            expected_principal_repayment = total_disbursements - total_principal_paid
            expected_interest_repayment = total_interest - total_interest_paid

            # Fetch customer name
            if loan.customer:
                customer_name = f"{loan.customer.first_name} {loan.customer.middle_name or ''} {loan.customer.last_name}".strip()
            else:
                customer_name = 'N/A'

            expected_repayments.append({
                'gl_no': loan.gl_no,
                'ac_no': loan.ac_no,
                'customer_name': customer_name,
                'loan_amount': loan.loan_amount,
                'total_disbursements': total_disbursements,
                'total_principal_paid': total_principal_paid,
                'total_interest_paid': total_interest_paid,
                'total_interest': total_interest,
                'expected_principal_repayment': expected_principal_repayment,
                'expected_interest_repayment': expected_interest_repayment,
                'expected_total_repayment': expected_principal_repayment + expected_interest_repayment,
            })

            # Update totals
            grand_total_repayment += expected_principal_repayment
            grand_total_interest += total_interest
            grand_total_principal_paid += total_principal_paid
            grand_total_interest_paid += total_interest_paid

    context = {
        'report_title': 'Expected Repayment Report',
        'expected_repayments': expected_repayments,
        'grand_total_repayment': grand_total_repayment,
        'grand_total_interest': grand_total_interest,
        'grand_total_principal_paid': grand_total_principal_paid,
        'grand_total_interest_paid': grand_total_interest_paid,
        'grand_total_expected': grand_total_repayment + grand_total_interest_paid,
        'current_date': timezone.now(),
        'reporting_date': reporting_date,
        'branches': branches,
        'gl_accounts': gl_accounts,
        'selected_branch': selected_branch,
        'selected_gl_no': selected_gl_no
    }

    return render(request, 'reports/loans/expected_repayment_report.html', context)




def active_loans_by_officer(request):
    # Initialize variables
    active_loans_by_officer = {}
    grand_total_loans = 0
    grand_total_amount = 0
    selected_officer = ''
    selected_branch = ''
    selected_gl_no = ''  # Changed from selected_product
    reporting_date = ''

    # Fetch all options for dropdowns
    officers = Account_Officer.objects.all()
    branches = Branch.objects.all()
    gl_accounts = Account.objects.all()  # Changed from products

    if request.method == 'POST':
        # Get form data
        selected_officer = request.POST.get('officer', '')
        selected_branch = request.POST.get('branch', '')
        selected_gl_no = request.POST.get('gl_no', '')  # Changed from product
        reporting_date = request.POST.get('reporting_date', '')

        # Initialize queryset
        loans = Loans.objects.filter(loan_amount__gt=0)  # Changed to gt 0 for active loans

        # Apply filters
        if selected_officer:
            loans = loans.filter(loan_officer__user=selected_officer)
        if selected_branch:
            loans = loans.filter(branch_id=selected_branch)
        if selected_gl_no:  # Changed from product filter
            loans = loans.filter(gl_no=selected_gl_no)  # Using the gl_no field directly
        if reporting_date:
            try:
                reporting_date = timezone.datetime.strptime(reporting_date, '%Y-%m-%d').date()
                loans = loans.filter(disbursement_date__lte=reporting_date)
            except ValueError:
                messages.error(request, "Invalid date format. Please use YYYY-MM-DD.")
                return redirect('active_loans_by_officer')

        # Group loans by loan officer and calculate totals
        for loan in loans:
            officer_name = loan.loan_officer.user if loan.loan_officer else 'Unassigned'
            
            if officer_name not in active_loans_by_officer:
                active_loans_by_officer[officer_name] = {
                    'loans': [],
                    'total_loans': 0,
                    'total_amount': 0,
                    'branch': loan.branch.branch_name if loan.branch else 'N/A'
                }

            customer_name = ''
            if loan.customer:
                customer_name = f"{loan.customer.first_name} {loan.customer.middle_name or ''} {loan.customer.last_name}".strip()

            active_loans_by_officer[officer_name]['loans'].append({
                'gl_no': loan.gl_no,
                'ac_no': loan.ac_no,
                'customer_name': customer_name or 'N/A',
                'loan_amount': loan.loan_amount,
                'disbursement_date': loan.disbursement_date,
                'account_name': loan.cust_gl_no  # Using available GL information
            })

            # Update totals
            active_loans_by_officer[officer_name]['total_loans'] += 1
            active_loans_by_officer[officer_name]['total_amount'] += loan.loan_amount
            grand_total_loans += 1
            grand_total_amount += loan.loan_amount

    context = {
        'report_title': 'Active Loans by Loan Officer',
        'active_loans_by_officer': dict(sorted(active_loans_by_officer.items())),
        'officers': officers,
        'branches': branches,
        'gl_accounts': gl_accounts,  # Changed from products
        'selected_officer': selected_officer,
        'selected_branch': selected_branch,
        'selected_gl_no': selected_gl_no,  # Changed from selected_product
        'reporting_date': reporting_date,
        'grand_total_loans': grand_total_loans,
        'grand_total_amount': grand_total_amount,
        'current_date': timezone.now(),
    }
    return render(request, 'reports/loans/active_loans_by_officer.html', context)

from .forms import LoanTillSheetForm





def loan_till_sheet(request):
    # Initialize variables
    repayments_with_percentage = []
    grand_total_principal = grand_total_interest = grand_total_penalty = total_paid_sum = 0
    start_date = end_date = None
    selected_branch = None
    selected_gl_no = None

    # Initialize form
    form = LoanTillSheetForm(request.POST or None)

    if request.method == 'POST' and form.is_valid():
        try:
            # Get cleaned form data
            start_date = form.cleaned_data['start_date']
            end_date = form.cleaned_data['end_date']
            selected_branch = form.cleaned_data.get('branch')
            selected_gl_no = form.cleaned_data.get('gl_no')

            # Fetch loans disbursed before or on the end date
            loans = Loans.objects.filter(disbursement_date__lte=end_date) if end_date else Loans.objects.all()
            
            if selected_gl_no:
                loans = loans.filter(gl_no=selected_gl_no.gl_no)

            # Filter loan repayments
            repayments = LoanHist.objects.filter(
                trx_date__range=[start_date, end_date],
                trx_type='LP',
                gl_no__in=loans.values_list('gl_no', flat=True)
            )

            # Apply branch filter if selected
            if selected_branch:
                repayments = repayments.filter(branch=selected_branch.branch_code)

            # Aggregate repayments
            repayment_summaries = repayments.values('gl_no', 'ac_no', 'cycle').annotate(
                total_principal=Sum('principal'),
                total_interest=Sum('interest'),
                total_penalty=Sum('penalty'),
                total_paid=Sum(F('principal') + F('interest') + F('penalty'))
            )

            # Calculate grand totals
            grand_total_principal = repayment_summaries.aggregate(Sum('total_principal'))['total_principal__sum'] or 0
            grand_total_interest = repayment_summaries.aggregate(Sum('total_interest'))['total_interest__sum'] or 0
            grand_total_penalty = repayment_summaries.aggregate(Sum('total_penalty'))['total_penalty__sum'] or 0

            # Process repayment data
            for summary in repayment_summaries:
                loan = loans.filter(
                    gl_no=summary['gl_no'],
                    ac_no=summary['ac_no'],
                    cycle=summary['cycle']
                ).first()
                
                if loan:
                    loan_amount = loan.loan_amount
                    total_loan_interest = loan.total_interest
                    customer = Customer.objects.filter(gl_no=loan.gl_no, ac_no=loan.ac_no).first()
                    customer_name = customer.get_full_name() if customer else 'Unknown'
                else:
                    loan_amount = total_loan_interest = 0
                    customer_name = 'Unknown'

                percentage_paid = (summary['total_principal'] / loan_amount * 100) if loan_amount > 0 else 0

                repayments_with_percentage.append({
                    'gl_no': summary['gl_no'],
                    'ac_no': summary['ac_no'],
                    'cycle': summary['cycle'],
                    'total_principal': summary['total_principal'],
                    'total_interest': summary['total_interest'],
                    'total_penalty': summary['total_penalty'],
                    'total_paid': summary['total_paid'],
                    'loan_amount': loan_amount,
                    'total_loan_interest': total_loan_interest,
                    'customer_name': customer_name,
                    'percentage_paid': round(percentage_paid, 2)
                })

            total_paid_sum = sum(item['total_paid'] for item in repayments_with_percentage)

        except Exception as e:
            messages.error(request, f"An error occurred: {str(e)}")
            return redirect('loan_till_sheet')

    context = {
        'form': form,
        'report_title': 'Loan Repayment Till Sheet',
        'repayments': repayments_with_percentage,
        'selected_branch': selected_branch,
        'selected_gl_no': selected_gl_no,
        'grand_total_principal': grand_total_principal,
        'grand_total_interest': grand_total_interest,
        'grand_total_penalty': grand_total_penalty,
        'total_paid_sum': total_paid_sum,
        'current_date': timezone.now(),
    }
    return render(request, 'reports/loans/loan_till_sheet.html', context)



def loan_due_vs_repayment_report(request):
    # Initialize variables
    expected_repayments = []
    grand_totals = {
        'loan_amount': 0,
        'total_interest': 0,
        'total_principal_paid': 0,
        'total_interest_paid': 0,
        'expected_principal_repayment': 0,
        'expected_interest_repayment': 0,
    }
    
    reporting_date = timezone.now().date()
    selected_branch = ''
    selected_gl_no = ''

    if request.method == 'POST':
        # Get form data
        reporting_date_str = request.POST.get('reporting_date', '')
        selected_branch = request.POST.get('branch', '')
        selected_gl_no = request.POST.get('gl_no', '')

        # Validate reporting_date
        if reporting_date_str:
            try:
                reporting_date = timezone.datetime.strptime(reporting_date_str, '%Y-%m-%d').date()
            except ValueError:
                return HttpResponseBadRequest("Invalid date format. Please use YYYY-MM-DD.")

        # Initialize queryset with filters
        loans = Loans.objects.filter(disbursement_date__lte=reporting_date)
        
        # Apply branch filter if provided
        if selected_branch:
            loans = loans.filter(
                Q(branch_id=selected_branch) | 
                Q(branch__branch_code=selected_branch)
            )
        
        # Apply GL filter if provided
        if selected_gl_no:
            loans = loans.filter(gl_no=selected_gl_no)

        # Prepare report data
        for loan in loans:
            # Get all disbursements
            disbursements = LoanHist.objects.filter(
                gl_no=loan.gl_no,
                ac_no=loan.ac_no,
                cycle=loan.cycle,
                trx_type='LD',
                trx_date__lte=reporting_date
            ).aggregate(
                total_principal=Sum('principal'),
                total_interest=Sum('interest')
            )
            
            # Get all repayments
            repayments = LoanHist.objects.filter(
                gl_no=loan.gl_no,
                ac_no=loan.ac_no,
                cycle=loan.cycle,
                trx_type='LP',
                trx_date__lte=reporting_date
            ).aggregate(
                total_principal_paid=Sum('principal'),
                total_interest_paid=Sum('interest')
            )

            # Calculate values (handle None cases)
            total_disbursements = disbursements['total_principal'] or 0
            total_interest = disbursements['total_interest'] or 0
            total_principal_paid = repayments['total_principal_paid'] or 0
            total_interest_paid = repayments['total_interest_paid'] or 0
            
            expected_principal_repayment = total_disbursements + total_principal_paid
            expected_interest_repayment = total_interest + total_interest_paid

            # Update grand totals
            grand_totals['loan_amount'] += loan.loan_amount
            grand_totals['total_interest'] += total_interest
            grand_totals['total_principal_paid'] += total_principal_paid
            grand_totals['total_interest_paid'] += total_interest_paid
            grand_totals['expected_principal_repayment'] += expected_principal_repayment
            grand_totals['expected_interest_repayment'] += expected_interest_repayment

            expected_repayments.append({
                'gl_no': loan.gl_no,
                'ac_no': loan.ac_no,
                'customer_name': f"{loan.customer.first_name} {loan.customer.middle_name or ''} {loan.customer.last_name}" if loan.customer else 'N/A',
                'loan_amount': loan.loan_amount,
                'total_disbursements': total_disbursements,
                'total_interest': total_interest,
                'total_principal_paid': total_principal_paid,
                'total_interest_paid': total_interest_paid,
                'expected_principal_repayment': expected_principal_repayment,
                'expected_interest_repayment': expected_interest_repayment,
            })

    context = {
        'report_title': 'Loan Dues vs Repayment Report',
        'expected_repayments': expected_repayments if request.method == 'POST' else None,
        'grand_totals': grand_totals,
        'current_date': timezone.now(),
        'reporting_date': reporting_date,
        'branches': Branch.objects.all(),
        'gl_accounts': Account.objects.all(),
        'selected_branch': selected_branch,
        'selected_gl_no': selected_gl_no
    }

    return render(request, 'reports/loans/loan_due_vs_repayment_report.html', context)


from django.shortcuts import render
from datetime import date
from accounts_admin.models import LoanProvision
from django.db.models import Sum
from django.utils.dateparse import parse_date  # Helper to parse date from string

def portfolio_at_risk_report_view(request):
    # Fetch the user's branch to get the associated Company
    user_branch = request.user.branch

    # Get the reporting_date from the user input, default to None if not supplied
    reporting_date_str = request.GET.get('reporting_date')  # Get date as string from query parameters
    reporting_date = None

    if reporting_date_str:
        try:
            # Try to parse the user-supplied date
            reporting_date = parse_date(reporting_date_str)
            if not reporting_date:
                reporting_date = date.today()  # Fallback to today if parsing fails
        except ValueError:
            reporting_date = date.today()  # Handle invalid date input gracefully

    # Initialize the report structure
    report = {
        "total_loans": 0,
        "total_outstanding": 0,
        "loans": [],
        "gl_nos": Account.objects.all(),
        "account_officers": Account_Officer.objects.all(),
    }

    # Fetch loans and apply filters only if the form is submitted (i.e., reporting_date is provided)
    if reporting_date:
        selected_product = request.GET.get('product')
        selected_officer = request.GET.get('account_officer')
        exclude_ac_no_one = request.GET.get('exclude_ac_no_one') == 'on'

        # Fetch loans disbursed before or on the reporting date
        loans = Loans.objects.filter(disb_status='T', disbursement_date__lte=reporting_date)  # Filter by reporting date

        if selected_product:
            loans = loans.filter(gl_no=selected_product)
        if selected_officer:
            loans = loans.filter(loan_officer_id=selected_officer)
        if exclude_ac_no_one:
            # Exclude internal accounts based on specific criteria
            non_financial_gl_nos = Account.objects.filter(is_non_financial=True).values_list('gl_no', flat=True)
            loans = loans.exclude(gl_no__in=non_financial_gl_nos)

        for loan in loans:
            # Loan expected repayment (LD) and actual payment (LP) till reporting date
            loan_hist_ld = LoanHist.objects.filter(
                gl_no=loan.gl_no, ac_no=loan.ac_no, cycle=loan.cycle, trx_type='LD', trx_date__lte=reporting_date
            ).aggregate(total_expected=Sum('principal'))['total_expected'] or 0

            loan_hist_lp = LoanHist.objects.filter(
                gl_no=loan.gl_no, ac_no=loan.ac_no, cycle=loan.cycle, trx_type='LP', trx_date__lte=reporting_date
            ).aggregate(total_paid=Sum('principal'))['total_paid'] or 0

            outstanding_balance = loan_hist_ld - loan_hist_lp  # Total expected minus total paid

            if outstanding_balance > 0:
                # Find the most recent LP transaction date for days in arrears calculation
                last_payment = LoanHist.objects.filter(
                    gl_no=loan.gl_no, ac_no=loan.ac_no, cycle=loan.cycle, trx_type='LP', trx_date__lte=reporting_date
                ).order_by('-trx_date').first()

                if last_payment:
                    last_payment_date = last_payment.trx_date

                    # Calculate days in arrears using the reporting date
                    days_in_arrears = (reporting_date - last_payment_date).days

                    # Find the appropriate provision category based on days in arrears
                    provision_category = LoanProvision.objects.filter(
                        min_days__lte=days_in_arrears, max_days__gte=days_in_arrears
                    ).first()

                    if provision_category:
                        category_name = provision_category.name
                        provision_amount = (provision_category.rate / 100) * outstanding_balance  # Calculate provision amount

                        # Add loan details and category to the report
                        report["loans"].append({
                            "gl_no": loan.gl_no,
                            "ac_no": loan.ac_no,
                            "cycle": loan.cycle,
                            "branch": loan.branch,
                            "outstanding_balance": outstanding_balance,
                            "days_in_arrears": days_in_arrears,
                            "category": category_name,
                            "provision_amount": provision_amount
                        })

                        # Update totals
                        report["total_loans"] += 1
                        report["total_outstanding"] += outstanding_balance

    # Render the report in the template
    return render(request, 'reports/loan_provision/par_report.html', {
        'report': report,
        'reporting_date': reporting_date  # Pass reporting date to template for display
    })



# Add these imports to your views.py
import io
from django.http import HttpResponse
from django.template.loader import get_template
from xhtml2pdf import pisa
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import csv
from datetime import datetime

class StatementView(View):
    def get(self, request):
        # Handle export requests
        export_format = request.GET.get('export')
        if export_format:
            return self.handle_export(request, export_format)
        
        # Regular statement view logic here
        context = {
            'start_date': request.GET.get('start_date', ''),
            'end_date': request.GET.get('end_date', ''),
            'gl_no': request.GET.get('gl_no', ''),
            'ac_no': request.GET.get('ac_no', ''),
        }
        return render(request, 'statement_of_account.html', context)
    
    def post(self, request):
        # Handle statement generation
        start_date = request.POST.get('start_date')
        end_date = request.POST.get('end_date')
        gl_no = request.POST.get('gl_no')
        ac_no = request.POST.get('ac_no')
        
        # Your existing statement generation logic here
        # ...
        
        context = {
            'statement_data': statement_data,
            'start_date': start_date,
            'end_date': end_date,
            'gl_no': gl_no,
            'ac_no': ac_no,
            'opening_balance': opening_balance,
            'closing_balance': closing_balance,
            'debit_amount': debit_amount,
            'credit_amount': credit_amount,
            # ... other context data
        }
        
        return render(request, 'statement_of_account.html', context)
    
    def handle_export(self, request, format_type):
        """Handle different export formats"""
        # Get statement data
        statement_data = self.get_statement_data(request)
        
        if format_type == 'pdf':
            return self.export_pdf(request, statement_data)
        elif format_type == 'word':
            return self.export_word(request, statement_data)
        elif format_type == 'excel':
            return self.export_excel(request, statement_data)
        elif format_type == 'csv':
            return self.export_csv(request, statement_data)
        else:
            return HttpResponse('Invalid export format', status=400)
    
    def get_statement_data(self, request):
        """Get statement data for export"""
        start_date = request.GET.get('start_date')
        end_date = request.GET.get('end_date')
        gl_no = request.GET.get('gl_no')
        ac_no = request.GET.get('ac_no')
        
        # Your logic to fetch statement data
        # This should return the same data structure as your main view
        
        return {
            'statement_data': statement_data,
            'start_date': start_date,
            'end_date': end_date,
            'gl_no': gl_no,
            'ac_no': ac_no,
            'opening_balance': opening_balance,
            'closing_balance': closing_balance,
            'debit_amount': debit_amount,
            'credit_amount': credit_amount,
            'full_name': full_name,
            'company': company,
            'branch': branch,
        }
    
    def export_pdf(self, request, data):
        """Export statement as PDF"""
        template_path = 'exports/statement_pdf.html'
        context = data
        
        # Create a Django response object, and specify content_type as pdf
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="statement_{data["ac_no"]}_{datetime.now().strftime("%Y%m%d")}.pdf"'
        
        # Find the template and render it
        template = get_template(template_path)
        html = template.render(context)
        
        # Create PDF
        pisa_status = pisa.CreatePDF(html, dest=response)
        
        # If error then show some funny view
        if pisa_status.err:
            return HttpResponse('We had some errors <pre>' + html + '</pre>')
        return response
    
    def export_word(self, request, data):
        """Export statement as Word document"""
        # Create a document
        doc = Document()
        
        # Add title
        title = doc.add_heading('STATEMENT OF ACCOUNT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subtitle
        subtitle = doc.add_heading(f'{data["company"].company_name.upper()} | {data["branch"].branch_name.upper()}', level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add client info
        doc.add_paragraph(f'Client: {data["full_name"]}')
        doc.add_paragraph(f'GL No: {data["gl_no"]} | AC No: {data["ac_no"]}')
        doc.add_paragraph(f'Period: {data["start_date"]} to {data["end_date"]}')
        doc.add_paragraph('')
        
        # Add table
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        # Add headers
        headers = ['Date', 'Description', 'Trx No', 'Debit', 'Credit', 'Balance']
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Add opening balance
        row_cells = table.add_row().cells
        row_cells[0].text = data['start_date']
        row_cells[1].text = 'Opening Balance'
        row_cells[2].text = '-'
        row_cells[3].text = '-'
        row_cells[4].text = '-'
        row_cells[5].text = f"₦{data['opening_balance']:,.2f}"
        
        # Add data rows
        for entry in data['statement_data']:
            row_cells = table.add_row().cells
            row_cells[0].text = entry['date'].strftime('%b %d, %Y')
            row_cells[1].text = entry['description']
            row_cells[2].text = entry['trx_no']
            row_cells[3].text = f"₦{entry['debit']:,.2f}" if entry['debit'] else '-'
            row_cells[4].text = f"₦{entry['credit']:,.2f}" if entry['credit'] else '-'
            row_cells[5].text = f"₦{entry['running_balance']:,.2f}"
        
        # Add totals row
        row_cells = table.add_row().cells
        row_cells[0].text = ''
        row_cells[1].text = 'TOTALS'
        row_cells[2].text = ''
        row_cells[3].text = f"₦{data['debit_amount']:,.2f}"
        row_cells[4].text = f"₦{data['credit_amount']:,.2f}"
        row_cells[5].text = f"₦{data['closing_balance']:,.2f}"
        
        # Make totals row bold
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Save to response
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="statement_{data["ac_no"]}_{datetime.now().strftime("%Y%m%d")}.docx"'
        
        doc.save(response)
        return response
    
    def export_excel(self, request, data):
        """Export statement as Excel file"""
        # Create a Pandas DataFrame
        statement_rows = []
        
        # Add opening balance
        statement_rows.append({
            'Date': data['start_date'],
            'Description': 'Opening Balance',
            'Trx No': '',
            'Debit': '',
            'Credit': '',
            'Balance': data['opening_balance']
        })
        
        # Add statement data
        for entry in data['statement_data']:
            statement_rows.append({
                'Date': entry['date'].strftime('%Y-%m-%d'),
                'Description': entry['description'],
                'Trx No': entry['trx_no'],
                'Debit': entry['debit'] if entry['debit'] else '',
                'Credit': entry['credit'] if entry['credit'] else '',
                'Balance': entry['running_balance']
            })
        
        # Add totals row
        statement_rows.append({
            'Date': '',
            'Description': 'TOTALS',
            'Trx No': '',
            'Debit': data['debit_amount'],
            'Credit': data['credit_amount'],
            'Balance': data['closing_balance']
        })
        
        df = pd.DataFrame(statement_rows)
        
        # Create Excel file in memory
        output = io.BytesIO()
        
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            # Write main statement
            df.to_excel(writer, sheet_name='Statement', index=False)
            
            # Write summary sheet
            summary_data = {
                'Item': ['Opening Balance', 'Total Debits', 'Total Credits', 'Closing Balance'],
                'Amount': [data['opening_balance'], data['debit_amount'], 
                          data['credit_amount'], data['closing_balance']]
            }
            summary_df = pd.DataFrame(summary_data)
            summary_df.to_excel(writer, sheet_name='Summary', index=False)
            
            # Get workbook and worksheet objects
            workbook = writer.book
            worksheet = writer.sheets['Statement']
            
            # Format headers
            header_format = workbook.add_format({
                'bold': True,
                'bg_color': '#366092',
                'font_color': 'white',
                'align': 'center'
            })
            
            # Apply header format
            for col_num, value in enumerate(df.columns.values):
                worksheet.write(0, col_num, value, header_format)
        
        output.seek(0)
        
        response = HttpResponse(
            output.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = f'attachment; filename="statement_{data["ac_no"]}_{datetime.now().strftime("%Y%m%d")}.xlsx"'
        
        return response
    
    def export_csv(self, request, data):
        """Export statement as CSV file"""
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename="statement_{data["ac_no"]}_{datetime.now().strftime("%Y%m%d")}.csv"'
        
        writer = csv.writer(response)
        
        # Write header info
        writer.writerow(['STATEMENT OF ACCOUNT'])
        writer.writerow([f'{data["company"].company_name.upper()} | {data["branch"].branch_name.upper()}'])
        writer.writerow([f'Client: {data["full_name"]}'])
        writer.writerow([f'GL No: {data["gl_no"]} | AC No: {data["ac_no"]}'])
        writer.writerow([f'Period: {data["start_date"]} to {data["end_date"]}'])
        writer.writerow([])  # Empty row
        
        # Write table headers
        writer.writerow(['Date', 'Description', 'Trx No', 'Debit', 'Credit', 'Balance'])
        
        # Write opening balance
        writer.writerow([
            data['start_date'],
            'Opening Balance',
            '',
            '',
            '',
            f"{data['opening_balance']:,.2f}"
        ])
        
        # Write statement data
        for entry in data['statement_data']:
            writer.writerow([
                entry['date'].strftime('%Y-%m-%d'),
                entry['description'],
                entry['trx_no'],
                f"{entry['debit']:,.2f}" if entry['debit'] else '',
                f"{entry['credit']:,.2f}" if entry['credit'] else '',
                f"{entry['running_balance']:,.2f}"
            ])
        
        # Write totals
        writer.writerow([
            '',
            'TOTALS',
            '',
            f"{data['debit_amount']:,.2f}",
            f"{data['credit_amount']:,.2f}",
            f"{data['closing_balance']:,.2f}"
        ])
        
        return response


















from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages

@login_required
def cbn_returns(request):
    """Display CBN Returns dashboard with all available forms"""
    
    cbn_forms = [
        {
            'code': 'MFBR_771',
            'title': 'Summary of Non-Performing Loans (NPL) / Classified Loans & Advances',
            'description': 'Central Bank of Nigeria',
            'category': 'Loan Analysis'
        },
        {
            'code': 'MMFBR_762',
            'title': 'Sectoral Analysis of Loans & Advances',
            'description': 'Central Bank of Nigeria',
            'category': 'Loan Analysis'
        },
        {
            'code': 'MMFBR_763',
            'title': 'Loan Structure & Maturity Profile',
            'description': 'Central Bank of Nigeria',
            'category': 'Loan Analysis'
        },
        {
            'code': 'MMFBR_764',
            'title': 'Interest Rate Schedule (Loan interest rates by category)',
            'description': 'Central Bank of Nigeria',
            'category': 'Loan Analysis'
        },
        {
            'code': 'MMFBR_811',
            'title': 'Other Assets',
            'description': 'Central Bank of Nigeria',
            'category': 'Assets'
        },
        {
            'code': 'MMFBR_141',
            'title': 'Other Deposits',
            'description': 'Central Bank of Nigeria',
            'category': 'Deposits'
        },
        {
            'code': 'MMFBR_201',
            'title': 'Deposit Structure & Maturity Profile',
            'description': 'Central Bank of Nigeria',
            'category': 'Deposits'
        },
        {
            'code': 'MMFBR_202',
            'title': 'Insured Deposits',
            'description': 'Central Bank of Nigeria',
            'category': 'Deposits'
        },
        {
            'code': 'MMFBR_212',
            'title': 'Takings from Banks in Nigeria',
            'description': 'Central Bank of Nigeria',
            'category': 'Takings'
        },
        {
            'code': 'MMFBR_322',
            'title': 'Takings from Other Institutions',
            'description': 'Central Bank of Nigeria',
            'category': 'Takings'
        },
        {
            'code': 'MMFBR_451',
            'title': 'Re-Financing Facilities',
            'description': 'Central Bank of Nigeria',
            'category': 'Facilities'
        },
        {
            'code': 'MMFBR_501',
            'title': 'Other Liabilities',
            'description': 'Central Bank of Nigeria',
            'category': 'Liabilities'
        },
        {
            'code': 'MMFBR_642',
            'title': 'Borrowings from Foreign Agencies',
            'description': 'Central Bank of Nigeria',
            'category': 'Borrowings'
        },
        {
            'code': 'MMFBR_651',
            'title': 'Borrowings from Other Agencies',
            'description': 'Central Bank of Nigeria',
            'category': 'Borrowings'
        },
        {
            'code': 'MMFBR_933',
            'title': 'Deferred Grants and Donations',
            'description': 'Central Bank of Nigeria',
            'category': 'Capital'
        },
        {
            'code': 'MMFBR_951',
            'title': 'Other Reserves',
            'description': 'Central Bank of Nigeria',
            'category': 'Capital'
        },
        {
            'code': 'MMFBR_996',
            'title': 'Off-Balance Sheet Engagements',
            'description': 'Central Bank of Nigeria',
            'category': 'Off-Balance Sheet'
        },
        {
            'code': 'MMFBR_980',
            'title': 'Gap Analysis (Assets vs Liabilities, maturity mismatch)',
            'description': 'Central Bank of Nigeria',
            'category': 'Analysis'
        },
        {
            'code': 'MMFBR_I',
            'title': 'Anti-Money Laundering and KYC (Know-Your-Customer) Schedule',
            'description': 'Central Bank of Nigeria',
            'category': 'Compliance'
        }
    ]
    
    context = {
        'cbn_forms': cbn_forms,
        'total_forms': len(cbn_forms)
    }
    
    return render(request, 'cbn_returns/cbn_returns.html', context)

@login_required
def cbn_return_detail(request, form_code):
    """Handle specific CBN return form"""
    
    # Dictionary mapping form codes to their details
    form_details = {
        'MFBR_771': {
            'title': 'Summary of Non-Performing Loans (NPL) / Classified Loans & Advances',
            'description': 'Generate and submit NPL classification reports',
            'fields': ['Loan Category', 'Outstanding Amount', 'NPL Amount', 'Provision Amount']
        },
        'MMFBR_762': {
            'title': 'Sectoral Analysis of Loans & Advances',
            'description': 'Analyze loans by economic sectors',
            'fields': ['Sector', 'Amount Disbursed', 'Outstanding Balance', 'Interest Rate']
        },
        # Add more form details as needed
    }
    
    form_info = form_details.get(form_code, {
        'title': 'CBN Return Form',
        'description': 'Central Bank of Nigeria regulatory return',
        'fields': []
    })
    
    context = {
        'form_code': form_code,
        'form_info': form_info
    }
    
    return render(request, 'cbn_returns/cbn_return_detail.html', context)



def handle_statement_export(request, export_format):
    """Handle different export formats for statement"""
    try:
        # Get parameters from GET request
        start_date = request.GET.get('start_date')
        end_date = request.GET.get('end_date') 
        gl_no = request.GET.get('gl_no')
        ac_no = request.GET.get('ac_no')
        
        if not all([start_date, end_date, gl_no, ac_no]):
            return HttpResponse('Missing required parameters', status=400)
            
        # Convert string dates to date objects
        from datetime import datetime
        start_date = datetime.strptime(start_date, '%Y-%m-%d').date()
        end_date = datetime.strptime(end_date, '%Y-%m-%d').date()
        
        # Get statement data
        statement_data = get_statement_data_for_export(request, start_date, end_date, gl_no, ac_no)
        
        if export_format == 'pdf':
            return export_statement_pdf(statement_data)
        elif export_format == 'word':
            return export_statement_word(statement_data)
        elif export_format == 'excel':
            return export_statement_excel(statement_data)
        elif export_format == 'csv':
            return export_statement_csv(statement_data)
        else:
            return HttpResponse('Invalid export format', status=400)
            
    except Exception as e:
        return HttpResponse(f'Export error: {str(e)}', status=500)

def get_statement_data_for_export(request, start_date, end_date, gl_no, ac_no):
    """Get statement data for export"""
    from transactions.models import Memtrans
    from customers.models import Customer
    from django.db.models import Sum
    
    # Get branch info
    branch = request.user.branch
    
    # Retrieve transactions within the specified date range
    transactions = Memtrans.objects.filter(
        ses_date__range=[start_date, end_date],
        gl_no=gl_no,
        ac_no=ac_no
    ).exclude(error='H').order_by('ses_date', 'trx_no')
    
    # Opening balance
    opening_balance = Memtrans.objects.filter(
        ses_date__lt=start_date,
        gl_no=gl_no,
        ac_no=ac_no
    ).exclude(error='H').aggregate(
        opening_balance=Sum('amount')
    )['opening_balance'] or 0
    
    # Closing balance  
    closing_balance = Memtrans.objects.filter(
        ses_date__lte=end_date,
        gl_no=gl_no,
        ac_no=ac_no
    ).exclude(error='H').aggregate(
        closing_balance=Sum('amount')
    )['closing_balance'] or 0
    
    # Debit & Credit totals
    debit_amount = abs(transactions.filter(type='D').aggregate(
        debit_amount=Sum('amount')
    )['debit_amount'] or 0)
    
    credit_amount = transactions.filter(type='C').aggregate(
        credit_amount=Sum('amount')
    )['credit_amount'] or 0
    
    # Get customer name
    try:
        customer = Customer.objects.filter(gl_no=gl_no, ac_no=ac_no).first()
        full_name = customer.get_full_name() if customer else f'Account {ac_no}'
    except:
        full_name = f'Account {ac_no}'
    
    # Statement details
    statement_data = []
    running_balance = opening_balance
    
    for transaction in transactions:
        if transaction.type == 'D':
            debit = abs(transaction.amount)
            credit = 0
            running_balance += transaction.amount
        else:
            debit = 0  
            credit = transaction.amount
            running_balance += transaction.amount
            
        entry = {
            'date': transaction.ses_date,
            'trx_no': transaction.trx_no,
            'description': transaction.description or '',
            'debit': debit,
            'credit': credit, 
            'running_balance': running_balance,
        }
        statement_data.append(entry)
    
    return {
        'start_date': start_date,
        'end_date': end_date,
        'gl_no': gl_no,
        'ac_no': ac_no,
        'opening_balance': opening_balance,
        'closing_balance': closing_balance,
        'debit_amount': debit_amount,
        'credit_amount': credit_amount,
        'statement_data': statement_data,
        'full_name': full_name,
        'branch': branch,
        'company': branch.company,
    }

def export_statement_pdf(data):
    """Export statement as PDF using WeasyPrint"""
    # Create simple HTML content
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>Statement of Account</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 20px; }}
            .header {{ text-align: center; margin-bottom: 30px; }}
            .company-info {{ text-align: center; margin-bottom: 20px; }}
            .account-info {{ margin-bottom: 20px; }}
            table {{ width: 100%; border-collapse: collapse; margin-bottom: 20px; }}
            th, td {{ border: 1px solid #ccc; padding: 8px; text-align: left; }}
            th {{ background-color: #f5f5f5; font-weight: bold; }}
            .amount {{ text-align: right; }}
            .total-row {{ font-weight: bold; background-color: #f9f9f9; }}
        </style>
    </head>
    <body>
        <div class="header">
            <h1>STATEMENT OF ACCOUNT</h1>
            <div class="company-info">
                <h2>{data['company'].company_name.upper()}</h2>
                <p>{data['branch'].branch_name}</p>
            </div>
        </div>
        
        <div class="account-info">
            <p><strong>Account Holder:</strong> {data['full_name']}</p>
            <p><strong>Account Number:</strong> {data['ac_no']}</p>
            <p><strong>GL Number:</strong> {data['gl_no']}</p>
            <p><strong>Period:</strong> {data['start_date'].strftime('%B %d, %Y')} to {data['end_date'].strftime('%B %d, %Y')}</p>
        </div>
        
        <table>
            <thead>
                <tr>
                    <th>Date</th>
                    <th>Description</th>
                    <th>Trx No</th>
                    <th class="amount">Debit</th>
                    <th class="amount">Credit</th>
                    <th class="amount">Balance</th>
                </tr>
            </thead>
            <tbody>
                <tr>
                    <td>{data['start_date'].strftime('%Y-%m-%d')}</td>
                    <td>Opening Balance</td>
                    <td>-</td>
                    <td class="amount">-</td>
                    <td class="amount">-</td>
                    <td class="amount">₦{data['opening_balance']:,.2f}</td>
                </tr>
    """
    
    # Add transaction rows
    for entry in data['statement_data']:
        html_content += f"""
                <tr>
                    <td>{entry['date'].strftime('%Y-%m-%d')}</td>
                    <td>{entry['description']}</td>
                    <td>{entry['trx_no']}</td>
                    <td class="amount">{f"₦{entry['debit']:,.2f}" if entry['debit'] else '-'}</td>
                    <td class="amount">{f"₦{entry['credit']:,.2f}" if entry['credit'] else '-'}</td>
                    <td class="amount">₦{entry['running_balance']:,.2f}</td>
                </tr>
        """
    
    # Add totals and close HTML
    html_content += f"""
                <tr class="total-row">
                    <td colspan="3">TOTALS</td>
                    <td class="amount">₦{data['debit_amount']:,.2f}</td>
                    <td class="amount">₦{data['credit_amount']:,.2f}</td>
                    <td class="amount">₦{data['closing_balance']:,.2f}</td>
                </tr>
            </tbody>
        </table>
        
        <div style="margin-top: 30px; text-align: center; font-size: 12px; color: #666;">
            <p>Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}</p>
        </div>
    </body>
    </html>
    """
    
    # Generate PDF
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="statement_{data["ac_no"]}_{datetime.now().strftime("%Y%m%d")}.pdf"'
    
    try:
        # Use WeasyPrint to generate PDF
        html_doc = weasyprint.HTML(string=html_content)
        html_doc.write_pdf(response)
    except Exception as e:
        # Fallback - return HTML if PDF generation fails
        response = HttpResponse(html_content, content_type='text/html')
        
    return response

def export_statement_word(data):
    """Export statement as Word document"""
    # Create a document
    doc = Document()
    
    # Add title
    title = doc.add_heading('STATEMENT OF ACCOUNT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add company info
    company_para = doc.add_paragraph()
    company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_run = company_para.add_run(f'{data["company"].company_name.upper()}\n{data["branch"].branch_name}')
    company_run.bold = True
    
    doc.add_paragraph('')  # Space
    
    # Add account info
    doc.add_paragraph(f'Account Holder: {data["full_name"]}')
    doc.add_paragraph(f'Account Number: {data["ac_no"]}')
    doc.add_paragraph(f'GL Number: {data["gl_no"]}')
    doc.add_paragraph(f'Period: {data["start_date"].strftime("%B %d, %Y")} to {data["end_date"].strftime("%B %d, %Y")}')
    
    doc.add_paragraph('')  # Space
    
    # Add table
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    # Add headers
    headers = ['Date', 'Description', 'Trx No', 'Debit', 'Credit', 'Balance']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Add opening balance
    row_cells = table.add_row().cells
    row_cells[0].text = data['start_date'].strftime('%Y-%m-%d')
    row_cells[1].text = 'Opening Balance'
    row_cells[2].text = '-'
    row_cells[3].text = '-'
    row_cells[4].text = '-'
    row_cells[5].text = f"₦{data['opening_balance']:,.2f}"
    
    # Add data rows
    for entry in data['statement_data']:
        row_cells = table.add_row().cells
        row_cells[0].text = entry['date'].strftime('%Y-%m-%d')
        row_cells[1].text = entry['description']
        row_cells[2].text = entry['trx_no']
        row_cells[3].text = f"₦{entry['debit']:,.2f}" if entry['debit'] else '-'
        row_cells[4].text = f"₦{entry['credit']:,.2f}" if entry['credit'] else '-'
        row_cells[5].text = f"₦{entry['running_balance']:,.2f}"
    
    # Add totals row
    row_cells = table.add_row().cells
    row_cells[0].text = ''
    row_cells[1].text = 'TOTALS'
    row_cells[2].text = ''
    row_cells[3].text = f"₦{data['debit_amount']:,.2f}"
    row_cells[4].text = f"₦{data['credit_amount']:,.2f}"
    row_cells[5].text = f"₦{data['closing_balance']:,.2f}"
    
    # Make totals row bold
    for cell in row_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Add generated timestamp
    doc.add_paragraph('')
    timestamp_para = doc.add_paragraph(f'Generated on {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
    timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save to response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="statement_{data["ac_no"]}_{datetime.now().strftime("%Y%m%d")}.docx"'
    
    doc.save(response)
    return response

def export_statement_excel(data):
    """Export statement as Excel file"""
    # Create a Pandas DataFrame
    statement_rows = []
    
    # Add opening balance
    statement_rows.append({
        'Date': data['start_date'].strftime('%Y-%m-%d'),
        'Description': 'Opening Balance',
        'Trx No': '',
        'Debit': '',
        'Credit': '',
        'Balance': data['opening_balance']
    })
    
    # Add statement data
    for entry in data['statement_data']:
        statement_rows.append({
            'Date': entry['date'].strftime('%Y-%m-%d'),
            'Description': entry['description'],
            'Trx No': entry['trx_no'],
            'Debit': entry['debit'] if entry['debit'] else '',
            'Credit': entry['credit'] if entry['credit'] else '',
            'Balance': entry['running_balance']
        })
    
    # Add totals row
    statement_rows.append({
        'Date': '',
        'Description': 'TOTALS',
        'Trx No': '',
        'Debit': data['debit_amount'],
        'Credit': data['credit_amount'],
        'Balance': data['closing_balance']
    })
    
    df = pd.DataFrame(statement_rows)
    
    # Create Excel file in memory
    output = io.BytesIO()
    
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        # Add header info to the first sheet
        header_info = pd.DataFrame([
            ['STATEMENT OF ACCOUNT', ''],
            [f'{data["company"].company_name.upper()}', ''],
            [f'{data["branch"].branch_name}', ''],
            ['', ''],
            ['Account Holder:', data['full_name']],
            ['Account Number:', data['ac_no']],
            ['GL Number:', data['gl_no']],
            ['Period:', f'{data["start_date"].strftime("%B %d, %Y")} to {data["end_date"].strftime("%B %d, %Y")}'],
            ['', '']
        ])
        
        # Write header info first
        header_info.to_excel(writer, sheet_name='Statement', index=False, header=False, startrow=0)
        
        # Write main statement starting after header
        df.to_excel(writer, sheet_name='Statement', index=False, startrow=len(header_info) + 1)
        
        # Write summary sheet
        summary_data = pd.DataFrame({
            'Item': ['Opening Balance', 'Total Debits', 'Total Credits', 'Closing Balance'],
            'Amount': [data['opening_balance'], data['debit_amount'], 
                      data['credit_amount'], data['closing_balance']]
        })
        summary_data.to_excel(writer, sheet_name='Summary', index=False)
    
    output.seek(0)
    
    response = HttpResponse(
        output.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename="statement_{data["ac_no"]}_{datetime.now().strftime("%Y%m%d")}.xlsx"'
    
    return response

def export_statement_csv(data):
    """Export statement as CSV file"""
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="statement_{data["ac_no"]}_{datetime.now().strftime("%Y%m%d")}.csv"'
    
    writer = csv.writer(response)
    
    # Write header info
    writer.writerow(['STATEMENT OF ACCOUNT'])
    writer.writerow([f'{data["company"].company_name.upper()}'])
    writer.writerow([f'{data["branch"].branch_name}'])
    writer.writerow([])
    writer.writerow(['Account Holder:', data['full_name']])
    writer.writerow(['Account Number:', data['ac_no']])
    writer.writerow(['GL Number:', data['gl_no']])
    writer.writerow(['Period:', f'{data["start_date"].strftime("%B %d, %Y")} to {data["end_date"].strftime("%B %d, %Y")}'])
    writer.writerow([])  # Empty row
    
    # Write table headers
    writer.writerow(['Date', 'Description', 'Trx No', 'Debit', 'Credit', 'Balance'])
    
    # Write opening balance
    writer.writerow([
        data['start_date'].strftime('%Y-%m-%d'),
        'Opening Balance',
        '',
        '',
        '',
        f"{data['opening_balance']:,.2f}"
    ])
    
    # Write statement data
    for entry in data['statement_data']:
        writer.writerow([
            entry['date'].strftime('%Y-%m-%d'),
            entry['description'],
            entry['trx_no'],
            f"{entry['debit']:,.2f}" if entry['debit'] else '',
            f"{entry['credit']:,.2f}" if entry['credit'] else '',
            f"{entry['running_balance']:,.2f}"
        ])
    
    # Write totals
    writer.writerow([
        '',
        'TOTALS',
        '',
        f"{data['debit_amount']:,.2f}",
        f"{data['credit_amount']:,.2f}",
        f"{data['closing_balance']:,.2f}"
    ])
    
    return response