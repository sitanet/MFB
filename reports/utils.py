"""
Report Export Utilities
Provides PDF, Word, Excel, and CSV export functionality for all reports
"""
from django.http import HttpResponse
from django.template.loader import render_to_string
from datetime import datetime
from io import BytesIO
import csv


def export_to_pdf(html_content, filename):
    """Export HTML content to PDF using xhtml2pdf"""
    from xhtml2pdf import pisa
    
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf"'
    
    try:
        result = BytesIO()
        pdf = pisa.pisaDocument(BytesIO(html_content.encode("UTF-8")), result)
        
        if not pdf.err:
            response.write(result.getvalue())
        else:
            response = HttpResponse(f'PDF Generation Error: {pdf.err}', content_type='text/html')
    except Exception as e:
        response = HttpResponse(f'PDF Generation Error: {str(e)}', content_type='text/html')
        
    return response


def export_to_word(title, headers, data_rows, filename, company_name='', branch_name='', period=''):
    """Export data to Word document"""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    
    doc = Document()
    
    # Add title
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add company info if provided
    if company_name:
        company_para = doc.add_paragraph()
        company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_run = company_para.add_run(f'{company_name.upper()}')
        company_run.bold = True
        if branch_name:
            company_para.add_run(f'\n{branch_name}')
    
    # Add period if provided
    if period:
        period_para = doc.add_paragraph()
        period_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        period_para.add_run(f'Period: {period}')
    
    doc.add_paragraph('')  # Space
    
    # Create table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Add headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = str(header)
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Add data rows
    for row in data_rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value) if value is not None else '-'
    
    # Add timestamp
    doc.add_paragraph('')
    timestamp_para = doc.add_paragraph(f'Generated on {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
    timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save to response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{filename}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx"'
    
    doc.save(response)
    return response


def export_to_excel(title, headers, data_rows, filename, company_name='', branch_name='', period=''):
    """Export data to Excel file"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    from openpyxl.utils import get_column_letter
    
    wb = Workbook()
    ws = wb.active
    ws.title = 'Report'
    
    # Styles
    header_font = Font(bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='2563EB', end_color='2563EB', fill_type='solid')
    center_align = Alignment(horizontal='center', vertical='center')
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    current_row = 1
    
    # Add title
    ws.merge_cells(f'A{current_row}:{get_column_letter(len(headers))}{current_row}')
    ws[f'A{current_row}'] = title
    ws[f'A{current_row}'].font = Font(bold=True, size=16)
    ws[f'A{current_row}'].alignment = center_align
    current_row += 1
    
    # Add company info
    if company_name:
        ws.merge_cells(f'A{current_row}:{get_column_letter(len(headers))}{current_row}')
        ws[f'A{current_row}'] = f'{company_name.upper()} - {branch_name}'
        ws[f'A{current_row}'].font = Font(bold=True, size=12)
        ws[f'A{current_row}'].alignment = center_align
        current_row += 1
    
    # Add period
    if period:
        ws.merge_cells(f'A{current_row}:{get_column_letter(len(headers))}{current_row}')
        ws[f'A{current_row}'] = f'Period: {period}'
        ws[f'A{current_row}'].alignment = center_align
        current_row += 1
    
    current_row += 1  # Empty row
    
    # Add headers
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=current_row, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center_align
        cell.border = thin_border
    current_row += 1
    
    # Add data rows
    for row in data_rows:
        for col, value in enumerate(row, 1):
            cell = ws.cell(row=current_row, column=col, value=value if value is not None else '-')
            cell.border = thin_border
            # Right-align numeric values
            if isinstance(value, (int, float)):
                cell.alignment = Alignment(horizontal='right')
        current_row += 1
    
    # Auto-adjust column widths
    for col in range(1, len(headers) + 1):
        max_length = 0
        column = get_column_letter(col)
        for cell in ws[column]:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column].width = adjusted_width
    
    # Add timestamp
    current_row += 1
    ws.merge_cells(f'A{current_row}:{get_column_letter(len(headers))}{current_row}')
    ws[f'A{current_row}'] = f'Generated on {datetime.now().strftime("%B %d, %Y at %I:%M %p")}'
    ws[f'A{current_row}'].alignment = center_align
    ws[f'A{current_row}'].font = Font(italic=True, color='666666')
    
    # Save to response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="{filename}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx"'
    
    wb.save(response)
    return response


def export_to_csv(headers, data_rows, filename):
    """Export data to CSV file"""
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="{filename}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.csv"'
    
    writer = csv.writer(response)
    writer.writerow(headers)
    
    for row in data_rows:
        writer.writerow([str(value) if value is not None else '' for value in row])
    
    return response


def get_report_pdf_template(title, headers, data_rows, company_name='', branch_name='', period='', 
                            summary_data=None, footer_text=''):
    """Generate a styled HTML template for PDF export"""
    
    # Generate table rows
    table_rows_html = ''
    for row in data_rows:
        table_rows_html += '<tr>'
        for value in row:
            # Format numbers with commas if numeric
            if isinstance(value, (int, float)):
                display_value = f'₦{value:,.2f}' if isinstance(value, float) else f'{value:,}'
            else:
                display_value = str(value) if value is not None else '-'
            table_rows_html += f'<td>{display_value}</td>'
        table_rows_html += '</tr>'
    
    # Generate headers
    headers_html = ''.join([f'<th>{h}</th>' for h in headers])
    
    # Generate summary section if provided
    summary_html = ''
    if summary_data:
        summary_html = '<div class="summary-section"><h3>Summary</h3><table class="summary-table">'
        for label, value in summary_data.items():
            if isinstance(value, (int, float)):
                display_value = f'₦{value:,.2f}' if isinstance(value, float) else f'{value:,}'
            else:
                display_value = str(value)
            summary_html += f'<tr><td><strong>{label}:</strong></td><td>{display_value}</td></tr>'
        summary_html += '</table></div>'
    
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>{title}</title>
        <style>
            @page {{
                size: A4;
                margin: 1.5cm;
            }}
            body {{
                font-family: Arial, sans-serif;
                font-size: 10pt;
                margin: 0;
                padding: 0;
                color: #333;
            }}
            .header {{
                text-align: center;
                margin-bottom: 20px;
                padding-bottom: 15px;
                border-bottom: 3px solid #f97316;
            }}
            .header h1 {{
                color: #16a34a;
                margin: 0 0 5px 0;
                font-size: 18pt;
            }}
            .header h2 {{
                color: #f97316;
                margin: 0 0 5px 0;
                font-size: 14pt;
            }}
            .header p {{
                margin: 5px 0;
                color: #666;
            }}
            .period {{
                text-align: center;
                margin-bottom: 15px;
                font-weight: bold;
                color: #16a34a;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin-bottom: 15px;
            }}
            th {{
                background-color: #16a34a;
                color: white;
                padding: 8px 5px;
                text-align: left;
                font-size: 9pt;
                border: 1px solid #ccc;
            }}
            td {{
                padding: 6px 5px;
                border: 1px solid #ddd;
                font-size: 9pt;
            }}
            tr:nth-child(even) {{
                background-color: #f9f9f9;
            }}
            tr:hover {{
                background-color: #f0fdf4;
            }}
            .amount {{
                text-align: right;
                font-family: 'Courier New', monospace;
            }}
            .total-row {{
                font-weight: bold;
                background-color: #fff7ed !important;
                border-top: 2px solid #f97316;
            }}
            .summary-section {{
                margin-top: 20px;
                padding: 15px;
                background-color: #f0fdf4;
                border-left: 4px solid #16a34a;
            }}
            .summary-section h3 {{
                margin: 0 0 10px 0;
                color: #16a34a;
            }}
            .summary-table {{
                width: auto;
            }}
            .summary-table td {{
                padding: 5px 15px 5px 0;
                border: none;
            }}
            .footer {{
                margin-top: 20px;
                text-align: center;
                font-size: 8pt;
                color: #666;
                border-top: 1px solid #ddd;
                padding-top: 10px;
            }}
            .text-success {{ color: #16a34a; }}
            .text-danger {{ color: #ef4444; }}
            .text-warning {{ color: #f97316; }}
        </style>
    </head>
    <body>
        <div class="header">
            <h1>{company_name.upper() if company_name else 'FINANCIAL REPORT'}</h1>
            <h2>{title}</h2>
            {f'<p>{branch_name}</p>' if branch_name else ''}
        </div>
        
        {f'<div class="period">Period: {period}</div>' if period else ''}
        
        <table>
            <thead>
                <tr>{headers_html}</tr>
            </thead>
            <tbody>
                {table_rows_html}
            </tbody>
        </table>
        
        {summary_html}
        
        <div class="footer">
            <p>Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}</p>
            {f'<p>{footer_text}</p>' if footer_text else ''}
            <p>FinanceFlex Banking System</p>
        </div>
    </body>
    </html>
    """
    
    return html_content
